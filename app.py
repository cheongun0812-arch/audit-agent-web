import streamlit as st
import streamlit.components.v1 as components
import os
import google.generativeai as genai
from docx import Document
import PyPDF2
from youtube_transcript_api import YouTubeTranscriptApi
import requests
from bs4 import BeautifulSoup
import time
import glob
import tempfile
import hashlib
import hmac
import base64
import html
import json
import datetime
import pytz
import pandas as pd
import re
from urllib.parse import urlparse, parse_qs

import plotly.graph_objects as go
import plotly.express as px

# Plotly: 확대/축소 후 "원점 복원" 가능하도록 모드바 항상 표시
PLOTLY_CONFIG = {
    "displayModeBar": True,
    "displaylogo": False,
    "responsive": True,
    "scrollZoom": False,          # 스크롤로 의도치 않은 확대 방지
    "doubleClick": "reset",       # 더블클릭/더블탭 시 원점 복원
}

# [필수] 구글 시트 라이브러리 체크
try:
    import gspread
    from oauth2client.service_account import ServiceAccountCredentials
except ImportError:
    gspread = None
    ServiceAccountCredentials = None
    st.error("❌ 구글 시트 라이브러리가 없습니다. requirements.txt를 확인하세요.")

# [필수] yt_dlp 라이브러리 체크
try:
    import yt_dlp
except ImportError:
    yt_dlp = None

# ==========================================
# 1. 페이지 설정
# ==========================================
st.set_page_config(
    page_title="SMART WORK AI AGENT",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ==========================================
# 2. 🎨 디자인 테마 (사이드바/토글 강제 표시 포함)
#    + 전체 텍스트 0.2px 증가
#    + ✅ (요청 반영) 자율점검 탭(#audit-tab) 내 Expander 헤더/입력라벨/셀렉트 가독성 강화
# ==========================================
st.markdown("""
<style>
/* 🔥 Expander 제목 가독성 강제 개선 */
details > summary {
    font-size: 1.15rem !important;
    font-weight: 900 !important;
    color: #1565C0 !important;  /* 📜 서약 타이틀과 동일 색상 */
}

/* 펼쳐졌을 때도 동일하게 유지 */
details[open] > summary {
    font-size: 1.15rem !important;
    font-weight: 900 !important;
    color: #1565C0 !important;
}

/* summary 안의 span도 같이 잡아줌 (환경 차이 대응) */
details > summary,
details > summary span,
details[open] > summary,
details[open] > summary span {
    font-size: 1.5rem !important;   /* ← 여기 숫자만 조절 */
    font-weight: 900 !important;
    color: #1565C0 !important;
}

/* ✅ 전체 글자 크기 +0.1px */
html { font-size: 16.2px; }

.stApp { background-color: #F4F6F9; }
[data-testid="stSidebar"] { background-color: #2C3E50; }
[data-testid="stSidebar"] * { color: #FFFFFF !important; }

/* ✅ 사이드바 텍스트 입력의 아이콘(눈/지우기 등)을 항상 검정색으로 */
[data-testid="stSidebar"] div[data-testid="stTextInput"] button,
[data-testid="stSidebar"] div[data-testid="stTextInput"] button:hover,
[data-testid="stSidebar"] div[data-testid="stTextInput"] button:focus,
[data-testid="stSidebar"] div[data-testid="stTextInput"] button:active {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    color: #000000 !important;
    opacity: 1 !important;
}

[data-testid="stSidebar"] div[data-testid="stTextInput"] button svg,
[data-testid="stSidebar"] div[data-testid="stTextInput"] button svg *,
[data-testid="stSidebar"] div[data-testid="stTextInput"] button svg path {
    fill: #000000 !important;
    stroke: #000000 !important;
    opacity: 1 !important;
}

/* aria-label이 환경/언어에 따라 달라도 적용되도록, 패스워드 토글 버튼도 강제 */
div[data-testid="stTextInput"] button[aria-label],
div[data-testid="stTextInput"] button[aria-label] svg,
div[data-testid="stTextInput"] button[aria-label] svg * {
    fill: #000000 !important;
    stroke: #000000 !important;
    color: #000000 !important;
    opacity: 1 !important;
}

.stTextInput input, .stTextArea textarea {
    background-color: #FFFFFF !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    border: 1px solid #BDC3C7 !important;
}

/* ✅ 버튼 스타일 (일반 버튼 + 폼 제출 버튼) */
.stButton > button,
div[data-testid="stFormSubmitButton"] > button {
    background: linear-gradient(to right, #2980B9, #2C3E50) !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 0.6rem 1rem !important;
    font-weight: 800 !important;
    width: 100% !important;
    opacity: 1 !important;
}

/* ✅ disabled여도 텍스트가 흐려지지 않도록 */
.stButton > button:disabled,
div[data-testid="stFormSubmitButton"] > button:disabled {
    background: linear-gradient(to right, #2980B9, #2C3E50) !important;
    color: #FFFFFF !important;
    opacity: 1 !important;
    filter: none !important;
}

/* ✅ 버튼 내부 텍스트/아이콘도 상시 선명 */
.stButton > button *,
div[data-testid="stFormSubmitButton"] > button * {
    color: #FFFFFF !important;
    opacity: 1 !important;
}

/* (서약 우측 카운트다운 표시용) */
.pledge-right {
  display:flex;
  align-items:center;
  justify-content:flex-end;
  gap: 8px;
  font-weight: 900;
  color: #0B5ED7;
  min-width: 90px;
}

/* =========================================================
   ✅ (요청 1,3,4) 자율점검 탭 전용 가독성 강화
   - 다른 탭/영역 영향 최소화: #audit-tab 내부에서만 적용
   ========================================================= */
#audit-tab [data-testid="stExpander"] summary {
    font-weight: 900 !important;
    font-size: 1.12rem !important;
    color: #1565C0 !important;                 /* 📜 타이틀 색상과 동일 */
}
#audit-tab [data-testid="stExpander"] summary * {
    font-weight: 900 !important;
    color: #1565C0 !important;
}

/* 입력 라벨(사번/성명/총괄/본부/단/상세 부서명) 굵게 */
#audit-tab div[data-testid="stTextInput"] label,
#audit-tab div[data-testid="stSelectbox"] label {
    font-weight: 900 !important;
    color: #2C3E50 !important;
}

/* ✅ 메인 화면의 Selectbox(총괄/본부/단) 선택값 가독성 강제 */
section.main div[data-testid="stSelectbox"] div[data-baseweb="select"] {
    font-size: 1.08rem !important;    /* ← 원하면 더 키우세요 */
    font-weight: 900 !important;
}

/* 선택값이 들어있는 실제 박스(콤보박스) */
section.main div[data-testid="stSelectbox"] div[role="combobox"] {
    background: #FFFFFF !important;
    border: 1px solid #90A4AE !important;
}

/* 선택된 텍스트(대부분 span에 들어감) */
section.main div[data-testid="stSelectbox"] div[role="combobox"] span {
    color: #2C3E50 !important;
    font-weight: 900 !important;
    opacity: 1 !important;
}

/* 어떤 환경에서는 input에 값이 들어가므로 같이 처리 */
section.main div[data-testid="stSelectbox"] div[role="combobox"] input {
    color: #2C3E50 !important;
    -webkit-text-fill-color: #2C3E50 !important;
    font-weight: 900 !important;
    opacity: 1 !important;
}

/* 드롭다운 화살표(아이콘)도 선명하게 */
section.main div[data-testid="stSelectbox"] svg,
section.main div[data-testid="stSelectbox"] svg * {
    fill: #2C3E50 !important;
    stroke: #2C3E50 !important;
    opacity: 1 !important;
}

/* 드롭다운 옵션 목록도 굵게 */
div[role="listbox"] * {
    font-weight: 850 !important;
}
/* ✅ 메인 영역 selectbox를 텍스트 입력창처럼 보이게 (흰박스 + 동일 톤) */
section.main div[data-testid="stSelectbox"] div[role="combobox"]{
  background:#FFFFFF !important;
  border:1px solid #CBD5E1 !important;
  border-radius:6px !important;
  min-height: 42px !important;
  box-shadow: none !important;
}

/* ✅ 선택값 텍스트(진하게) */
section.main div[data-testid="stSelectbox"] div[role="combobox"] span{
  color:#2C3E50 !important;
  font-weight: 800 !important;
  opacity: 1 !important;
}

/* ✅ '선택/placeholder'처럼 보이는 텍스트(옅은 회색) */
/* Streamlit/브라우저마다 placeholder가 input에 들어가거나 span으로 들어가서 둘 다 커버 */
section.main div[data-testid="stSelectbox"] div[role="combobox"] input{
  color:#94A3B8 !important;                 /* search box 느낌의 회색 */
  -webkit-text-fill-color:#94A3B8 !important;
  font-weight: 700 !important;
  opacity: 1 !important;
}

/* ✅ 드롭다운 화살표도 선명하게 */
section.main div[data-testid="stSelectbox"] svg,
section.main div[data-testid="stSelectbox"] svg *{
  fill:#64748B !important;
  stroke:#64748B !important;
  opacity:1 !important;
}
</style>
""", unsafe_allow_html=True)

# ✅ PC에서는 사이드바 기본 펼침, 모바일에서는 기본 접힘
st.markdown("""
<script>
(function() {
  const KEY = "__sidebar_autopen_done__";
  const isDesktop = () => (window.innerWidth || 0) >= 900;
  let tries = 0;
  const maxTries = 25;

  function clickToggleIfNeeded() {
    try {
      if (!isDesktop()) return;
      if (window.sessionStorage.getItem(KEY) === "1") return;

      const doc = window.parent?.document || document;
      const candidates = [
        '[data-testid="stSidebarCollapsedControl"] button',
        '[data-testid="stSidebarCollapsedControl"]',
        'button[title="Open sidebar"]',
        'button[aria-label="Open sidebar"]'
      ];

      for (const sel of candidates) {
        const el = doc.querySelector(sel);
        if (el) {
          el.click();
          window.sessionStorage.setItem(KEY, "1");
          return;
        }
      }
    } catch (e) {}
  }

  const timer = setInterval(() => {
    tries += 1;
    clickToggleIfNeeded();
    if (tries >= maxTries) clearInterval(timer);
  }, 250);
})();
</script>
""", unsafe_allow_html=True)

# ==========================================
# 3. 로그인 및 세션 관리
# ==========================================
def _set_query_param_key(clean_key: str) -> None:
    encoded_key = base64.b64encode(clean_key.encode()).decode()
    try:
        st.query_params["k"] = encoded_key
    except Exception:
        st.experimental_set_query_params(k=encoded_key)

def _clear_query_params() -> None:
    try:
        st.query_params.clear()
    except Exception:
        st.experimental_set_query_params()

def _validate_and_store_key(clean_key: str) -> None:
    genai.configure(api_key=clean_key)
    list(genai.list_models())
    st.session_state["api_key"] = clean_key
    st.session_state["login_error"] = None
    _set_query_param_key(clean_key)

def try_login_from_session_key(key_name: str) -> None:
    raw_key = st.session_state.get(key_name, "")
    clean_key = "".join(str(raw_key).split())
    if not clean_key:
        st.session_state["login_error"] = "⚠️ 키를 입력해주세요."
        return
    try:
        _validate_and_store_key(clean_key)
    except Exception as e:
        st.session_state["login_error"] = f"❌ 인증 실패: {e}"

def perform_logout():
    st.session_state["logout_anim"] = True

# ==========================================
# 4. 자동 로그인 복구 (URL 파라미터)
# ==========================================
if "api_key" not in st.session_state:
    try:
        qp = st.query_params
        if "k" in qp:
            k_val = qp["k"] if isinstance(qp["k"], str) else qp["k"][0]
            restored_key = base64.b64decode(k_val).decode("utf-8")
            _validate_and_store_key(restored_key)
            st.toast("🔄 세션이 복구되었습니다.", icon="✨")
            st.rerun()
    except Exception:
        pass

# ==========================================
# 5. 사이드바 (로그인/로그아웃)
# ==========================================
with st.sidebar:
    st.markdown("### 🏛️ Control Center")
    st.markdown("---")

    if "api_key" not in st.session_state:
        with st.form(key="login_form"):
            st.markdown("<h4 style='color:white;'>🔐 Access Key</h4>", unsafe_allow_html=True)
            st.text_input(
                "Key",
                type="password",
                placeholder="API 키를 입력해 주세요",
                label_visibility="collapsed",
                key="login_input_key",
            )
            st.form_submit_button(
                label="시스템 접속 (Login)",
                on_click=try_login_from_session_key,
                args=("login_input_key",),
                use_container_width=True,
            )

        if st.session_state.get("login_error"):
            st.error(st.session_state["login_error"])
    else:
        st.success("🟢 정상 가동 중")
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("로그아웃 (Logout)", type="primary", use_container_width=True):
            perform_logout()
            st.rerun()

    st.markdown("---")
    st.markdown(
        "<div style='color:white; text-align:center; font-size:12px; opacity:0.8;'>ktMOS북부 Smart Work AI Solution © 2026<br>Engine: Gemini 2.5 / Search Grounding Ready</div>",
        unsafe_allow_html=True,
    )

# ==========================================
# 7. 로그아웃 애니메이션
# ==========================================
if st.session_state.get("logout_anim"):
    st.markdown("""
<div style="background:#0B1B2B; padding:44px 26px; border-radius:18px; text-align:center; border:1px solid rgba(255,255,255,0.12);">
  <div style="font-size: 78px; margin-bottom: 12px; line-height:1.1;">🎆✨</div>
  <div style="font-size: 22px; font-weight: 900; color: #FFFFFF; margin-bottom: 8px;">새해 복 많이 받으세요!</div>
  <div style="font-size: 15px; color: rgba(255,255,255,0.85); line-height: 1.55;">
    올해도 건강과 행운이 가득하시길 바랍니다.<br>
    안전하게 로그아웃되었습니다.
  </div>
  <div style="margin-top:18px; font-size: 12px; color: rgba(255,255,255,0.65);">
    ktMOS북부 Smart Work AI Solution © 2026
  </div>
</div>
""", unsafe_allow_html=True)
    time.sleep(3.0)
    _clear_query_params()
    st.session_state.clear()
    st.rerun()

# ==========================================
# 8. 핵심 기능 함수 (구글시트, AI, 파일처리)
# ==========================================
@st.cache_resource
def init_google_sheet_connection():
    if gspread is None or ServiceAccountCredentials is None:
        return None
    try:
        scope = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        creds = ServiceAccountCredentials.from_json_keyfile_dict(st.secrets["gcp_service_account"], scope)
        return gspread.authorize(creds)
    except Exception:
        return None

def _korea_now():
    try:
        kst = pytz.timezone("Asia/Seoul")
        return datetime.datetime.now(kst)
    except Exception:
        return datetime.datetime.now()

def _campaign_key(dt: datetime.datetime) -> str:
    return f"{dt.year}-{dt.month:02d}"

def _ensure_campaign_config_sheet(spreadsheet):
    try:
        ws = spreadsheet.worksheet("Campaign_Config")
        return ws
    except Exception:
        ws = spreadsheet.add_worksheet(title="Campaign_Config", rows=200, cols=10)
        ws.append_row(["campaign_key", "title", "sheet_name", "start_date"])
        return ws

def _default_campaign_title(dt: datetime.datetime) -> str:
    if dt.month == 1:
        return "1월 자율점검(윤리경영원칙 실천지침 실천 서약)"
    return f"{dt.month}월 자율점검(윤리경영원칙 실천지침 실천서약)"

def _default_campaign_sheet_name(dt: datetime.datetime, spreadsheet=None) -> str:
    if spreadsheet is not None and dt.year == 2026 and dt.month == 1:
        try:
            spreadsheet.worksheet("2026_윤리경영_실천서약")
            return "2026_윤리경영_실천서약"
        except Exception:
            pass
    return f"{dt.year}_{dt.month:02d}_자율점검"

def get_current_campaign_info(spreadsheet, now_dt: datetime.datetime | None = None) -> dict:
    now_dt = now_dt or _korea_now()
    key = _campaign_key(now_dt)
    cfg_ws = _ensure_campaign_config_sheet(spreadsheet)
    records = cfg_ws.get_all_records()
    for r in records:
        if str(r.get("campaign_key", "")).strip() == key:
            title = str(r.get("title") or "").strip() or _default_campaign_title(now_dt)
            sheet_name = str(r.get("sheet_name") or "").strip() or _default_campaign_sheet_name(now_dt, spreadsheet)
            start_date = str(r.get("start_date") or "").strip()
            return {"key": key, "title": title, "sheet_name": sheet_name, "start_date": start_date}

    title = _default_campaign_title(now_dt)
    sheet_name = _default_campaign_sheet_name(now_dt, spreadsheet)
    start_date = now_dt.strftime("%Y.%m.%d")
    cfg_ws.append_row([key, title, sheet_name, start_date])
    return {"key": key, "title": title, "sheet_name": sheet_name, "start_date": start_date}

def set_current_campaign_info(spreadsheet, title: str | None = None, sheet_name: str | None = None, now_dt: datetime.datetime | None = None) -> dict:
    now_dt = now_dt or _korea_now()
    key = _campaign_key(now_dt)
    cfg_ws = _ensure_campaign_config_sheet(spreadsheet)
    all_rows = cfg_ws.get_all_values()
    row_idx = None
    for i in range(2, len(all_rows) + 1):
        if len(all_rows[i-1]) >= 1 and str(all_rows[i-1][0]).strip() == key:
            row_idx = i
            break
    if row_idx is None:
        _ = get_current_campaign_info(spreadsheet, now_dt)
        row_idx = len(all_rows) + 1

    cur = get_current_campaign_info(spreadsheet, now_dt)
    new_title = (title or cur["title"]).strip()
    new_sheet = (sheet_name or cur["sheet_name"]).strip()
    new_start = cur.get("start_date") or now_dt.strftime("%Y.%m.%d")
    cfg_ws.update(f"B{row_idx}:D{row_idx}", [[new_title, new_sheet, new_start]])
    return {"key": key, "title": new_title, "sheet_name": new_sheet, "start_date": new_start}

def save_audit_result(emp_id, name, unit, dept, answer, sheet_name):
    client = init_google_sheet_connection()
    if not client:
        return False, "구글 시트 연결 실패 (Secrets 확인)"
    try:
        spreadsheet = client.open("Audit_Result_2026")
        try:
            sheet = spreadsheet.worksheet(sheet_name)
        except Exception:
            sheet = spreadsheet.add_worksheet(title=sheet_name, rows=2000, cols=10)
            sheet.append_row(["저장시간", "사번", "성명", "총괄/본부/단", "부서", "답변", "비고"])

        # ==========================================
        # ✅ 중복 검증 로직 개선 (사번 + 성명 조합)
        # ==========================================
        all_records = sheet.get_all_records()
        emp_id_str = str(emp_id).strip()
        name_str = str(name).strip()

        for record in all_records:
            # 시트의 사번과 성명 데이터를 가져옴
            existing_emp_id = str(record.get("사번", "")).strip()
            existing_name = str(record.get("성명", "")).strip()

            if emp_id_str == "00000000":
                # 예외 사번(00000000)인 경우: 사번과 성명이 모두 같아야 중복
                if existing_emp_id == "00000000" and existing_name == name_str:
                    return False, f"'{name_str}'님은 이미 '00000000' 사번으로 참여하셨습니다."
            else:
                # 일반 사번인 경우: 사번만 같아도 중복 처리
                if existing_emp_id == emp_id_str:
                    return False, f"사번 {emp_id_str}은(는) 이미 참여한 기록이 있습니다."
        # ==========================================

        korea_tz = pytz.timezone("Asia/Seoul")
        now = datetime.datetime.now(korea_tz).strftime("%Y-%m-%d %H:%M:%S")
        sheet.append_row([now, emp_id, name, unit, dept, answer, "완료"])
        return True, "성공"
    except Exception as e:
        return False, str(e)

def _clean_model_name(model_name: str) -> str:
    """Gemini SDK가 반환하는 'models/...' 형식과 순수 모델명을 모두 안전하게 처리합니다."""
    return str(model_name or "").replace("models/", "").strip()


def _select_available_model(task: str = "balanced") -> str:
    """업무 성격별 우선 모델을 선택하되, 계정에서 지원하지 않으면 자동 fallback 합니다."""
    task = (task or "balanced").lower()
    preference_map = {
        "legal": ["gemini-2.5-pro", "gemini-2.0-flash", "gemini-1.5-pro", "gemini-1.5-flash"],
        "report": ["gemini-2.5-pro", "gemini-2.0-flash", "gemini-1.5-pro", "gemini-1.5-flash"],
        "summary": ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"],
        "chat": ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"],
        "balanced": ["gemini-2.5-flash", "gemini-2.5-pro", "gemini-2.0-flash", "gemini-1.5-pro", "gemini-1.5-flash"],
    }
    preferred = preference_map.get(task, preference_map["balanced"])

    try:
        available_models = [
            _clean_model_name(m.name)
            for m in genai.list_models()
            if "generateContent" in getattr(m, "supported_generation_methods", [])
        ]
        for target in preferred:
            for model_name in available_models:
                if target in model_name:
                    return model_name
        if available_models:
            return available_models[0]
    except Exception:
        # list_models 실패 시에도 아래 기본값으로 시도합니다.
        pass

    return preferred[-1]


def get_model(task: str = "balanced", temperature: float = 0.2):
    """기존 get_model 호환성을 유지하면서 최신 Gemini 모델을 우선 사용합니다."""
    if "api_key" in st.session_state:
        genai.configure(api_key=st.session_state["api_key"])

    model_name = _select_available_model(task)
    try:
        return genai.GenerativeModel(
            model_name,
            generation_config={
                "temperature": temperature,
                "top_p": 0.9,
                "max_output_tokens": 8192,
            },
        )
    except Exception:
        return genai.GenerativeModel("gemini-1.5-flash")


def _extract_response_text(response) -> str:
    """Gemini 응답 객체에서 텍스트를 최대한 안전하게 추출합니다."""
    try:
        return response.text or ""
    except Exception:
        pass
    try:
        parts = response.candidates[0].content.parts
        return "\n".join(str(getattr(p, "text", "")) for p in parts if getattr(p, "text", ""))
    except Exception:
        return ""


def _extract_grounding_sources(response) -> list[dict]:
    """Google Search Grounding 결과의 출처를 사용자에게 보여주기 좋은 형태로 정리합니다."""
    sources = []
    try:
        metadata = getattr(response.candidates[0], "grounding_metadata", None)
        chunks = getattr(metadata, "grounding_chunks", []) if metadata else []
        for chunk in chunks:
            web = getattr(chunk, "web", None)
            if not web:
                continue
            title = getattr(web, "title", "") or "출처"
            uri = getattr(web, "uri", "") or ""
            if uri and not any(s.get("uri") == uri for s in sources):
                sources.append({"title": title, "uri": uri})
    except Exception:
        pass
    return sources[:8]


def generate_ai_response(content, task: str = "balanced", use_search: bool = False, temperature: float = 0.2):
    """공통 AI 호출 함수: 검색 보강 요청 시 Google Search Grounding을 우선 시도하고 실패하면 일반 생성으로 fallback 합니다."""
    model = get_model(task=task, temperature=temperature)
    search_warning = None

    if use_search:
        # google-generativeai 구버전/신버전 호환을 위해 두 가지 표기를 순차 시도합니다.
        for tool_name in ("google_search_retrieval", "google_search"):
            try:
                response = model.generate_content(content, tools=tool_name)
                return response, True, None
            except Exception as e:
                search_warning = str(e)

    response = model.generate_content(content)
    return response, False, search_warning


def render_ai_response(response, grounded: bool = False, warning: str | None = None) -> None:
    """AI 결과와 검색 출처를 공통 UI로 출력합니다."""
    answer = _extract_response_text(response)
    if answer:
        st.markdown(answer)
    else:
        st.warning("AI 응답 텍스트를 추출하지 못했습니다. 입력 자료나 모델 응답 제한 여부를 확인해 주세요.")

    sources = _extract_grounding_sources(response)
    if grounded and sources:
        with st.expander("🔎 검색 기반 참고 출처", expanded=False):
            for i, src in enumerate(sources, 1):
                st.markdown(f"{i}. [{src['title']}]({src['uri']})")
    elif warning:
        st.caption("ℹ️ 검색 보강 호출이 실패하여 일반 AI 분석으로 대체되었습니다. 패키지 버전 또는 API 권한을 확인해 주세요.")


def truncate_text(text: str, limit: int = 45000) -> str:
    text = str(text or "")
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[※ 입력 자료가 길어 앞부분 기준으로 일부만 반영되었습니다.]"


def read_file(uploaded_file):
    content = ""
    try:
        name = uploaded_file.name.lower()
        if name.endswith(".txt"):
            raw = uploaded_file.getvalue()
            for enc in ("utf-8", "cp949", "euc-kr"):
                try:
                    content = raw.decode(enc)
                    break
                except Exception:
                    continue
        elif name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(uploaded_file)
            for idx, page in enumerate(reader.pages, 1):
                page_text = page.extract_text() or ""
                content += f"\n\n[Page {idx}]\n{page_text}"
        elif name.endswith(".docx"):
            doc = Document(uploaded_file)
            content = "\n".join([para.text for para in doc.paragraphs if para.text])
        elif name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
            content = df.head(200).to_markdown(index=False)
        elif name.endswith((".xlsx", ".xls")):
            df = pd.read_excel(uploaded_file)
            content = df.head(200).to_markdown(index=False)
    except Exception:
        return None
    return content.strip() if content else None


def read_multiple_files(files, max_each: int = 12000) -> str:
    """여러 파일을 감사보고서/법률검토 프롬프트에 안전하게 합칩니다."""
    chunks = []
    for f in files or []:
        body = read_file(f)
        if body:
            chunks.append(f"\n\n===== 파일: {getattr(f, 'name', 'uploaded')} =====\n{truncate_text(body, max_each)}")
        else:
            chunks.append(f"\n\n===== 파일: {getattr(f, 'name', 'uploaded')} =====\n[텍스트 추출 실패 또는 지원되지 않는 형식]")
    return "\n".join(chunks).strip()


def process_media_file(uploaded_file):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name

        st.toast("🤖 AI에게 분석 자료를 전달하고 있습니다...", icon="📂")
        myfile = genai.upload_file(tmp_path)
        with st.spinner("🎧 AI가 데이터를 분석하고 있습니다..."):
            while myfile.state.name == "PROCESSING":
                time.sleep(2)
                myfile = genai.get_file(myfile.name)

        os.remove(tmp_path)
        if myfile.state.name == "FAILED":
            return None
        return myfile
    except Exception:
        return None


def download_and_upload_youtube_audio(url):
    if yt_dlp is None:
        return None
    try:
        ydl_opts = {"format": "bestaudio/best", "outtmpl": "temp_audio.%(ext)s", "quiet": True}
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])
        audio_files = glob.glob("temp_audio.*")
        if not audio_files:
            return None
        audio_path = audio_files[0]
        myfile = genai.upload_file(audio_path)
        with st.spinner("🎧 유튜브 분석 중..."):
            while myfile.state.name == "PROCESSING":
                time.sleep(2)
                myfile = genai.get_file(myfile.name)
        os.remove(audio_path)
        return myfile
    except Exception:
        return None


def extract_youtube_id(url: str) -> str | None:
    """watch?v=, youtu.be, shorts, embed 형식까지 처리합니다."""
    try:
        parsed = urlparse(url)
        host = parsed.netloc.lower()
        path = parsed.path.strip("/")
        if "youtube.com" in host:
            qs = parse_qs(parsed.query)
            if qs.get("v"):
                return qs["v"][0]
            parts = path.split("/")
            if parts and parts[0] in {"shorts", "embed", "live"} and len(parts) > 1:
                return parts[1]
        if "youtu.be" in host and path:
            return path.split("/")[0]
    except Exception:
        pass
    m = re.search(r"(?:v=|youtu\.be/|shorts/|embed/)([A-Za-z0-9_-]{8,})", url or "")
    return m.group(1) if m else None


def get_youtube_transcript(url):
    try:
        video_id = extract_youtube_id(url)
        if not video_id:
            return None
        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=["ko", "en"])
        return " ".join([t.get("text", "") for t in transcript])
    except Exception:
        return None


def get_web_content(url):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120 Safari/537.36"
        }
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
            tag.decompose()
        title = soup.title.get_text(" ", strip=True) if soup.title else ""
        meta_desc = ""
        desc = soup.find("meta", attrs={"name": "description"})
        if desc and desc.get("content"):
            meta_desc = desc.get("content", "")
        text = soup.get_text("\n", strip=True)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return f"URL: {url}\n제목: {title}\n설명: {meta_desc}\n\n본문:\n{text[:25000]}"
    except Exception:
        return None


def build_legal_review_prompt(content: str, analysis_depth: str, doc_type: str, focus_area: str, company_position: str) -> str:
    return f"""[역할]
당신은 대한민국 기업 감사실을 보조하는 법률·컴플라이언스 검토 전문가입니다.

[검토 대상]
- 문서 유형: {doc_type}
- 회사 입장: {company_position}
- 분석 수준: {analysis_depth}
- 중점 검토: {focus_area}

[작성 원칙]
1. 업로드 문서에서 확인되는 사실과 AI의 법률적 판단/추정을 명확히 구분하세요.
2. 대한민국 법령·판례·공정거래/하도급/개인정보/근로관계 등 관련 기준을 고려하되, 근거가 불충분하면 '추가 확인 필요'로 표시하세요.
3. 실무자가 바로 사용할 수 있도록 '위험 조항', '리스크 이유', '개선 문안'을 구체적으로 제시하세요.
4. 과도하게 단정하지 말고, 감사·법무 검토 문서에 적합한 객관적 문체로 작성하세요.

[출력 형식]
## 1. 핵심 결론
- 즉시 수정 필요 / 협의 필요 / 수용 가능 항목을 요약

## 2. 조항별 리스크 검토표
| 우선순위 | 원문 또는 쟁점 | 리스크 등급 | 문제점 | 관련 법령·판례·가이드라인 방향 | 개선 의견 |

## 3. 상대방에게 제시할 수정 문안
- 조항별로 대체 문구 작성

## 4. 내부 검토 메모
- 감사실/법무/사업부가 추가 확인해야 할 사항

## 5. 한계 및 추가 확인 필요사항
- 문서에 없는 사실, 최신 법령 확인 필요사항, 외부 변호사 검토 필요사항

[입력 문서]
{truncate_text(content, 55000)}
"""


def build_audit_report_prompt(mode: str, case_title: str, case_scope: str, report_tone: str, materials: str, regulations_text: str, refs_text: str) -> str:
    if "초안" in mode:
        task = "감사보고서 초안을 생성"
        output = "사건개요, 확인자료, 주요 사실관계, 쟁점, 판단, 리스크, 조치의견, 후속관리 항목을 포함한 공식 감사보고서 초안"
    else:
        task = "감사보고서 초안을 검증·교정"
        output = "오탈자·논리비약·근거부족·표현위험·형식오류를 지적하고, 개선본과 수정 사유표를 제시"

    return f"""[역할]
당신은 기업 감사실의 감사보고서 품질관리 담당자입니다.

[작업]
- 작업 모드: {mode}
- 수행 작업: {task}
- 사건명: {case_title}
- 문서 톤: {report_tone}

[사건 개요]
{case_scope}

[작성 원칙]
1. 사실관계, 판단, 의견을 구분하세요.
2. 업로드 자료에 없는 사실은 새로 만들지 말고 '자료상 확인 불가'로 표시하세요.
3. 피조사자·관련자의 명예, 개인정보, 노동관계 리스크를 고려하여 표현을 중립적으로 조정하세요.
4. 내부 결재문서로 활용 가능한 수준의 문장으로 작성하세요.
5. 불리한 단정 표현은 '확인됨/확인 필요/소명 필요/자료상 불명확' 등으로 정리하세요.

[출력 형식]
## 1. {output}
## 2. 핵심 쟁점 및 증거 연결표
| 쟁점 | 확인자료 | 판단 가능 수준 | 보완 필요자료 |
## 3. 표현 리스크 점검
| 문장/표현 | 리스크 | 권장 표현 |
## 4. 후속 조치안
## 5. 추가 확인 필요사항

[감사 자료]
{truncate_text(materials, 50000)}

[회사 규정/판단 기준]
{truncate_text(regulations_text, 25000)}

[참고 보고서 형식]
{truncate_text(refs_text, 15000)}
"""


def build_chat_prompt(user_input: str, history: list[dict], mode: str) -> str:
    recent = history[-10:] if history else []
    history_text = "\n".join([f"{m.get('role')}: {m.get('content')}" for m in recent])
    return f"""[시스템 역할]
당신은 대한민국 기업 감사실을 지원하는 Professional Legal & Audit Assistant입니다.
감사, 컴플라이언스, 계약 검토, 개인정보, 하도급, 공정거래, 직장 내 괴롭힘, 내부통제 이슈에 대해 실무형으로 답변합니다.

[응답 원칙]
1. 결론을 먼저 제시하고, 근거와 실무 조치 순서로 설명하세요.
2. 법률·판례·행정해석이 필요한 질문은 '확인된 근거'와 '추가 확인 필요'를 구분하세요.
3. 단정이 위험한 사안은 리스크와 방어 전략을 함께 제시하세요.
4. 최종 법률 판단은 사내 법무/외부 변호사 검토가 필요하다는 점을 짧게 고지하세요.
5. 답변 마지막에는 '추가 확인 필요사항'을 1~2문장 포함하세요.

[현재 모드]
{mode}

[최근 대화]
{history_text}

[사용자 질문]
{user_input}
"""


def build_summary_prompt(summary_mode: str, output_style: str, source_hint: str, body_text: str) -> str:
    return f"""[역할]
당신은 기업 감사실과 컴플라이언스 부서를 위한 스마트 브리핑 분석가입니다.

[요약 대상]
- 요약 모드: {summary_mode}
- 출력 방식: {output_style}
- 출처/입력: {source_hint}

[작성 원칙]
1. 사실, 주장, 의견, 추정을 구분하세요.
2. 날짜·기관·당사자·수치가 있으면 빠뜨리지 마세요.
3. 회사 업무상 영향, 컴플라이언스 리스크, 후속 조치 필요사항을 별도로 정리하세요.
4. 원문에서 확인되지 않는 내용은 만들지 말고 '원문상 확인 불가'로 표시하세요.

[출력 형식]
## 1. 5줄 핵심 요약
## 2. 상세 내용
## 3. 업무상 의미/리스크
## 4. 후속 조치 체크리스트
## 5. 원문 한계 및 추가 확인사항

[입력 내용]
{truncate_text(body_text, 55000)}
"""

# ==========================================
# ✅ (요청 2) 사번 검증 유틸
# ==========================================
def validate_emp_id(emp_id: str) -> tuple[bool, str]:
    """
    규칙:
    - 기본: 8자리 숫자, '10'으로 시작 (10******)
    - 예외: 사번 미부여자는 '00000000' 허용(제출 가능)
    """
    s = (emp_id or "").strip()

    if not s:
        return False, "⚠️ 사번을 입력해 주세요. (사번 미부여 시 '00000000')"

    # ✅ 예외 허용: 사번 미부여
    if s == "00000000":
        return True, "ℹ️ 사번 미부여: '00000000'으로 제출됩니다. 제출 후 관리자에게 연락해 주세요."

    # 기본 형식 체크
    if (len(s) != 8) or (not s.isdigit()):
        return False, "⚠️ 사번이 8자리 숫자가 아닙니다. 사번을 정확히 입력했는지 다시 확인해 주세요."

    # 기본 규칙: 10으로 시작
    if not s.startswith("10"):
        return False, "⚠️ 사번을 정확히 입력했는지 확인해 주세요. 사번이 '10********'이 아니라면 '00000000'을 입력해 제출 후 관리자에게 연락해 주세요."

    return True, ""


# ==========================================
# ✅ 2026년 6월 컴플라이언스 인식제고 교육 저장 유틸
#    - 기존 윤리경영 실천서약 저장 로직과 분리
#    - Google Sheet: Audit_Result_2026 / 2026_06_컴플라이언스_인식제고교육
# ==========================================
JUNE_TRAINING_SHEET_NAME = "2026_06_컴플라이언스_인식제고교육"
JUNE_TRAINING_HEADERS = [
    "저장시간", "수료ID", "사번", "성명", "총괄/본부/단", "부서",
    "교육대상", "Theme1_청렴공정_확인", "Theme1_점수", "Theme2_협력정보_확인", "Theme2_점수",
    "이벤트퀴즈_선택", "이벤트퀴즈_정답여부", "퀴즈점수", "참여점수", "최종점수", "수료상태", "이벤트추첨대상", "비고"
]


def save_june_compliance_training_result(record: dict) -> tuple[bool, str]:
    """6월 컴플라이언스 인식제고 교육 수료 내역을 Google Sheet에 저장합니다."""
    client = init_google_sheet_connection()
    if not client:
        return False, "구글 시트 연결 실패 (Secrets 확인)"

    try:
        spreadsheet = client.open("Audit_Result_2026")
        try:
            sheet = spreadsheet.worksheet(JUNE_TRAINING_SHEET_NAME)
        except Exception:
            sheet = spreadsheet.add_worksheet(title=JUNE_TRAINING_SHEET_NAME, rows=3000, cols=len(JUNE_TRAINING_HEADERS) + 2)
            sheet.append_row(JUNE_TRAINING_HEADERS)

        # ✅ 동시 수료 제출 안정화: 저장 직전 전체 데이터 읽기(get_all_records)를 하지 않습니다.
        # 1,000명 동시 접속 상황에서 읽기 요청 한도(429)를 유발하던 중복검사는 운영 종료 후 Google Sheet에서 사번 기준으로 확인합니다.
        emp_id_str = str(record.get("사번", "")).strip()
        name_str = str(record.get("성명", "")).strip()
        dept_str = str(record.get("부서", "")).strip()

        now = _korea_now().strftime("%Y-%m-%d %H:%M:%S")
        completion_seed = f"{now}|{emp_id_str}|{name_str}|{dept_str}|2026-06-compliance"
        completion_id = hashlib.sha256(completion_seed.encode("utf-8")).hexdigest()[:12]

        row = [
            now,
            completion_id,
            record.get("사번", ""),
            record.get("성명", ""),
            record.get("총괄/본부/단", ""),
            record.get("부서", ""),
            record.get("교육대상", "전 임직원"),
            record.get("Theme1_청렴공정_확인", "완료"),
            record.get("Theme1_점수", 0),
            record.get("Theme2_협력정보_확인", "완료"),
            record.get("Theme2_점수", 0),
            record.get("이벤트퀴즈_선택", ""),
            record.get("이벤트퀴즈_정답여부", ""),
            record.get("퀴즈점수", 0),
            record.get("참여점수", 0),
            record.get("최종점수", 0),
            record.get("수료상태", "수료"),
            record.get("이벤트추첨대상", "대상"),
            record.get("비고", ""),
        ]
        last_error = None
        for attempt in range(5):
            try:
                sheet.append_row(row, value_input_option="USER_ENTERED")
                return True, "6월 컴플라이언스 인식제고 교육 수료 내역이 저장되었습니다."
            except Exception as append_error:
                last_error = append_error
                msg = str(append_error)
                if "429" in msg or "Quota exceeded" in msg or "RESOURCE_EXHAUSTED" in msg:
                    time.sleep(1.2 * (attempt + 1))
                    continue
                raise
        return False, f"Google Sheet 저장 요청이 일시적으로 집중되어 저장하지 못했습니다. 잠시 후 현재 화면에서 다시 제출해 주세요. ({last_error})"
    except Exception as e:
        return False, str(e)

# ==========================================
# 8-2. 국사 전원시설 정밀점검 및 Google Sheets 저장
#      - 기존 Google 서비스 계정/스프레드시트 연결 재사용
#      - 점검 1건을 Google Sheet 1행으로 저장
# ==========================================
POWER_INSPECTION_SPREADSHEET_NAME = "Audit_Result_2026"
POWER_INSPECTION_SHEET_NAME = "국사_전원시설_정밀점검"

POWER_REGION_DATA = {'1권역 · 파주·문산·동두천 등': {'담당자': ['이철순', '김수창'],
                       '모국_국소': {'동두천': ['은현', '신산', '상수'],
                                 '문산': ['문산', '파주', '마산', '마정', '통일촌(N3211)', '장현', '웅담', '적성', '파평', '당동상가BBH', '봉암1리마을회관BBH'],
                                 '법원리': ['법원리', '금파리마을회관BBH', '어유지리BBH'],
                                 '연천': ['연천', '대광', '삼곳', '내산', '고문'],
                                 '일산': ['중산BBH'],
                                 '전곡': ['원당', '백학', '동이', '궁평', '왕림', '초성', '동중', '진상', '북삼', '늘목', '양원', '대전', '전곡', '학곡'],
                                 '파주': ['영장',
                                        '장곡',
                                        '위전',
                                        '법흥',
                                        '문발',
                                        '용미',
                                        '발랑',
                                        '파주 연다산(N3218)',
                                        '광탄',
                                        '탄현',
                                        '운정',
                                        '(파주)마이프라자1층BBH',
                                        '분수3리마을회관BBH',
                                        '새말BBH',
                                        '토우프라자BBH',
                                        '통일프라자BBH',
                                        '한라비발디상가BBH',
                                        '대성동마을회관BBH']}},
 '2권역 · 고양·덕양·의정부 등': {'담당자': ['소순고', '정청운'],
                       '모국_국소': {'고양IBS고양BBH': ['백석12블럭BBH', '마두21블럭BBH', '마두25블럭BBH', '백석8블럭BBH', '풍동에이스타워BBH'],
                                 '능곡': ['소만8단지(소만풍림8단지)BBH', '햇빛주공22단지BBH', '화정동상가(별빛건영10단지) BBH'],
                                 '덕양': ['덕양', '덕양 벽제(N3055)', '고양', '달빛3단지상가지하BBH', '상곡BBH'],
                                 '서대문': ['서대문(최적화분기국사)', '화전'],
                                 '신촌': ['신촌분기국사', '망원동BBH#1'],
                                 '아현': ['마포분기국사'],
                                 '용산': ['청파3가BBH'],
                                 '은평': ['삼송'],
                                 '의정부': ['장흥', '송추', '백석', '광적', '의정부 덕도(N3162)', '석우', '삼하', '비암'],
                                 '일산': ['송포',
                                        '성석',
                                        '고봉산중계소',
                                        '백마5단지상가BBH(N3173)',
                                        '장항동BBH',
                                        '북일산최적화분기국사(북일산)',
                                        '고양BBH(N3170)',
                                        '가좌BBH',
                                        '백석13블럭BBH',
                                        '백석7블럭BBH',
                                        '정발BBH',
                                        '일산(EBS)']}},
 '3권역 · 광화문·중앙·광진 등': {'담당자': ['김태수', '이학원'],
                       '모국_국소': {'광진': ['구의동BBH', '중곡동BBH'],
                                 '광화문': ['독립문통신구BBH(N995)', '무교동BBH', '종로5가통신구-1BBH(통신구내)'],
                                 '노원': ['공릉최적화분기국사'],
                                 '도봉': ['방학최적화분기국사', '수유3국사(수유동269-16단독주택)BBH'],
                                 '방학': ['행운빌라BBH'],
                                 '성북': ['안암BBH', '대광BBH', '성북BBH', '보문시장BBH'],
                                 '신내': ['중랑최적화분기국사'],
                                 '용산': ['경찰청'],
                                 '은평': ['홍제최적화분기국사', '평창아파트BBH', '효자BBH(N977)'],
                                 '을지': ['을지로3가BBH', '을지로6가BBH', '을지로7가BBH(통신구내)'],
                                 '중랑': ['망우최적화분기국사(중랑)'],
                                 '중앙': ['BBH(후암동68-6-BBH)', '을지메인통신구', '을지입구B1통신구내 BBH(N944)'],
                                 '청량': ['청량최적화BBH'],
                                 '행당': ['약수역BBH', '동대문최적화BBH(통신구내)']}},
 '4권역 · 동의정부·동두천·철원 등': {'담당자': ['박동희', '신진우'],
                         '모국_국소': {'동두천': ['동두천', '덕정', '덕계', '소요', '광암'],
                                   '동의정부': ['경중앙(통신구내)BBH', '남방', '광사', '자일', '청학', '금오BBH'],
                                   '송우': ['송우', '가산', '내촌', '신팔', '이곡'],
                                   '의정부': ['녹양'],
                                   '철원': ['철원', '동송분기국사', '문혜', '내대', '관전', '장흥', '오지', '지경', '자등', '마현', '양지', '근남', '잠곡', '와수', '도창'],
                                   '청평': ['마일', '봉수', '율길', '임초', '조종', '상판', '대보'],
                                   '퇴계원': ['광릉', '금곡BBH'],
                                   '포천': ['포천',
                                          '자작',
                                          '직두',
                                          '화현',
                                          '일동',
                                          '사직',
                                          '수입',
                                          '장암',
                                          '도평',
                                          '산정',
                                          '영북',
                                          '관인',
                                          '운산',
                                          '창수',
                                          '신북',
                                          '남청산',
                                          '고소성',
                                          '만세교',
                                          '양문',
                                          '대회산']}},
 '5권역 · 가평·남양주·양평 등': {'담당자': ['강만식', '이민우'],
                       '모국_국소': {'가평': ['개곡', '가평', '상색', '산유', '북면', '화악', '백둔', '도대', '적목'],
                                 '남양주': ['남양주', '일패', '답내', '운수', '외방', '호평BBH'],
                                 '덕소': ['덕소', '조안', '송촌', '월문', '팔당'],
                                 '양평': ['양평',
                                        '노문',
                                        '정배',
                                        '목왕',
                                        '서종',
                                        '양수',
                                        '국수',
                                        '강하',
                                        '강상',
                                        '회현',
                                        '개군',
                                        '일신',
                                        '양동',
                                        '계정',
                                        '금왕',
                                        '고송',
                                        '용문',
                                        '신점',
                                        '단월',
                                        '산음',
                                        '명성',
                                        '용두',
                                        '옥천',
                                        '지평',
                                        '용문산중계소'],
                                 '청평': ['청평', '고성', '대성', '미사', '방일', '삼회', '설악', '회곡', '에덴성회BBH'],
                                 '퇴계원': ['퇴계원', '진접', '진건', '오남', '별내']}}}


def _build_power_station_map() -> dict[str, list[str]]:
    """모든 권역의 모국·국소를 합친 하위 호환용 전체 목록입니다."""
    station_map: dict[str, list[str]] = {}
    for region in POWER_REGION_DATA.values():
        for mother, locals_ in region.get("모국_국소", {}).items():
            station_map.setdefault(mother, [])
            for local in locals_:
                if local not in station_map[mother]:
                    station_map[mother].append(local)
    return station_map


POWER_STATION_MAP = _build_power_station_map()
POWER_INSPECTOR_OPTIONS = [
    person
    for region in POWER_REGION_DATA.values()
    for person in region.get("담당자", [])
]
POWER_INSPECTOR_MAJOR_AREA_MAP = {
    person: area
    for area, region in POWER_REGION_DATA.items()
    for person in region.get("담당자", [])
}

# ✅ 현장 기본정보 단축 입력용: 권역별 담당자 2명을 한 묶음으로 표시합니다.
POWER_AREA_INSPECTOR_DISPLAY = {
    area: ", ".join(
        str(person).strip()
        for person in region.get("담당자", [])
        if str(person).strip()
    )
    for area, region in POWER_REGION_DATA.items()
}
POWER_INSPECTOR_GROUP_OPTIONS = [
    display for display in POWER_AREA_INSPECTOR_DISPLAY.values() if display
]
POWER_INSPECTOR_DISPLAY_AREA_MAP = {
    display: area
    for area, display in POWER_AREA_INSPECTOR_DISPLAY.items()
    if display
}


def _build_power_station_search_entries() -> list[dict]:
    """국소명 하나로 권역·담당자·모국·국소를 찾을 수 있는 역색인을 만듭니다."""
    entries: list[dict] = []
    serial = 0
    for area, region in POWER_REGION_DATA.items():
        inspectors = POWER_AREA_INSPECTOR_DISPLAY.get(area, "")
        station_map = region.get("모국_국소", {}) if isinstance(region, dict) else {}
        for mother, locals_ in station_map.items():
            for local in locals_:
                serial += 1
                entries.append({
                    "id": f"power_station_{serial:03d}",
                    "area": str(area).strip(),
                    "inspectors": inspectors,
                    "mother": str(mother).strip(),
                    "local": str(local).strip(),
                })
    return entries


POWER_STATION_SEARCH_ENTRIES = _build_power_station_search_entries()
POWER_STATION_SEARCH_BY_ID = {
    entry["id"]: entry for entry in POWER_STATION_SEARCH_ENTRIES
}



# 담당자에 따라 소속 조를 자동 표시합니다.
# 이름과 조 정보가 추가되면 이 사전만 확장하면 됩니다.
POWER_INSPECTOR_GROUP_MAP = {
    "정청운": "덕양관리조",
    "정철선": "덕양관리조",
    "정철순": "덕양관리조",
    "소순고": "덕양관리조",
    "이철순": "고양관리조",
    "김수창": "고양관리조",
    # 영문 입력 보조
    "JEONGCHEONGWOON": "덕양관리조",
    "JEONGCHEOLSUN": "덕양관리조",
    "SOSOONGO": "덕양관리조",
    "LEECHEOLSOON": "고양관리조",
    "KIMSOOCHANG": "고양관리조",
}


def _normalize_inspector_name(name: str) -> str:
    normalized = re.sub(r"[\s\-_.]", "", str(name or "").strip())
    return normalized.upper() if re.search(r"[A-Za-z]", normalized) else normalized


def _inspector_group_for_name(name: str) -> str:
    return POWER_INSPECTOR_GROUP_MAP.get(_normalize_inspector_name(name), "")


def _major_areas_for_inspector(name: str) -> list[str]:
    normalized = _normalize_inspector_name(name)
    areas: list[str] = []
    for person, area in POWER_INSPECTOR_MAJOR_AREA_MAP.items():
        if _normalize_inspector_name(person) == normalized and area not in areas:
            areas.append(area)
    return areas


def _power_area_station_map(area: str) -> dict[str, list[str]]:
    region = POWER_REGION_DATA.get(str(area or "").strip(), {})
    mapping = region.get("모국_국소", {}) if isinstance(region, dict) else {}
    return mapping if isinstance(mapping, dict) else {}


def _inspectors_for_major_area(area: str) -> list[str]:
    """선택한 권역에 배정된 담당자 2명만 반환합니다."""
    region = POWER_REGION_DATA.get(str(area or "").strip(), {})
    people = region.get("담당자", []) if isinstance(region, dict) else []
    return [str(person).strip() for person in people if str(person).strip()]


def _automatic_inspector_display(area: str) -> str:
    """권역을 선택하면 별도 선택 없이 담당자 2명을 가로로 표시·저장합니다."""
    return POWER_AREA_INSPECTOR_DISPLAY.get(str(area or "").strip(), "")


def _major_area_for_worker_value(worker: str) -> str:
    """개별 담당자 또는 '담당자1, 담당자2' 묶음값에서 권역을 찾습니다."""
    value = str(worker or "").strip()
    if value in POWER_INSPECTOR_DISPLAY_AREA_MAP:
        return POWER_INSPECTOR_DISPLAY_AREA_MAP[value]
    return POWER_INSPECTOR_MAJOR_AREA_MAP.get(value, "권역 선택")


def _power_worker_matches_area(worker: str, area: str) -> bool:
    """과거 개별 담당자 값과 신규 2인 묶음값을 모두 허용합니다."""
    worker_value = str(worker or "").strip()
    area_value = str(area or "").strip()
    if area_value not in POWER_REGION_DATA:
        return False
    if worker_value == _automatic_inspector_display(area_value):
        return True
    return worker_value in _inspectors_for_major_area(area_value)


def _normalize_power_station_search(text: str) -> str:
    return re.sub(r"[\s·._()#\-]", "", str(text or "").strip()).lower()


def _search_power_station_entries(query: str, limit: int = 40) -> list[dict]:
    """국소명을 우선으로 검색하고, 모국/권역명도 보조 검색합니다."""
    normalized = _normalize_power_station_search(query)
    if not normalized:
        return []

    scored: list[tuple[int, str, dict]] = []
    for entry in POWER_STATION_SEARCH_ENTRIES:
        local_key = _normalize_power_station_search(entry.get("local", ""))
        mother_key = _normalize_power_station_search(entry.get("mother", ""))
        area_key = _normalize_power_station_search(entry.get("area", ""))
        if normalized == local_key:
            rank = 0
        elif local_key.startswith(normalized):
            rank = 1
        elif normalized in local_key:
            rank = 2
        elif normalized in mother_key:
            rank = 3
        elif normalized in area_key:
            rank = 4
        else:
            continue
        scored.append((rank, local_key, entry))

    scored.sort(key=lambda item: (item[0], item[1], item[2].get("mother", ""), item[2].get("area", "")))
    return [entry for _, _, entry in scored[:max(1, int(limit))]]


def _power_station_search_label(entry_id: str) -> str:
    entry = POWER_STATION_SEARCH_BY_ID.get(str(entry_id or ""), {})
    if not entry:
        return "검색 결과 없음"
    return (
        f"{entry.get('local', '')}  |  모국 {entry.get('mother', '')}  |  "
        f"담당자 {entry.get('inspectors', '')}  |  {entry.get('area', '')}"
    )


def _apply_power_station_search_entry(entry: dict) -> None:
    """검색으로 확정한 국사의 기본정보를 기존 입력 상태에 안전하게 반영합니다."""
    if not entry:
        return

    _preserve_current_power_measurements()
    selected_area = str(entry.get("area", "")).strip()
    st.session_state["power_worker"] = _automatic_inspector_display(selected_area)
    st.session_state["power_major_area"] = selected_area
    st.session_state["power_inspector_group"] = _inspector_group_for_area(selected_area)
    st.session_state["power_mother"] = str(entry.get("mother", "")).strip()
    st.session_state["power_local"] = str(entry.get("local", "")).strip()
    _clear_power_history_state()
    _mark_power_basic_info_changed()

    # 검색 반영 여부를 별도 보존하여 아래 기본정보 선택값을 자주색으로 강조합니다.
    st.session_state["power_station_search_applied"] = True
    st.session_state["power_station_search_candidates"] = []
    # duplicate radio key는 해당 위젯이 생성된 실행 중에는 직접 변경하지 않습니다.
    st.session_state["power_station_search_status"] = "applied"
    st.session_state["power_station_search_notice"] = (
        f"✅ {entry.get('local', '')} 국사 자동입력 완료 · "
        f"담당자 {entry.get('inspectors', '')} · {selected_area} · "
        f"모국 {entry.get('mother', '')}"
    )


def _run_power_station_search() -> None:
    """확인 버튼/Enter 제출 시 검색합니다. 1건이면 즉시 반영하고, 중복이면 후보만 표시합니다."""
    query = str(st.session_state.get("power_station_search_query", "") or "").strip()
    st.session_state["power_station_search_notice"] = ""
    st.session_state["power_station_search_status"] = ""
    st.session_state["power_station_search_candidates"] = []
    st.session_state["power_station_search_choice"] = ""

    if not query:
        st.session_state["power_station_search_status"] = "empty"
        return

    normalized_query = _normalize_power_station_search(query)
    matches = _search_power_station_entries(query)

    # '송포'처럼 국소명이 정확히 일치하면 부분검색 결과보다 정확일치를 우선합니다.
    exact_matches = [
        entry for entry in matches
        if _normalize_power_station_search(entry.get("local", "")) == normalized_query
    ]
    candidates = exact_matches if exact_matches else matches

    if len(candidates) == 1:
        _apply_power_station_search_entry(candidates[0])
        return

    if len(candidates) > 1:
        candidate_ids = [entry["id"] for entry in candidates]
        st.session_state["power_station_search_candidates"] = candidate_ids
        st.session_state["power_station_search_choice"] = candidate_ids[0]
        st.session_state["power_station_search_status"] = "multiple"
        st.session_state["power_station_search_applied"] = False
        return

    st.session_state["power_station_search_status"] = "none"
    st.session_state["power_station_search_applied"] = False


def _confirm_power_station_search_choice() -> None:
    """동일 이름/부분검색 후보 중 사용자가 선택한 한 곳을 최종 반영합니다."""
    selected_id = str(st.session_state.get("power_station_search_choice", "") or "").strip()
    entry = POWER_STATION_SEARCH_BY_ID.get(selected_id)
    if not entry:
        st.session_state["power_station_search_status"] = "choice_required"
        return
    _apply_power_station_search_entry(entry)


def _inspector_group_for_area(area: str) -> str:
    groups: list[str] = []
    for person in _inspectors_for_major_area(area):
        group = _inspector_group_for_name(person)
        if group and group not in groups:
            groups.append(group)
    return " · ".join(groups)


def _sync_power_inspectors_from_area() -> None:
    selected_area = st.session_state.get("power_major_area", "권역 선택")
    st.session_state["power_worker"] = _automatic_inspector_display(selected_area)
    st.session_state["power_inspector_group"] = _inspector_group_for_area(selected_area)


def _update_power_inspector_group() -> None:
    selected_worker = st.session_state.get("power_worker", "담당자 선택")
    selected_area = _major_area_for_worker_value(selected_worker)
    if selected_area in POWER_REGION_DATA:
        st.session_state["power_inspector_group"] = _inspector_group_for_area(selected_area)
    else:
        st.session_state["power_inspector_group"] = _inspector_group_for_name(selected_worker)


def _clear_power_history_state() -> None:
    """국사 선택과 연동된 과거기록 조회 상태만 초기화합니다."""
    for key in (
        "power_history_records", "power_history_message", "power_history_station",
        "power_history_selected_index", "power_loaded_message",
    ):
        st.session_state.pop(key, None)


def _preserve_current_power_measurements() -> None:
    """기본정보 또는 메뉴를 변경하기 전에 현재 측정값과 특이사항을 영구 임시저장소에 보존합니다."""
    current_theme = st.session_state.get("power_current_theme", POWER_THEME_ORDER[0])
    if current_theme in POWER_THEME_ORDER:
        _save_power_theme_to_draft(current_theme)


def _mark_power_basic_info_changed() -> None:
    st.session_state["power_basic_changed_notice"] = True
    st.session_state["power_draft_saved_at"] = _korea_now().strftime("%H:%M:%S")


def _on_power_worker_change() -> None:
    """담당자를 선택하면 권역을 자동 지정하고 기존 측정값은 보존합니다."""
    st.session_state["power_station_search_applied"] = False
    st.session_state["power_station_search_notice"] = ""
    _preserve_current_power_measurements()
    selected_worker = st.session_state.get("power_worker", "담당자 선택")
    previous_area = st.session_state.get("power_major_area", "권역 선택")
    selected_area = _major_area_for_worker_value(selected_worker)
    st.session_state["power_major_area"] = selected_area
    _update_power_inspector_group()
    if selected_area != previous_area:
        st.session_state["power_mother"] = "모국 선택"
        st.session_state["power_local"] = "국소 선택"
        _clear_power_history_state()
    _mark_power_basic_info_changed()


def _on_power_major_area_change() -> None:
    """권역 변경 시 담당자 2명을 자동 설정하고 측정값은 유지합니다."""
    _preserve_current_power_measurements()
    _sync_power_inspectors_from_area()
    st.session_state["power_mother"] = "모국 선택"
    st.session_state["power_local"] = "국소 선택"
    _clear_power_history_state()
    _mark_power_basic_info_changed()


def _on_power_mother_change() -> None:
    """모국 변경 시 국소와 과거조회 상태만 초기화하고 측정값은 유지합니다."""
    st.session_state["power_station_search_applied"] = False
    st.session_state["power_station_search_notice"] = ""
    _preserve_current_power_measurements()
    st.session_state["power_local"] = "국소 선택"
    _clear_power_history_state()
    _mark_power_basic_info_changed()


def _power_headers() -> list[str]:
    headers = [
        "저장일시", "점검ID", "점검자", "운용조", "주요점검권역", "모국", "국소", "전원구분", "축전지조수",
        "입력방식", "원본점검ID", "원본저장일시",
        "입력완료율(%)", "누락항목수", "누락항목", "부분입력확인",
        "삼상전압_R-S(V)", "삼상전압_S-T(V)", "삼상전압_T-R(V)", "삼상전압_R-N(V)",
        "삼상전류_R(A)", "삼상전류_S(A)", "삼상전류_T(A)", "삼상전류_N(A)",
        "단상전압(V)", "단상전류(A)",
        "1조_측정셀수", "1조_방전후_Total전류(A)", "1조_방전후_Total전압(V)",
        "1조_최저전압(V)", "1조_최고전압(V)", "1조_방전종료전압(V)",
    ]
    headers.extend([f"1조_셀{i:02d}(V)" for i in range(1, 25)])
    headers.extend([
        "2조_측정셀수", "2조_방전후_Total전류(A)", "2조_방전후_Total전압(V)",
        "2조_최저전압(V)", "2조_최고전압(V)", "2조_방전종료전압(V)",
    ])
    headers.extend([f"2조_셀{i:02d}(V)" for i in range(1, 25)])
    headers.extend([
        "보안접지_1종(Ω)", "보안접지_2종(Ω)", "보안접지_3종(Ω)",
        "통신용접지_메인(Ω)", "피뢰침접지(Ω)", "특이사항",
        "사진수", "사진파일ID목록", "사진파일명목록",
    ])
    return headers


POWER_INSPECTION_HEADERS = _power_headers()


def _column_letter(column_number: int) -> str:
    if column_number < 1:
        raise ValueError("열 번호는 1 이상이어야 합니다.")
    letters = ""
    number = column_number
    while number:
        number, remainder = divmod(number - 1, 26)
        letters = chr(65 + remainder) + letters
    return letters


def _ensure_worksheet_grid_capacity(ws, required_rows: int = 1, required_cols: int = 1):
    """Google Sheet의 행·열 크기를 저장에 필요한 만큼 자동 확장합니다.

    기존 시트가 과거 버전의 열 수로 생성되어 있어도 새 측정항목이 추가되면
    헤더를 쓰기 전에 필요한 열까지 자동으로 늘립니다.
    """
    current_rows = int(getattr(ws, "row_count", 0) or 0)
    current_cols = int(getattr(ws, "col_count", 0) or 0)
    target_rows = max(current_rows, int(required_rows or 1))
    target_cols = max(current_cols, int(required_cols or 1))

    if target_rows == current_rows and target_cols == current_cols:
        return ws

    try:
        ws.resize(rows=target_rows, cols=target_cols)
    except TypeError:
        # 일부 gspread 버전의 위치 인자 방식도 지원합니다.
        ws.resize(target_rows, target_cols)
    return ws


def _power_sheet_has_measurement_rows(ws) -> bool:
    """헤더 아래에 실제 측정 데이터가 한 건이라도 있는지 확인합니다."""
    try:
        values = ws.get_all_values()
    except Exception:
        return True

    if len(values) <= 1:
        return False
    return any(
        any(str(cell or "").strip() for cell in row)
        for row in values[1:]
    )


def _rewrite_power_sheet_in_standard_order(ws) -> list[str]:
    """기존 데이터까지 헤더명 기준으로 재배열해 최종 표준 열 순서를 즉시 적용합니다.

    기존 시트에 데이터가 남아 있어도 각 행을 헤더명으로 다시 매핑하므로 값과 항목이
    어긋나지 않습니다. 삼상전류는 반드시 R → S → T → N 순서로 배치됩니다.
    표준 목록에 없는 기존 사용자 정의 열은 오른쪽 끝에 보존합니다.
    """
    try:
        values = ws.get_all_values()
    except Exception as read_error:
        raise RuntimeError(f"기존 Google Sheets 데이터를 읽지 못했습니다: {read_error}")

    current_headers = [str(value or "").strip() for value in (values[0] if values else [])]
    standard_headers = POWER_INSPECTION_HEADERS.copy()
    extra_headers = [
        header for header in current_headers
        if header and header not in standard_headers
    ]
    target_headers = standard_headers + extra_headers

    # 이미 정확한 표준 순서이면 불필요한 전체 재작성을 하지 않습니다.
    if current_headers == target_headers:
        _ensure_worksheet_grid_capacity(
            ws,
            required_rows=max(int(getattr(ws, "row_count", 0) or 0), 10000),
            required_cols=max(len(target_headers), 100),
        )
        return target_headers

    # 중복 헤더가 있더라도 최초 열의 값을 기준으로 안전하게 재배열합니다.
    source_index = {}
    for index, header in enumerate(current_headers):
        if header and header not in source_index:
            source_index[header] = index

    reordered_rows = [target_headers]
    for source_row in values[1:]:
        reordered_rows.append([
            source_row[source_index[header]]
            if header in source_index and source_index[header] < len(source_row)
            else ""
            for header in target_headers
        ])

    required_rows = max(
        int(getattr(ws, "row_count", 0) or 0),
        len(reordered_rows),
        10000,
    )
    required_cols = max(
        int(getattr(ws, "col_count", 0) or 0),
        len(current_headers),
        len(target_headers),
        100,
    )
    _ensure_worksheet_grid_capacity(
        ws,
        required_rows=required_rows,
        required_cols=required_cols,
    )

    # 기존 값 영역을 비운 뒤 헤더와 모든 행을 표준 순서로 다시 기록합니다.
    clear_last_row = max(len(values), 1)
    clear_range = f"A1:{_column_letter(required_cols)}{clear_last_row}"
    try:
        ws.batch_clear([clear_range])
    except Exception:
        try:
            ws.update(
                range_name=clear_range,
                values=[[""] * required_cols for _ in range(clear_last_row)],
            )
        except TypeError:
            ws.update(
                clear_range,
                [[""] * required_cols for _ in range(clear_last_row)],
            )

    end_col = _column_letter(len(target_headers))
    end_row = len(reordered_rows)
    try:
        ws.update(
            range_name=f"A1:{end_col}{end_row}",
            values=reordered_rows,
            value_input_option="USER_ENTERED",
        )
    except TypeError:
        ws.update(
            f"A1:{end_col}{end_row}",
            reordered_rows,
            value_input_option="USER_ENTERED",
        )

    # 재작성 결과를 다시 확인하여 R/S/T/N 순서가 실제 시트에 반영됐는지 검증합니다.
    verified_headers = [str(value or "").strip() for value in ws.row_values(1)]
    expected_sequence = [
        "삼상전류_R(A)",
        "삼상전류_S(A)",
        "삼상전류_T(A)",
        "삼상전류_N(A)",
    ]
    sequence_start = target_headers.index("삼상전류_R(A)")
    if verified_headers[sequence_start:sequence_start + 4] != expected_sequence:
        raise RuntimeError("Google Sheets의 삼상전류 R/S/T/N 열 순서 재배치에 실패했습니다.")
    return target_headers


def _ensure_power_inspection_sheet(spreadsheet):
    try:
        ws = spreadsheet.worksheet(POWER_INSPECTION_SHEET_NAME)
    except Exception:
        ws = spreadsheet.add_worksheet(
            title=POWER_INSPECTION_SHEET_NAME,
            rows=10000,
            cols=max(len(POWER_INSPECTION_HEADERS) + 5, 100),
        )

    # 데이터 유무와 관계없이 기존 행을 헤더명으로 안전하게 재매핑하여
    # 삼상전류 R → S → T → N 표준 순서를 실제 시트에 즉시 적용합니다.
    current_headers = _rewrite_power_sheet_in_standard_order(ws)

    _ensure_worksheet_grid_capacity(
        ws,
        required_rows=max(int(getattr(ws, "row_count", 0) or 0), 10000),
        required_cols=max(len(current_headers), 100),
    )

    try:
        ws.freeze(rows=1)
        ws.format(
            f"A1:{_column_letter(len(current_headers))}1",
            {
                "backgroundColor": {"red": 0.86, "green": 0.92, "blue": 0.98},
                "textFormat": {"bold": True},
                "horizontalAlignment": "CENTER",
            },
        )
    except Exception:
        pass

    return ws, current_headers


def _extract_appended_row_number(append_response) -> int | None:
    """gspread append_row 응답에서 실제 추가된 행 번호를 추출합니다."""
    if not isinstance(append_response, dict):
        return None
    updated_range = ""
    updates = append_response.get("updates")
    if isinstance(updates, dict):
        updated_range = str(updates.get("updatedRange", "") or "")
    if not updated_range:
        updated_range = str(append_response.get("updatedRange", "") or "")
    match = re.search(r"![A-Z]+(\d+):[A-Z]+(\d+)$", updated_range)
    if not match:
        return None
    return int(match.group(1))


def _locate_saved_inspection_row(ws, sheet_headers: list[str], inspection_id: str) -> int | None:
    """append 응답에 행 번호가 없을 때 점검ID 열에서 저장 행을 찾습니다."""
    if not inspection_id or "점검ID" not in sheet_headers:
        return None
    id_col = sheet_headers.index("점검ID") + 1
    try:
        values = ws.col_values(id_col)
    except Exception:
        return None
    for row_number in range(len(values), 1, -1):
        if str(values[row_number - 1]).strip() == inspection_id:
            return row_number
    return None


def _ensure_n_phase_current_saved(
    ws,
    sheet_headers: list[str],
    append_response,
    inspection_id: str,
    n_phase_value,
) -> None:
    """삼상 N상 전류를 저장 직후 확인하고 누락 시 정확한 셀에 보정 저장합니다.

    과거 버전 시트와 최종 표준 시트 모두에서 헤더명으로 실제 열을 찾고,
    저장된 행의 해당 셀을 직접 검증하여 N상 전류 누락을 방지합니다.
    """
    if n_phase_value in ("", None):
        return

    target_header = "삼상전류_N(A)"
    if target_header not in sheet_headers:
        raise RuntimeError("Google Sheets에 '삼상전류_N(A)' 헤더가 생성되지 않았습니다.")

    row_number = _extract_appended_row_number(append_response)
    if row_number is None:
        row_number = _locate_saved_inspection_row(ws, sheet_headers, inspection_id)
    if row_number is None:
        raise RuntimeError("저장된 행을 찾지 못해 N상 전류를 확인할 수 없습니다.")

    column_number = sheet_headers.index(target_header) + 1
    _ensure_worksheet_grid_capacity(
        ws,
        required_rows=max(int(getattr(ws, "row_count", 0) or 0), row_number),
        required_cols=max(int(getattr(ws, "col_count", 0) or 0), column_number),
    )
    cell_ref = f"{_column_letter(column_number)}{row_number}"

    existing_value = ""
    try:
        existing_value = str(ws.acell(cell_ref).value or "").strip()
    except Exception:
        existing_value = ""

    if existing_value:
        return

    last_error = None
    for attempt in range(3):
        try:
            ws.update(
                range_name=cell_ref,
                values=[[n_phase_value]],
                value_input_option="USER_ENTERED",
            )
            verified = str(ws.acell(cell_ref).value or "").strip()
            if verified:
                return
            last_error = RuntimeError("N상 전류 셀의 저장값이 비어 있습니다.")
        except TypeError:
            try:
                ws.update(cell_ref, [[n_phase_value]], value_input_option="USER_ENTERED")
                verified = str(ws.acell(cell_ref).value or "").strip()
                if verified:
                    return
                last_error = RuntimeError("N상 전류 셀의 저장값이 비어 있습니다.")
            except Exception as update_error:
                last_error = update_error
        except Exception as update_error:
            last_error = update_error
        time.sleep(0.4 * (attempt + 1))

    raise RuntimeError(f"N상 전류 저장 확인에 실패했습니다: {last_error}")


def _parse_power_number(value, implicit_decimals: int | None = None):
    """숫자만 입력한 측정값을 실제 숫자로 변환합니다.

    예시:
    - 전압·전류 3800, decimals=1 → 380.0
    - 축전지 셀 215, decimals=2 → 2.15
    - 접지저항 123, decimals=2 → 1.23
    """
    raw = str(value or "").strip().replace(",", "")
    if not raw:
        return ""

    cleaned = re.sub(r"[^0-9.+-]", "", raw)
    if cleaned in {"", "+", "-", ".", "+.", "-."}:
        return ""

    try:
        if implicit_decimals is not None and "." not in cleaned:
            sign = -1 if cleaned.startswith("-") else 1
            digits = cleaned.lstrip("+-")
            if not digits.isdigit():
                return ""
            return sign * (int(digits) / (10 ** implicit_decimals))
        return float(cleaned)
    except (TypeError, ValueError, OverflowError):
        return ""


def _parse_battery_cell_number(value):
    """방전 후 셀 전압은 현장 입력 자릿수에 따라 2~3자리 소수를 허용합니다.

    - 215  → 2.15V
    - 3507 → 3.507V
    - 0.00 / 0.000처럼 직접 입력한 소수점은 그대로 숫자로 저장
    """
    raw = str(value or "").strip().replace(",", "")
    if not raw:
        return ""
    cleaned = re.sub(r"[^0-9.+-]", "", raw)
    if cleaned in {"", "+", "-", ".", "+.", "-."}:
        return ""
    try:
        if "." in cleaned:
            return float(cleaned)
        sign = -1 if cleaned.startswith("-") else 1
        digits = cleaned.lstrip("+-")
        if not digits.isdigit():
            return ""
        decimals = 3 if len(digits) >= 4 else 2
        return sign * (int(digits) / (10 ** decimals))
    except (TypeError, ValueError, OverflowError):
        return ""


def _format_power_display(value, decimals: int) -> str:
    raw = str(value or "").strip().replace(",", "")
    if not raw:
        return ""
    cleaned = re.sub(r"[^0-9.-]", "", raw)
    if not cleaned:
        return ""
    try:
        number = float(cleaned)
        return f"{number:.{decimals}f}"
    except Exception:
        return raw


def _format_battery_cell_display(value) -> str:
    """Google Sheets의 셀 전압 자릿수를 가능한 한 보존합니다."""
    raw = str(value or "").strip().replace(",", "")
    if not raw:
        return ""
    cleaned = re.sub(r"[^0-9.-]", "", raw)
    if not cleaned:
        return ""
    try:
        number = float(cleaned)
        if "." in cleaned:
            fraction_len = len(cleaned.split(".", 1)[1])
            decimals = max(2, min(3, fraction_len))
        else:
            decimals = 2
        return f"{number:.{decimals}f}"
    except Exception:
        return raw


def _power_value_is_blank(value) -> bool:
    return value is None or (isinstance(value, str) and not value.strip())


def _power_draft() -> dict:
    draft = st.session_state.get("power_draft")
    if not isinstance(draft, dict):
        draft = {}
        st.session_state["power_draft"] = draft
    return draft


def _power_widget_key(data_key: str) -> str:
    """화면 위젯 키와 영구 임시저장 키를 분리합니다.

    Streamlit은 현재 화면에서 사라진 위젯 키를 정리할 수 있으므로,
    측정값은 power_draft와 data_key에 별도로 보존하고 화면에는 _ui_ 키를 사용합니다.
    """
    return f"_ui_{data_key}"


def _power_get(key: str, default=""):
    """화면의 최신값 → 임시저장값 → 영구 세션값 순으로 값을 반환합니다."""
    ui_key = _power_widget_key(key)
    if ui_key in st.session_state:
        return st.session_state.get(ui_key, default)
    draft = _power_draft()
    if key in draft:
        return draft.get(key, default)
    return st.session_state.get(key, default)


def _power_set(key: str, value) -> None:
    """영구 임시저장값과 현재 렌더링된 화면값을 함께 갱신합니다."""
    _power_draft()[key] = value
    st.session_state[key] = value
    ui_key = _power_widget_key(key)
    if ui_key in st.session_state:
        st.session_state[ui_key] = value


def _persist_power_widget(key: str) -> None:
    """화면 위젯값을 영구 임시저장소로 즉시 복사합니다."""
    ui_key = _power_widget_key(key)
    value = st.session_state.get(ui_key, st.session_state.get(key, ""))
    _power_draft()[key] = value
    st.session_state[key] = value
    st.session_state["power_draft_saved_at"] = _korea_now().strftime("%H:%M:%S")


def _hydrate_power_widget(key: str, default="") -> None:
    """매 렌더링 시 임시저장값으로 화면 위젯을 복원합니다."""
    draft = _power_draft()
    value = draft.get(key, st.session_state.get(key, default))
    st.session_state[key] = value
    # 위젯이 만들어지기 전에 항상 화면 키를 임시저장값으로 맞춥니다.
    st.session_state[_power_widget_key(key)] = value


def _power_text_input(label: str, key: str, **kwargs):
    _hydrate_power_widget(key, "")
    ui_key = _power_widget_key(key)
    result = st.text_input(
        label,
        key=ui_key,
        on_change=_persist_power_widget,
        args=(key,),
        **kwargs,
    )
    # 버튼 클릭 등 다른 이벤트로 재실행되더라도 현재 화면값을 놓치지 않습니다.
    current_value = st.session_state.get(ui_key, result)
    _power_draft()[key] = current_value
    st.session_state[key] = current_value
    return current_value


def _power_text_area(label: str, key: str, **kwargs):
    _hydrate_power_widget(key, "")
    ui_key = _power_widget_key(key)
    result = st.text_area(
        label,
        key=ui_key,
        on_change=_persist_power_widget,
        args=(key,),
        **kwargs,
    )
    current_value = st.session_state.get(ui_key, result)
    _power_draft()[key] = current_value
    st.session_state[key] = current_value
    return current_value


def _power_theme_keys(theme: str) -> list[str]:
    if theme == "전압·전류 측정":
        return [
            "power_phase_type",
            "power_three_voltage_rs", "power_three_voltage_st",
            "power_three_voltage_tr", "power_three_voltage_rn",
            "power_three_current_r", "power_three_current_s", "power_three_current_t", "power_three_current_n",
            "power_single_voltage", "power_single_current",
        ]
    if theme == "축전지 측정":
        keys = ["power_battery_set", "power_battery2_enabled"]
        for group in (1, 2):
            keys.extend([
                f"power_battery{group}_total_current",
                f"power_battery{group}_total_voltage",
                f"power_battery{group}_min_voltage",
                f"power_battery{group}_max_voltage",
                f"power_battery{group}_end_voltage",
            ])
            keys.extend(f"power_battery_{group}_{index:02d}" for index in range(1, 25))
        return keys
    if theme == "접지저항 측정":
        return [
            "power_security_ground_1", "power_security_ground_2", "power_security_ground_3",
            "power_telecom_ground", "power_lightning_ground",
        ]
    if theme == "최종 확인·전송":
        return ["power_notes"]
    return []


def _save_power_theme_to_draft(theme: str) -> None:
    """현재 화면값을 shadow UI 키에서 읽어 영구 임시저장소에 스냅샷합니다."""
    draft = _power_draft()
    for key in _power_theme_keys(theme):
        ui_key = _power_widget_key(key)
        if ui_key in st.session_state:
            value = st.session_state.get(ui_key, "")
        elif key in draft:
            value = draft.get(key, "")
        else:
            value = st.session_state.get(key, "")
        draft[key] = value
        st.session_state[key] = value
    st.session_state["power_draft_saved_at"] = _korea_now().strftime("%H:%M:%S")


def _hydrate_power_theme_from_draft(theme: str) -> None:
    """선택한 테마의 모든 값을 영구 임시저장소에서 복원합니다."""
    draft = _power_draft()
    for key in _power_theme_keys(theme):
        if key in draft:
            st.session_state[key] = draft[key]


def _on_power_phase_change() -> None:
    _persist_power_widget("power_phase_type")


def _on_power_battery_set_change() -> None:
    _persist_power_widget("power_battery_set")
    _save_power_theme_to_draft("축전지 측정")
    selected = _power_get("power_battery_set", "1조 셀 측정")
    if selected == "2조 셀 측정":
        _power_set("power_battery2_enabled", True)


def _power_basic_missing() -> list[str]:
    missing: list[str] = []
    worker = str(st.session_state.get("power_worker", "담당자 선택")).strip()
    major_area = str(st.session_state.get("power_major_area", "권역 선택")).strip()
    if not _power_worker_matches_area(worker, major_area):
        missing.append("담당자")
    if major_area == "권역 선택" or major_area not in POWER_REGION_DATA:
        missing.append("주요 점검권역")
    if st.session_state.get("power_mother", "모국 선택") == "모국 선택":
        missing.append("모국")
    if st.session_state.get("power_local", "국소 선택") == "국소 선택":
        missing.append("국소")
    return missing


def _power_battery2_enabled() -> bool:
    explicit_enabled = _power_get("power_battery2_enabled", None)
    if explicit_enabled is not None:
        return bool(explicit_enabled)
    if _power_get("power_battery_set", "1조 셀 측정") == "2조 셀 측정":
        return True
    keys = [
        "power_battery2_total_current", "power_battery2_total_voltage",
        "power_battery2_min_voltage", "power_battery2_max_voltage", "power_battery2_end_voltage",
    ]
    keys.extend(f"power_battery_2_{index:02d}" for index in range(1, 25))
    return any(not _power_value_is_blank(_power_get(key, "")) for key in keys)


def _measured_cell_count(cells: list) -> int:
    """마지막으로 값이 입력된 셀 번호를 실제 측정 셀 수로 사용합니다."""
    highest = 0
    for index, value in enumerate(list(cells or [])[:24], 1):
        if not _power_value_is_blank(value):
            highest = index
    return highest


def _power_payload_missing_items(payload: dict) -> list[str]:
    expected: list[tuple[str, object]] = []
    phase_type = str(payload.get("phase_type", "")).strip()

    if phase_type == "삼상":
        expected.extend([
            ("삼상 R-S 전압", payload.get("three_voltage_rs")),
            ("삼상 S-T 전압", payload.get("three_voltage_st")),
            ("삼상 T-R 전압", payload.get("three_voltage_tr")),
            ("삼상 R-N 전압", payload.get("three_voltage_rn")),
            ("삼상 R상 전류", payload.get("three_current_r")),
            ("삼상 S상 전류", payload.get("three_current_s")),
            ("삼상 T상 전류", payload.get("three_current_t")),
            ("삼상 N상 전류", payload.get("three_current_n")),
        ])
    elif phase_type == "단상":
        expected.extend([
            ("단상 전압", payload.get("single_voltage")),
            ("단상 전류", payload.get("single_current")),
        ])

    expected.extend([
        ("1조 방전 후 Total 전류", payload.get("battery1_total_current")),
        ("1조 방전 후 Total 전압", payload.get("battery1_total_voltage")),
        ("1조 최저전압", payload.get("battery1_min_voltage")),
        ("1조 최고전압", payload.get("battery1_max_voltage")),
        ("1조 방전종료 전압", payload.get("battery1_end_voltage")),
    ])
    battery1_cells = list(payload.get("battery1_cells", []))[:24]
    battery1_cells.extend([""] * (24 - len(battery1_cells)))
    battery1_cell_count = int(payload.get("battery1_cell_count", 0) or 0)
    if battery1_cell_count <= 0:
        battery1_cell_count = _measured_cell_count(battery1_cells)
    expected.extend(
        (f"1조 {index}셀", battery1_cells[index - 1])
        for index in range(1, min(battery1_cell_count, 24) + 1)
    )

    if int(payload.get("battery_group_count", 1) or 1) == 2:
        expected.extend([
            ("2조 방전 후 Total 전류", payload.get("battery2_total_current")),
            ("2조 방전 후 Total 전압", payload.get("battery2_total_voltage")),
            ("2조 최저전압", payload.get("battery2_min_voltage")),
            ("2조 최고전압", payload.get("battery2_max_voltage")),
            ("2조 방전종료 전압", payload.get("battery2_end_voltage")),
        ])
        battery2_cells = list(payload.get("battery2_cells", []))[:24]
        battery2_cells.extend([""] * (24 - len(battery2_cells)))
        battery2_cell_count = int(payload.get("battery2_cell_count", 0) or 0)
        if battery2_cell_count <= 0:
            battery2_cell_count = _measured_cell_count(battery2_cells)
        expected.extend(
            (f"2조 {index}셀", battery2_cells[index - 1])
            for index in range(1, min(battery2_cell_count, 24) + 1)
        )

    expected.extend([
        ("보안접지 1종", payload.get("security_ground_1")),
        ("보안접지 2종", payload.get("security_ground_2")),
        ("보안접지 3종", payload.get("security_ground_3")),
        ("통신용접지(메인)", payload.get("telecom_ground")),
        ("피뢰침접지", payload.get("lightning_ground")),
    ])

    return [label for label, value in expected if _power_value_is_blank(value)]


def _power_expected_item_count(payload: dict) -> int:
    phase_count = 8 if str(payload.get("phase_type", "")).strip() == "삼상" else 2
    battery1_count = int(payload.get("battery1_cell_count", 0) or 0)
    battery_count = 5 + max(0, min(battery1_count, 24))
    if int(payload.get("battery_group_count", 1) or 1) == 2:
        battery2_count = int(payload.get("battery2_cell_count", 0) or 0)
        battery_count += 5 + max(0, min(battery2_count, 24))
    return phase_count + battery_count + 5


def _power_has_measurement(payload: dict) -> bool:
    measurement_keys = [
        "three_voltage_rs", "three_voltage_st", "three_voltage_tr", "three_voltage_rn",
        "three_current_r", "three_current_s", "three_current_t", "three_current_n",
        "single_voltage", "single_current",
        "battery1_total_current", "battery1_total_voltage", "battery1_min_voltage",
        "battery1_max_voltage", "battery1_end_voltage",
        "battery2_total_current", "battery2_total_voltage", "battery2_min_voltage",
        "battery2_max_voltage", "battery2_end_voltage",
        "security_ground_1", "security_ground_2", "security_ground_3",
        "telecom_ground", "lightning_ground",
    ]
    if any(payload.get(key, "") not in ("", None) for key in measurement_keys):
        return True
    if any(value not in ("", None) for value in payload.get("battery1_cells", [])):
        return True
    if any(value not in ("", None) for value in payload.get("battery2_cells", [])):
        return True
    return bool(str(payload.get("notes", "")).strip())


def save_power_inspection_result(payload: dict, photos: list | None = None) -> tuple[bool, str, str]:
    """국사 전원시설 정밀점검 결과와 선택 사진을 Google Sheet/Drive에 함께 저장합니다."""
    photos = list(photos or [])[:WORK_LOG_MAX_PHOTOS]
    client = init_google_sheet_connection()
    if not client:
        return False, "구글 시트 연결 실패: Streamlit Secrets의 gcp_service_account 설정을 확인하세요.", ""

    try:
        worker = str(payload.get("worker", "")).strip()
        major_area = str(payload.get("major_area", "")).strip()
        mother = str(payload.get("mother", "")).strip()
        local = str(payload.get("local", "")).strip()
        phase_type = str(payload.get("phase_type", "")).strip()

        if not worker or worker == "담당자 선택":
            return False, "담당자를 선택해 주세요.", ""
        if not _power_worker_matches_area(worker, major_area):
            return False, "선택한 담당자와 주요 점검권역 정보가 일치하지 않습니다.", ""
        area_map = _power_area_station_map(major_area)
        if mother not in area_map:
            return False, "선택한 권역에 포함된 모국을 선택해 주세요.", ""
        if local not in area_map.get(mother, []):
            return False, "선택한 권역·모국·국소의 조합이 올바르지 않습니다.", ""
        if phase_type not in {"삼상", "단상"}:
            return False, "삼상 또는 단상 측정 방식을 선택해 주세요.", ""
        if not _power_has_measurement(payload):
            return False, "측정값 또는 특이사항을 한 개 이상 입력해 주세요.", ""

        missing_items = _power_payload_missing_items(payload)
        if not bool(payload.get("final_confirmed", False)):
            return False, "최종 확인에 동의해 주세요.", ""

        expected_count = max(_power_expected_item_count(payload), 1)
        completed_count = expected_count - len(missing_items)
        completion_rate = round((completed_count / expected_count) * 100, 1)

        spreadsheet = client.open(POWER_INSPECTION_SPREADSHEET_NAME)
        ws, sheet_headers = _ensure_power_inspection_sheet(spreadsheet)

        now = _korea_now()
        saved_at = now.strftime("%Y-%m-%d %H:%M:%S")
        inspection_seed = f"{saved_at}|{worker}|{major_area}|{mother}|{local}|{time.time_ns()}"
        inspection_id = hashlib.sha256(inspection_seed.encode("utf-8")).hexdigest()[:14]
        source_id = str(payload.get("source_inspection_id", "")).strip()
        source_saved_at = str(payload.get("source_saved_at", "")).strip()

        # 정밀점검 사진은 WORK LOG와 동일한 비공개 Apps Script → Drive 경로로 저장합니다.
        # 사진 저장이 실패하면 Google Sheets 행도 만들지 않아 기록/사진이 어긋나지 않도록 합니다.
        photo_ids: list[str] = []
        photo_names: list[str] = []
        used_power_photo_names: set[str] = set()
        if photos:
            preflight_ok, preflight_message = _worklog_apps_script_healthcheck()
            if not preflight_ok:
                return False, f"사진 업로드 사전진단 실패: {preflight_message}", ""

            for photo_index, photo in enumerate(photos, 1):
                capture_stamp = _photo_capture_timestamp(photo, fallback_dt=now)
                compressed, _safe_name, mime_type, photo_error = _worklog_compress_image(photo)
                if not compressed:
                    return False, f"{photo_index}번째 사진 처리 실패: {photo_error}", ""
                base_drive_name = f"정밀점검_{capture_stamp}"
                drive_name = f"{base_drive_name}.jpg"
                duplicate_no = 2
                while drive_name in used_power_photo_names:
                    drive_name = f"{base_drive_name}_{duplicate_no:02d}.jpg"
                    duplicate_no += 1
                used_power_photo_names.add(drive_name)
                photo_ok, drive_meta, upload_error = _worklog_upload_drive_image(
                    compressed, drive_name, mime_type
                )
                if not photo_ok:
                    return False, f"{photo_index}번째 사진 저장 실패: {upload_error}", ""
                photo_ids.append(str(drive_meta.get("id", "") or ""))
                photo_names.append(str(drive_meta.get("name", drive_name) or drive_name))

        row_map = {
            "저장일시": saved_at,
            "점검ID": inspection_id,
            "점검자": worker,
            "운용조": str(payload.get("inspector_group", "")).strip() or _inspector_group_for_area(major_area),
            "주요점검권역": major_area,
            "모국": mother,
            "국소": local,
            "전원구분": phase_type,
            "축전지조수": int(payload.get("battery_group_count", 1) or 1),
            "입력방식": "기존값 불러오기 후 수정" if source_id else "신규 입력",
            "원본점검ID": source_id,
            "원본저장일시": source_saved_at,
            "입력완료율(%)": completion_rate,
            "누락항목수": len(missing_items),
            "누락항목": ", ".join(missing_items),
            "부분입력확인": "확인" if missing_items else "해당없음",
            "삼상전압_R-S(V)": payload.get("three_voltage_rs", "") if phase_type == "삼상" else "",
            "삼상전압_S-T(V)": payload.get("three_voltage_st", "") if phase_type == "삼상" else "",
            "삼상전압_T-R(V)": payload.get("three_voltage_tr", "") if phase_type == "삼상" else "",
            "삼상전압_R-N(V)": payload.get("three_voltage_rn", "") if phase_type == "삼상" else "",
            "삼상전류_R(A)": payload.get("three_current_r", "") if phase_type == "삼상" else "",
            "삼상전류_S(A)": payload.get("three_current_s", "") if phase_type == "삼상" else "",
            "삼상전류_T(A)": payload.get("three_current_t", "") if phase_type == "삼상" else "",
            "삼상전류_N(A)": payload.get("three_current_n", "") if phase_type == "삼상" else "",
            "단상전압(V)": payload.get("single_voltage", "") if phase_type == "단상" else "",
            "단상전류(A)": payload.get("single_current", "") if phase_type == "단상" else "",
            "1조_측정셀수": int(payload.get("battery1_cell_count", 0) or 0),
            "1조_방전후_Total전류(A)": payload.get("battery1_total_current", ""),
            "1조_방전후_Total전압(V)": payload.get("battery1_total_voltage", ""),
            "1조_최저전압(V)": payload.get("battery1_min_voltage", ""),
            "1조_최고전압(V)": payload.get("battery1_max_voltage", ""),
            "1조_방전종료전압(V)": payload.get("battery1_end_voltage", ""),
            "2조_측정셀수": int(payload.get("battery2_cell_count", 0) or 0) if int(payload.get("battery_group_count", 1) or 1) == 2 else "",
            "2조_방전후_Total전류(A)": payload.get("battery2_total_current", ""),
            "2조_방전후_Total전압(V)": payload.get("battery2_total_voltage", ""),
            "2조_최저전압(V)": payload.get("battery2_min_voltage", ""),
            "2조_최고전압(V)": payload.get("battery2_max_voltage", ""),
            "2조_방전종료전압(V)": payload.get("battery2_end_voltage", ""),
            "보안접지_1종(Ω)": payload.get("security_ground_1", ""),
            "보안접지_2종(Ω)": payload.get("security_ground_2", ""),
            "보안접지_3종(Ω)": payload.get("security_ground_3", ""),
            "통신용접지_메인(Ω)": payload.get("telecom_ground", ""),
            "피뢰침접지(Ω)": payload.get("lightning_ground", ""),
            "특이사항": str(payload.get("notes", "")).strip(),
            "사진수": len(photo_ids),
            "사진파일ID목록": "|".join(photo_ids),
            "사진파일명목록": "|".join(photo_names),
        }

        battery1_cells = list(payload.get("battery1_cells", []))[:24]
        battery1_cells.extend([""] * (24 - len(battery1_cells)))
        battery2_cells = list(payload.get("battery2_cells", []))[:24]
        battery2_cells.extend([""] * (24 - len(battery2_cells)))
        for index, value in enumerate(battery1_cells, 1):
            row_map[f"1조_셀{index:02d}(V)"] = value
        for index, value in enumerate(battery2_cells, 1):
            row_map[f"2조_셀{index:02d}(V)"] = value

        row = [row_map.get(header, "") for header in sheet_headers]
        last_error = None
        append_response = None
        for attempt in range(5):
            try:
                append_response = ws.append_row(row, value_input_option="USER_ENTERED")
                break
            except Exception as append_error:
                last_error = append_error
                error_text = str(append_error)
                if any(token in error_text for token in ("429", "Quota exceeded", "RESOURCE_EXHAUSTED")):
                    time.sleep(1.1 * (attempt + 1))
                    continue
                raise

        if append_response is None:
            return False, f"저장 요청이 집중되어 전송하지 못했습니다. 다시 전송해 주세요. ({last_error})", ""

        # 최종 표준 순서(R/S/T/N)에서도 N상 전류가 실제 저장됐는지 확인합니다.
        if phase_type == "삼상":
            n_phase_value = row_map.get("삼상전류_N(A)", "")
            try:
                _ensure_n_phase_current_saved(
                    ws,
                    sheet_headers,
                    append_response,
                    inspection_id,
                    n_phase_value,
                )
            except Exception as n_phase_error:
                return False, (
                    "기본 측정데이터 행은 저장되었으나 N상 전류 저장 확인에 실패했습니다. "
                    f"관리자에게 확인해 주세요. ({n_phase_error})"
                ), inspection_id

        photo_message = f" · 현장사진 {len(photo_ids)}장 Drive 저장" if photo_ids else ""
        return True, f"측정값과 N상 전류가 Google Sheets에 정상 저장되었습니다.{photo_message}", inspection_id
    except Exception as e:
        return False, str(e), ""


def load_recent_power_inspection(mother: str, local: str, within_days: int = 60) -> tuple[bool, str, dict]:
    """동일 모국·국소의 최근 측정값을 찾아 입력폼 재사용용으로 반환합니다."""
    client = init_google_sheet_connection()
    if not client:
        return False, "구글 시트 연결 실패: Secrets 설정을 확인하세요.", {}
    if mother not in POWER_STATION_MAP or local not in POWER_STATION_MAP.get(mother, []):
        return False, "먼저 모국과 국소를 정확히 선택해 주세요.", {}

    try:
        spreadsheet = client.open(POWER_INSPECTION_SPREADSHEET_NAME)
        try:
            ws = spreadsheet.worksheet(POWER_INSPECTION_SHEET_NAME)
        except Exception:
            return False, "아직 저장된 전원 정밀점검 기록이 없습니다.", {}

        values = ws.get_all_values()
        if len(values) < 2:
            return False, "아직 저장된 전원 정밀점검 기록이 없습니다.", {}

        headers = values[0]
        now_naive = _korea_now().replace(tzinfo=None)
        cutoff = now_naive - datetime.timedelta(days=max(1, int(within_days)))
        for row in reversed(values[1:]):
            record = {headers[index]: row[index] if index < len(row) else "" for index in range(len(headers))}
            if str(record.get("모국", "")).strip() != mother or str(record.get("국소", "")).strip() != local:
                continue
            saved_text = str(record.get("저장일시", "")).strip()
            try:
                saved_dt = datetime.datetime.strptime(saved_text, "%Y-%m-%d %H:%M:%S")
            except Exception:
                continue
            if saved_dt < cutoff:
                continue
            return True, f"최근 측정값을 불러왔습니다. ({saved_text})", record
        return False, f"최근 {within_days}일 이내 동일 국소의 저장 기록이 없습니다.", {}
    except Exception as e:
        return False, str(e), {}


def list_power_inspection_history(
    mother: str,
    local: str,
    within_days: int = 183,
    max_records: int = 100,
) -> tuple[bool, str, list[dict]]:
    """동일 모국·국소의 과거 측정기록을 최신순으로 반환합니다."""
    client = init_google_sheet_connection()
    if not client:
        return False, "구글 시트 연결 실패: Secrets 설정을 확인하세요.", []
    if mother not in POWER_STATION_MAP or local not in POWER_STATION_MAP.get(mother, []):
        return False, "먼저 모국과 국소를 정확히 선택해 주세요.", []

    try:
        spreadsheet = client.open(POWER_INSPECTION_SPREADSHEET_NAME)
        try:
            ws = spreadsheet.worksheet(POWER_INSPECTION_SHEET_NAME)
        except Exception:
            return False, "아직 저장된 전원 정밀점검 기록이 없습니다.", []

        values = ws.get_all_values()
        if len(values) < 2:
            return False, "아직 저장된 전원 정밀점검 기록이 없습니다.", []

        headers = values[0]
        now_naive = _korea_now().replace(tzinfo=None)
        cutoff = now_naive - datetime.timedelta(days=max(1, int(within_days)))
        records: list[dict] = []

        for row in reversed(values[1:]):
            record = {
                headers[index]: row[index] if index < len(row) else ""
                for index in range(len(headers))
            }
            if str(record.get("모국", "")).strip() != mother:
                continue
            if str(record.get("국소", "")).strip() != local:
                continue

            saved_text = str(record.get("저장일시", "")).strip()
            try:
                saved_dt = datetime.datetime.strptime(saved_text, "%Y-%m-%d %H:%M:%S")
            except Exception:
                continue
            if saved_dt < cutoff:
                continue

            records.append(record)
            if len(records) >= max(1, int(max_records)):
                break

        if not records:
            return False, f"최근 {within_days}일 이내 동일 국소의 저장 기록이 없습니다.", []
        return True, f"과거 측정기록 {len(records)}건을 조회했습니다.", records
    except Exception as e:
        return False, str(e), []


def _set_power_state_from_record(record: dict) -> None:
    draft = _power_draft()

    def set_value(key: str, header: str, decimals: int | None = None) -> None:
        value = str(record.get(header, "")).strip()
        if not value:
            formatted = ""
        elif decimals is None:
            formatted = value
        else:
            formatted = _format_power_display(value, decimals)
        draft[key] = formatted
        st.session_state[key] = formatted

    phase = str(record.get("전원구분", "삼상")).strip()
    phase_value = phase if phase in {"삼상", "단상"} else "삼상"
    draft["power_phase_type"] = phase_value
    st.session_state["power_phase_type"] = phase_value

    phase_map = [
        ("power_three_voltage_rs", "삼상전압_R-S(V)", 1),
        ("power_three_voltage_st", "삼상전압_S-T(V)", 1),
        ("power_three_voltage_tr", "삼상전압_T-R(V)", 1),
        ("power_three_voltage_rn", "삼상전압_R-N(V)", 1),
        ("power_three_current_r", "삼상전류_R(A)", 1),
        ("power_three_current_s", "삼상전류_S(A)", 1),
        ("power_three_current_t", "삼상전류_T(A)", 1),
        ("power_three_current_n", "삼상전류_N(A)", 1),
        ("power_single_voltage", "단상전압(V)", 1),
        ("power_single_current", "단상전류(A)", 1),
    ]
    for key, header, decimals in phase_map:
        set_value(key, header, decimals)

    for group in (1, 2):
        prefix = f"power_battery{group}"
        set_value(f"{prefix}_total_current", f"{group}조_방전후_Total전류(A)", 1)
        set_value(f"{prefix}_total_voltage", f"{group}조_방전후_Total전압(V)", 2)
        set_value(f"{prefix}_min_voltage", f"{group}조_최저전압(V)", 2)
        set_value(f"{prefix}_max_voltage", f"{group}조_최고전압(V)", 2)
        set_value(f"{prefix}_end_voltage", f"{group}조_방전종료전압(V)", 2)
        for index in range(1, 25):
            key = f"power_battery_{group}_{index:02d}"
            value = str(record.get(f"{group}조_셀{index:02d}(V)", "")).strip()
            formatted = _format_battery_cell_display(value) if value else ""
            draft[key] = formatted
            st.session_state[key] = formatted

    ground_map = [
        ("power_security_ground_1", "보안접지_1종(Ω)"),
        ("power_security_ground_2", "보안접지_2종(Ω)"),
        ("power_security_ground_3", "보안접지_3종(Ω)"),
        ("power_telecom_ground", "통신용접지_메인(Ω)"),
        ("power_lightning_ground", "피뢰침접지(Ω)"),
    ]
    for key, header in ground_map:
        set_value(key, header, 2)
    set_value("power_notes", "특이사항", None)

    group_count = str(record.get("축전지조수", "1")).strip()
    has_group2 = group_count == "2" or any(
        str(record.get(f"2조_셀{index:02d}(V)", "")).strip() for index in range(1, 25)
    )
    draft["power_battery2_enabled"] = has_group2
    draft["power_battery_set"] = "1조 셀 측정"
    st.session_state["power_battery2_enabled"] = has_group2
    st.session_state["power_battery_set"] = "1조 셀 측정"
    st.session_state["power_loaded_source_id"] = str(record.get("점검ID", "")).strip()
    st.session_state["power_loaded_source_saved_at"] = str(record.get("저장일시", "")).strip()
    st.session_state["power_loaded_notice"] = True
    st.session_state["power_draft_saved_at"] = _korea_now().strftime("%H:%M:%S")
    # 최근 측정값을 불러온 경우 모든 테마를 바로 확인·수정할 수 있도록 잠금을 해제합니다.
    st.session_state["power_unlocked_theme_index"] = len(POWER_THEME_ORDER) - 1
    st.session_state["power_theme_confirmations"] = {
        theme: {
            "answer": "기존값 불러오기",
            "missing_count": len(_power_theme_missing(theme)) if theme != "최종 확인·전송" else 0,
            "confirmed_at": _korea_now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        for theme in POWER_THEME_ORDER[:-1]
    }
    # 화면 위젯 shadow 값은 제거하여 불러온 최신 draft 값으로 다시 생성합니다.
    for session_key in list(st.session_state.keys()):
        if session_key.startswith("_ui_power_"):
            del st.session_state[session_key]
    st.session_state["power_panel_nonce"] = int(st.session_state.get("power_panel_nonce", 0) or 0) + 1


def _render_power_cell_inputs(group_number: int) -> list[str]:
    """셀 번호 1~24를 한 줄 최대 5개로 배치하고, 각 입력을 세션 임시저장소에 보존합니다."""
    values: list[str] = []
    for start in range(1, 25, 5):
        row_columns = st.columns(5, gap="small")
        for offset, column in enumerate(row_columns):
            cell_number = start + offset
            if cell_number > 24:
                break
            key = f"power_battery_{group_number}_{cell_number:02d}"
            with column:
                values.append(
                    _power_text_input(
                        str(cell_number),
                        key=key,
                        label_visibility="visible",
                    )
                )
    return values


def _render_power_battery_summary(group_number: int) -> None:
    prefix = f"power_battery{group_number}"
    st.markdown(f"**{group_number}조 측정값**")
    row1 = st.columns(2, gap="small")
    with row1[0]:
        _power_text_input("방전 후 Total 전류 (A)", key=f"{prefix}_total_current")
    with row1[1]:
        _power_text_input("방전 후 Total 전압 (V)", key=f"{prefix}_total_voltage")
    row2 = st.columns(2, gap="small")
    with row2[0]:
        _power_text_input("최저전압 (V)", key=f"{prefix}_min_voltage")
    with row2[1]:
        _power_text_input("최고전압 (V)", key=f"{prefix}_max_voltage")
    _power_text_input("방전종료 전압 (V)", key=f"{prefix}_end_voltage")
    st.markdown(f"**{group_number}조 방전 후 셀 전압 (V)**")
    st.caption("실제 설치된 셀 수만 입력해도 됩니다. 예: 10셀만 측정한 경우 1~10번까지만 입력하고 다음 단계로 진행할 수 있습니다.")
    _render_power_cell_inputs(group_number)


POWER_THEME_ORDER = ["전압·전류 측정", "축전지 측정", "접지저항 측정", "최종 확인·전송"]
POWER_THEME_ICON = {
    "전압·전류 측정": "⚡",
    "축전지 측정": "🔋",
    "접지저항 측정": "🛡️",
    "최종 확인·전송": "📤",
}


def _power_state_blank(key: str) -> bool:
    value = _power_get(key, "")
    return value is None or (isinstance(value, str) and not value.strip())


def _clear_power_measurements_after_station_change() -> None:
    """국소가 바뀌어도 측정값을 삭제하지 않고 현재값을 보존합니다.

    과거에는 이 함수가 power_draft까지 초기화하여, 최종 확인 단계에서
    기본정보를 보완하면 모든 측정값이 사라지는 문제가 있었습니다.
    이제 측정값 초기화는 최종 전송 성공 후 `_reset_power_inspection()`에서만 수행합니다.
    """
    st.session_state["power_station_search_applied"] = False
    st.session_state["power_station_search_notice"] = ""
    _preserve_current_power_measurements()
    _clear_power_history_state()
    _mark_power_basic_info_changed()


def _power_theme_missing(theme: str) -> list[str]:
    phase_type = _power_get("power_phase_type", "삼상")
    if theme == "전압·전류 측정":
        if phase_type == "삼상":
            checks = [
                ("power_three_voltage_rs", "R-S 전압"),
                ("power_three_voltage_st", "S-T 전압"),
                ("power_three_voltage_tr", "T-R 전압"),
                ("power_three_voltage_rn", "R-N 전압"),
                ("power_three_current_r", "R상 전류"),
                ("power_three_current_s", "S상 전류"),
                ("power_three_current_t", "T상 전류"),
                ("power_three_current_n", "N상 전류"),
            ]
        else:
            checks = [
                ("power_single_voltage", "단상 전압"),
                ("power_single_current", "단상 전류"),
            ]
        return [label for key, label in checks if _power_state_blank(key)]

    if theme == "축전지 측정":
        checks = [
            ("power_battery1_total_current", "1조 방전 후 Total 전류"),
            ("power_battery1_total_voltage", "1조 방전 후 Total 전압"),
            ("power_battery1_min_voltage", "1조 최저전압"),
            ("power_battery1_max_voltage", "1조 최고전압"),
            ("power_battery1_end_voltage", "1조 방전종료 전압"),
        ]
        missing = [label for key, label in checks if _power_state_blank(key)]
        battery1_values = [_power_get(f"power_battery_1_{index:02d}", "") for index in range(1, 25)]
        battery1_count = _measured_cell_count(battery1_values)
        missing.extend(
            f"1조 {index}셀" for index in range(1, battery1_count + 1)
            if _power_state_blank(f"power_battery_1_{index:02d}")
        )
        if _power_battery2_enabled():
            checks2 = [
                ("power_battery2_total_current", "2조 방전 후 Total 전류"),
                ("power_battery2_total_voltage", "2조 방전 후 Total 전압"),
                ("power_battery2_min_voltage", "2조 최저전압"),
                ("power_battery2_max_voltage", "2조 최고전압"),
                ("power_battery2_end_voltage", "2조 방전종료 전압"),
            ]
            missing.extend(label for key, label in checks2 if _power_state_blank(key))
            battery2_values = [_power_get(f"power_battery_2_{index:02d}", "") for index in range(1, 25)]
            battery2_count = _measured_cell_count(battery2_values)
            missing.extend(
                f"2조 {index}셀" for index in range(1, battery2_count + 1)
                if _power_state_blank(f"power_battery_2_{index:02d}")
            )
        return missing

    if theme == "접지저항 측정":
        checks = [
            ("power_security_ground_1", "보안접지 1종"),
            ("power_security_ground_2", "보안접지 2종"),
            ("power_security_ground_3", "보안접지 3종"),
            ("power_telecom_ground", "통신접지(메인)"),
            ("power_lightning_ground", "피뢰침접지"),
        ]
        return [label for key, label in checks if _power_state_blank(key)]
    return []


def _power_theme_started(theme: str) -> bool:
    if theme == "전압·전류 측정":
        keys = _power_theme_keys(theme)[1:]
        return any(not _power_state_blank(key) for key in keys)
    if theme == "축전지 측정":
        keys = [key for key in _power_theme_keys(theme) if key.startswith("power_battery") and key not in {"power_battery_set", "power_battery2_enabled"}]
        return any(not _power_state_blank(key) for key in keys)
    if theme == "접지저항 측정":
        return any(not _power_state_blank(key) for key in _power_theme_keys(theme))
    return False


def _power_unlocked_theme_index() -> int:
    """모든 측정 테마는 언제든 선택할 수 있으므로 마지막 인덱스를 반환합니다."""
    return len(POWER_THEME_ORDER) - 1


def _bump_power_panel_nonce() -> None:
    st.session_state["power_panel_nonce"] = int(st.session_state.get("power_panel_nonce", 0) or 0) + 1


def _clear_power_completion_prompt() -> None:
    for key in (
        "power_completion_prompt_theme", "power_completion_answer",
        "power_completion_validation_error", "power_battery2_measure_answer",
        "power_battery_move_error", "power_battery_exit_stage",
    ):
        st.session_state.pop(key, None)


def _move_to_power_theme(target_theme: str) -> None:
    if target_theme not in POWER_THEME_ORDER:
        return
    _hydrate_power_theme_from_draft(target_theme)
    st.session_state["power_current_theme"] = target_theme
    st.session_state["power_temp_saved_notice"] = True
    st.session_state.pop("power_pending_theme_switch", None)
    st.session_state.pop("power_pending_from_theme", None)
    st.session_state.pop("power_navigation_error", None)
    _clear_power_completion_prompt()
    _bump_power_panel_nonce()


def _activate_power_theme(target_theme: str) -> None:
    """측정 순서를 강제하지 않고, 다른 메뉴로 이동하기 전에 현재값 확인을 요청합니다."""
    if target_theme not in POWER_THEME_ORDER:
        return
    current_theme = st.session_state.get("power_current_theme", POWER_THEME_ORDER[0])
    if target_theme == current_theme:
        _hydrate_power_theme_from_draft(current_theme)
        return

    if current_theme in POWER_THEME_ORDER[:-1]:
        _save_power_theme_to_draft(current_theme)
        st.session_state["power_pending_from_theme"] = current_theme
        st.session_state["power_pending_theme_switch"] = target_theme
        st.session_state.pop("power_navigation_error", None)
        return

    _move_to_power_theme(target_theme)


def _confirm_power_theme_switch() -> None:
    """직접 메뉴 이동은 현재값만 임시저장하고 완료 상태는 변경하지 않습니다."""
    from_theme = st.session_state.get("power_pending_from_theme")
    target_theme = st.session_state.get("power_pending_theme_switch")
    if from_theme not in POWER_THEME_ORDER[:-1] or target_theme not in POWER_THEME_ORDER:
        st.session_state["power_navigation_error"] = "이동할 측정 메뉴를 다시 선택해 주세요."
        return

    _save_power_theme_to_draft(from_theme)
    st.session_state["power_temp_saved_notice"] = True
    _move_to_power_theme(target_theme)


def _cancel_power_theme_switch() -> None:
    current_theme = st.session_state.get("power_current_theme", POWER_THEME_ORDER[0])
    if current_theme in POWER_THEME_ORDER[:-1]:
        _save_power_theme_to_draft(current_theme)
    st.session_state.pop("power_pending_theme_switch", None)
    st.session_state.pop("power_pending_from_theme", None)
    st.session_state.pop("power_navigation_error", None)


def _finish_battery_navigation(measure_second_group: bool) -> None:
    target_theme = st.session_state.get("power_battery_navigation_target")
    _save_power_theme_to_draft("축전지 측정")
    if measure_second_group:
        _power_set("power_battery2_enabled", True)
        _power_set("power_battery_set", "2조 셀 측정")
        st.session_state["power_current_theme"] = "축전지 측정"
        _hydrate_power_theme_from_draft("축전지 측정")
        st.session_state["power_temp_saved_notice"] = True
    elif target_theme in POWER_THEME_ORDER:
        _move_to_power_theme(target_theme)
    st.session_state.pop("power_battery_navigation_target", None)
    st.session_state.pop("power_battery_exit_stage", None)
    _bump_power_panel_nonce()


def _request_power_completion(theme: str) -> None:
    if theme not in POWER_THEME_ORDER[:-1]:
        return
    _save_power_theme_to_draft(theme)
    st.session_state["power_completion_prompt_theme"] = theme
    st.session_state.pop("power_completion_answer", None)
    st.session_state.pop("power_completion_validation_error", None)


def _cancel_power_completion() -> None:
    current_theme = st.session_state.get("power_current_theme", POWER_THEME_ORDER[0])
    if current_theme in POWER_THEME_ORDER:
        _save_power_theme_to_draft(current_theme)
    _clear_power_completion_prompt()


def _next_power_theme_after_completion(theme: str) -> str:
    """현재 테마 다음부터 순환하며 아직 완료하지 않은 측정 테마를 찾습니다."""
    measurement_themes = POWER_THEME_ORDER[:-1]
    confirmations = dict(st.session_state.get("power_theme_confirmations", {}))
    if theme not in measurement_themes:
        return POWER_THEME_ORDER[-1]

    current_index = measurement_themes.index(theme)
    for offset in range(1, len(measurement_themes) + 1):
        candidate = measurement_themes[(current_index + offset) % len(measurement_themes)]
        if candidate not in confirmations:
            return candidate
    return POWER_THEME_ORDER[-1]


def _mark_power_theme_complete(theme: str, answer_note: str = "담당자 측정 완료 확인") -> None:
    if theme not in POWER_THEME_ORDER[:-1]:
        return
    _save_power_theme_to_draft(theme)
    confirmations = dict(st.session_state.get("power_theme_confirmations", {}))
    confirmations[theme] = {
        "answer": answer_note,
        "missing_count": len(_power_theme_missing(theme)),
        "confirmed_at": _korea_now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    st.session_state["power_theme_confirmations"] = confirmations
    target_theme = _next_power_theme_after_completion(theme)
    _move_to_power_theme(target_theme)


def _complete_current_power_theme() -> None:
    """담당자가 현재 테마의 측정 완료를 명시적으로 확정합니다."""
    theme = st.session_state.get("power_current_theme", POWER_THEME_ORDER[0])
    if theme not in POWER_THEME_ORDER[:-1]:
        return
    _save_power_theme_to_draft(theme)

    if (
        theme == "축전지 측정"
        and _power_get("power_battery_set", "1조 셀 측정") == "1조 셀 측정"
        and not _power_battery2_enabled()
    ):
        st.session_state["power_battery_exit_stage"] = "ask_group2_complete"
        st.session_state["power_temp_saved_notice"] = True
        return

    _mark_power_theme_complete(theme)


def _finish_battery_completion(measure_second_group: bool) -> None:
    """축전지 1조 완료 후 2조 측정 여부를 처리합니다."""
    _save_power_theme_to_draft("축전지 측정")
    if measure_second_group:
        _power_set("power_battery2_enabled", True)
        _power_set("power_battery_set", "2조 셀 측정")
        st.session_state["power_current_theme"] = "축전지 측정"
        st.session_state["power_temp_saved_notice"] = True
        st.session_state.pop("power_battery_exit_stage", None)
        _hydrate_power_theme_from_draft("축전지 측정")
        _bump_power_panel_nonce()
        return

    _power_set("power_battery2_enabled", False)
    _power_set("power_battery_set", "1조 셀 측정")
    st.session_state.pop("power_battery_exit_stage", None)
    _mark_power_theme_complete("축전지 측정", answer_note="1조 완료·2조 미측정 확인")

def _process_power_completion() -> bool:
    theme = st.session_state.get("power_completion_prompt_theme")
    if theme not in POWER_THEME_ORDER[:-1]:
        return False
    answer = st.session_state.get("power_completion_answer")
    if answer not in {"예", "아니오"}:
        st.session_state["power_completion_validation_error"] = "‘예’ 또는 ‘아니오’를 선택해 주세요."
        return False

    _save_power_theme_to_draft(theme)

    # '예'는 누락값이 있으므로 현재 테마를 유지하며 추가 입력합니다.
    if answer == "예":
        st.session_state["power_temp_saved_notice"] = True
        _clear_power_completion_prompt()
        return True

    # 1조 입력을 마친 경우에는 2조 측정 여부를 한 번 더 확인합니다.
    if theme == "축전지 측정":
        selected_group = 1 if _power_get("power_battery_set", "1조 셀 측정") == "1조 셀 측정" else 2
        if selected_group == 1 and not _power_battery2_enabled():
            st.session_state.pop("power_completion_prompt_theme", None)
            st.session_state.pop("power_completion_answer", None)
            st.session_state.pop("power_completion_validation_error", None)
            st.session_state["power_battery_exit_stage"] = "ask_group2"
            return True

    _mark_power_theme_complete(theme)
    return True


def _process_battery2_measure_confirmation() -> bool:
    answer = st.session_state.get("power_battery2_measure_answer")
    if answer not in {"예", "아니오"}:
        st.session_state["power_battery_move_error"] = "2조 축전지 측정 여부를 선택해 주세요."
        return False

    _save_power_theme_to_draft("축전지 측정")
    if answer == "예":
        _power_set("power_battery2_enabled", True)
        _power_set("power_battery_set", "2조 셀 측정")
        st.session_state["power_current_theme"] = "축전지 측정"
        st.session_state["power_temp_saved_notice"] = True
        _clear_power_completion_prompt()
        _bump_power_panel_nonce()
        return True

    _power_set("power_battery2_enabled", False)
    _power_set("power_battery_set", "1조 셀 측정")
    _mark_power_theme_complete("축전지 측정", answer_note="2조 미측정 확인")
    return True


def _build_power_payload_from_state(final_confirmed: bool = False) -> dict:
    current_theme = st.session_state.get("power_current_theme", POWER_THEME_ORDER[0])
    _save_power_theme_to_draft(current_theme)
    phase_type = _power_get("power_phase_type", "삼상")
    group_count = 2 if _power_battery2_enabled() else 1
    worker = str(st.session_state.get("power_worker", "")).strip()
    battery1_cells = [
        _parse_battery_cell_number(_power_get(f"power_battery_1_{index:02d}", ""))
        for index in range(1, 25)
    ]
    battery2_cells = [
        _parse_battery_cell_number(_power_get(f"power_battery_2_{index:02d}", ""))
        for index in range(1, 25)
    ] if group_count == 2 else [""] * 24
    return {
        "worker": worker,
        "inspector_group": st.session_state.get("power_inspector_group", "") or _inspector_group_for_area(st.session_state.get("power_major_area", "")) or _inspector_group_for_name(worker),
        "major_area": st.session_state.get("power_major_area", "권역 선택"),
        "mother": st.session_state.get("power_mother", "모국 선택"),
        "local": st.session_state.get("power_local", "국소 선택"),
        "phase_type": phase_type,
        "battery_group_count": group_count,
        "source_inspection_id": st.session_state.get("power_loaded_source_id", ""),
        "source_saved_at": st.session_state.get("power_loaded_source_saved_at", ""),
        "three_voltage_rs": _parse_power_number(_power_get("power_three_voltage_rs", ""), 1),
        "three_voltage_st": _parse_power_number(_power_get("power_three_voltage_st", ""), 1),
        "three_voltage_tr": _parse_power_number(_power_get("power_three_voltage_tr", ""), 1),
        "three_voltage_rn": _parse_power_number(_power_get("power_three_voltage_rn", ""), 1),
        "three_current_r": _parse_power_number(_power_get("power_three_current_r", ""), 1),
        "three_current_s": _parse_power_number(_power_get("power_three_current_s", ""), 1),
        "three_current_t": _parse_power_number(_power_get("power_three_current_t", ""), 1),
        "three_current_n": _parse_power_number(_power_get("power_three_current_n", ""), 1),
        "single_voltage": _parse_power_number(_power_get("power_single_voltage", ""), 1),
        "single_current": _parse_power_number(_power_get("power_single_current", ""), 1),
        "battery1_total_current": _parse_power_number(_power_get("power_battery1_total_current", ""), 1),
        "battery1_total_voltage": _parse_power_number(_power_get("power_battery1_total_voltage", ""), 2),
        "battery1_min_voltage": _parse_power_number(_power_get("power_battery1_min_voltage", ""), 2),
        "battery1_max_voltage": _parse_power_number(_power_get("power_battery1_max_voltage", ""), 2),
        "battery1_end_voltage": _parse_power_number(_power_get("power_battery1_end_voltage", ""), 2),
        "battery1_cell_count": _measured_cell_count(battery1_cells),
        "battery1_cells": battery1_cells,
        "battery2_total_current": _parse_power_number(_power_get("power_battery2_total_current", ""), 1) if group_count == 2 else "",
        "battery2_total_voltage": _parse_power_number(_power_get("power_battery2_total_voltage", ""), 2) if group_count == 2 else "",
        "battery2_min_voltage": _parse_power_number(_power_get("power_battery2_min_voltage", ""), 2) if group_count == 2 else "",
        "battery2_max_voltage": _parse_power_number(_power_get("power_battery2_max_voltage", ""), 2) if group_count == 2 else "",
        "battery2_end_voltage": _parse_power_number(_power_get("power_battery2_end_voltage", ""), 2) if group_count == 2 else "",
        "battery2_cell_count": _measured_cell_count(battery2_cells) if group_count == 2 else 0,
        "battery2_cells": battery2_cells,
        "security_ground_1": _parse_power_number(_power_get("power_security_ground_1", ""), 2),
        "security_ground_2": _parse_power_number(_power_get("power_security_ground_2", ""), 2),
        "security_ground_3": _parse_power_number(_power_get("power_security_ground_3", ""), 2),
        "telecom_ground": _parse_power_number(_power_get("power_telecom_ground", ""), 2),
        "lightning_ground": _parse_power_number(_power_get("power_lightning_ground", ""), 2),
        "notes": _power_get("power_notes", ""),
        "final_confirmed": bool(final_confirmed),
    }


def _reset_power_inspection() -> None:
    for key in list(st.session_state.keys()):
        if key.startswith("power_") or key.startswith("_ui_power_"):
            del st.session_state[key]
    st.session_state["power_current_theme"] = POWER_THEME_ORDER[0]
    st.session_state["power_unlocked_theme_index"] = len(POWER_THEME_ORDER) - 1
    st.session_state["power_theme_confirmations"] = {}
    st.session_state["power_panel_nonce"] = 0
    st.session_state["power_draft"] = {
        "power_phase_type": "삼상",
        "power_major_area": "권역 선택",
        "power_battery_set": "1조 셀 측정",
        "power_battery2_enabled": False,
    }
    st.session_state["power_phase_type"] = "삼상"
    st.session_state["power_battery_set"] = "1조 셀 측정"


def _render_power_auto_decimal_script() -> None:
    data_field_rules = {
        "power_three_voltage_rs": 1, "power_three_voltage_st": 1,
        "power_three_voltage_tr": 1, "power_three_voltage_rn": 1,
        "power_three_current_r": 1, "power_three_current_s": 1, "power_three_current_t": 1, "power_three_current_n": 1,
        "power_single_voltage": 1, "power_single_current": 1,
        "power_battery1_total_current": 1, "power_battery2_total_current": 1,
        "power_battery1_total_voltage": 2, "power_battery2_total_voltage": 2,
        "power_battery1_min_voltage": 2, "power_battery2_min_voltage": 2,
        "power_battery1_max_voltage": 2, "power_battery2_max_voltage": 2,
        "power_battery1_end_voltage": 2, "power_battery2_end_voltage": 2,
        "power_security_ground_1": 2, "power_security_ground_2": 2,
        "power_security_ground_3": 2, "power_telecom_ground": 2,
        "power_lightning_ground": 2,
    }
    for group in (1, 2):
        for cell_number in range(1, 25):
            data_field_rules[f"power_battery_{group}_{cell_number:02d}"] = {"mode": "battery_cell"}

    # 실제 화면에는 shadow UI key가 렌더링됩니다.
    field_rules = {
        _power_widget_key(data_key): rule
        for data_key, rule in data_field_rules.items()
    }
    rules_json = json.dumps(field_rules, ensure_ascii=False)
    explicit_next_data_keys = {
        "power_security_ground_1": "power_security_ground_2",
        "power_security_ground_2": "power_security_ground_3",
        "power_security_ground_3": "power_telecom_ground",
        "power_telecom_ground": "power_lightning_ground",
    }
    explicit_next_keys = {
        _power_widget_key(current_key): _power_widget_key(next_key)
        for current_key, next_key in explicit_next_data_keys.items()
    }
    next_keys_json = json.dumps(explicit_next_keys, ensure_ascii=False)
    script = r"""
        <script>
        (() => {
          const rules = __POWER_RULES_JSON__;
          const explicitNextKeys = __POWER_NEXT_KEYS_JSON__;
          const FOCUS_STORAGE_KEY = '__power_next_focus_key_v13__';
          const SCROLL_STORAGE_KEY = '__power_scroll_y_v13__';
          const RUNTIME_STORAGE_KEY = '__power_numeric_runtime_v13__';

          // Streamlit rerun이 반복되어도 이전 감시기/이벤트가 누적되지 않도록 먼저 정리합니다.
          try {
            const previousRuntime = window.parent[RUNTIME_STORAGE_KEY];
            if (previousRuntime && typeof previousRuntime.cleanup === 'function') previousRuntime.cleanup();
          } catch (e) {}

          function formatFixed(value, decimals, key) {
            if (value.includes('.')) {
              const pieces = value.split('.', 2);
              const integer = (pieces[0] || '0').replace(/\D/g, '') || '0';
              const fraction = (pieces[1] || '').replace(/\D/g, '').slice(0, decimals).padEnd(decimals, '0');
              return decimals > 0 ? `${integer}.${fraction}` : integer;
            }
            const digits = value.replace(/\D/g, '');
            if (!digits) return '';
            if (decimals <= 0) return digits;

            // 접지저항 현장 입력: 00→0.0, 000→0.00, 0000→00.00
            const isGroundResistance = String(key || '').includes('ground');
            if (isGroundResistance && digits.length === 2) {
              return `${digits.slice(0, 1)}.${digits.slice(1)}`;
            }
            const padded = digits.length <= decimals ? digits.padStart(decimals + 1, '0') : digits;
            return `${padded.slice(0, -decimals)}.${padded.slice(-decimals)}`;
          }

          function formatBatteryCell(value) {
            if (value.includes('.')) {
              const pieces = value.split('.', 2);
              const integer = (pieces[0] || '0').replace(/\D/g, '') || '0';
              const rawFraction = (pieces[1] || '').replace(/\D/g, '');
              const decimals = Math.max(2, Math.min(3, rawFraction.length || 2));
              const fraction = rawFraction.slice(0, decimals).padEnd(decimals, '0');
              return `${integer}.${fraction}`;
            }
            const digits = value.replace(/\D/g, '');
            if (!digits) return '';
            // 참고 시험성적서의 셀 전압은 소수 셋째 자리까지 사용합니다.
            // 215→2.15, 3507→3.507, 000→0.00, 0000→0.000
            const decimals = digits.length >= 4 ? 3 : 2;
            const padded = digits.length <= decimals ? digits.padStart(decimals + 1, '0') : digits;
            return `${padded.slice(0, -decimals)}.${padded.slice(-decimals)}`;
          }

          function formatted(raw, rule, key) {
            let value = String(raw || '').trim().replace(/,/g, '');
            if (!value) return '';
            value = value.replace(/[^0-9.]/g, '');
            if (!value) return '';
            if (rule && typeof rule === 'object' && rule.mode === 'battery_cell') {
              return formatBatteryCell(value);
            }
            return formatFixed(value, Number(rule || 0), key);
          }

          function parentDocument() {
            try { return window.parent.document; } catch (error) { return null; }
          }

          function setReactValue(input, value) {
            const view = input.ownerDocument.defaultView || window.parent;
            const proto = view.HTMLInputElement.prototype;
            const descriptor = Object.getOwnPropertyDescriptor(proto, 'value');
            if (descriptor && descriptor.set) descriptor.set.call(input, value);
            else input.value = value;
            input.dispatchEvent(new view.Event('input', { bubbles: true }));
            input.dispatchEvent(new view.Event('change', { bubbles: true }));
          }

          function isVisible(element) {
            if (!element) return false;
            const view = element.ownerDocument.defaultView || window.parent;
            const style = view.getComputedStyle(element);
            const rect = element.getBoundingClientRect();
            return style.display !== 'none' && style.visibility !== 'hidden'
              && Number(style.opacity || 1) !== 0 && rect.width > 0 && rect.height > 0;
          }

          function wrapperForKey(doc, key) {
            return doc.querySelector(`div.st-key-${key}`);
          }

          function visibleMeasurementInputs(doc) {
            const found = [];
            Object.keys(rules).forEach((key) => {
              const wrapper = wrapperForKey(doc, key);
              const input = wrapper ? wrapper.querySelector('input') : null;
              if (input && isVisible(input)) found.push({ key, input });
            });
            const NodeCtor = doc.defaultView.Node;
            return found.sort((a, b) => {
              const pos = a.input.compareDocumentPosition(b.input);
              if (pos & NodeCtor.DOCUMENT_POSITION_FOLLOWING) return -1;
              if (pos & NodeCtor.DOCUMENT_POSITION_PRECEDING) return 1;
              return 0;
            });
          }

          function rememberViewport() {
            try {
              const parentWindow = window.parent;
              parentWindow.sessionStorage.setItem(SCROLL_STORAGE_KEY, String(parentWindow.scrollY || 0));
            } catch (e) {}
          }

          function restoreViewport(lockDuration = 520) {
            let saved = null;
            try {
              const raw = window.parent.sessionStorage.getItem(SCROLL_STORAGE_KEY);
              if (raw !== null && raw !== '') saved = Number(raw);
            } catch (e) {}
            if (!Number.isFinite(saved)) return;

            const parentWindow = window.parent;
            const restore = () => parentWindow.scrollTo({ top: saved, left: 0, behavior: 'auto' });
            restore();
            const started = Date.now();
            const timer = parentWindow.setInterval(() => {
              restore();
              if (Date.now() - started >= lockDuration) {
                parentWindow.clearInterval(timer);
                try { parentWindow.sessionStorage.removeItem(SCROLL_STORAGE_KEY); } catch (e) {}
              }
            }, 24);
          }

          function rememberAndFocusNext(doc, input) {
            const currentKey = String(input.dataset.powerKey || '');
            const explicitNextKey = explicitNextKeys[currentKey] || '';
            let nextItem = null;

            // 보안접지 1종→2종→3종→통신접지→피뢰침접지는
            // 열 배치나 모바일 DOM 순서와 관계없이 지정된 순서로 이동합니다.
            if (explicitNextKey) {
              const explicitWrapper = wrapperForKey(doc, explicitNextKey);
              const explicitInput = explicitWrapper ? explicitWrapper.querySelector('input') : null;
              if (explicitInput && isVisible(explicitInput)) {
                nextItem = { key: explicitNextKey, input: explicitInput };
              }
            }

            if (!nextItem) {
              const ordered = visibleMeasurementInputs(doc);
              const currentIndex = ordered.findIndex((item) => item.input === input);
              nextItem = currentIndex >= 0 ? ordered[currentIndex + 1] : null;
            }

            try {
              if (nextItem) window.parent.sessionStorage.setItem(FOCUS_STORAGE_KEY, nextItem.key);
              else window.parent.sessionStorage.removeItem(FOCUS_STORAGE_KEY);
            } catch (e) {}
            return nextItem;
          }

          function prepareNumericInput(input, key) {
            if (!input || !key || !Object.prototype.hasOwnProperty.call(rules, key)) return;
            input.dataset.powerKey = key;
            input.setAttribute('inputmode', 'decimal');
            input.setAttribute('pattern', '[0-9.]*');
            input.setAttribute('autocomplete', 'off');
            input.setAttribute('autocapitalize', 'off');
            input.setAttribute('enterkeyhint', 'next');
            input.spellcheck = false;
          }

          function measurementKeyForInput(input) {
            if (!input) return '';
            for (const key of Object.keys(rules)) {
              const wrapper = wrapperForKey(input.ownerDocument, key);
              if (wrapper && wrapper.contains(input)) return key;
            }
            return '';
          }

          function bindInput(doc, key, rule) {
            const wrapper = wrapperForKey(doc, key);
            const input = wrapper ? wrapper.querySelector('input') : null;
            if (!input || input.dataset.powerDecimalBoundV13 === '1') return;

            input.dataset.powerDecimalBoundV13 = '1';
            prepareNumericInput(input, key);

            const applyFormat = () => {
              const next = formatted(input.value, rule, key);
              if (next !== input.value) setReactValue(input, next);
              return next;
            };

            input.addEventListener('blur', applyFormat, { passive: true });
            input.addEventListener('keydown', (event) => {
              if (event.key !== 'Enter' && event.keyCode !== 13) return;
              event.preventDefault();
              event.stopPropagation();

              // Enter 직전의 화면 위치와 다음 입력키를 먼저 보존합니다.
              rememberViewport();
              const nextItem = rememberAndFocusNext(doc, input);
              applyFormat();

              window.setTimeout(() => {
                restoreViewport();
                if (nextItem && isVisible(nextItem.input)) {
                  nextItem.input.focus({ preventScroll: true });
                  nextItem.input.select();
                } else {
                  input.blur();
                }
              }, 35);
            }, true);
          }

          function bindInputs() {
            const doc = parentDocument();
            if (!doc) return;
            Object.entries(rules).forEach(([key, rule]) => bindInput(doc, key, rule));
          }

          function restoreNextFocus() {
            const doc = parentDocument();
            if (!doc) return;
            let nextKey = '';
            try { nextKey = window.parent.sessionStorage.getItem(FOCUS_STORAGE_KEY) || ''; } catch (e) {}
            if (!nextKey) return;
            const wrapper = wrapperForKey(doc, nextKey);
            const input = wrapper ? wrapper.querySelector('input') : null;
            if (input && isVisible(input)) {
              restoreViewport();
              input.focus({ preventScroll: true });
              input.select();
              try { window.parent.sessionStorage.removeItem(FOCUS_STORAGE_KEY); } catch (e) {}
            }
          }

          bindInputs();
          restoreNextFocus();

          const doc = parentDocument();
          let observer = null;
          let timer = null;
          let prepareFromEvent = null;

          if (doc) {
            prepareFromEvent = (event) => {
              const input = event && event.target;
              if (!input || String(input.tagName || '').toLowerCase() !== 'input') return;
              const key = measurementKeyForInput(input);
              if (key) prepareNumericInput(input, key);
            };
            // 모바일에서 사용자가 새로 열린 2조 셀을 즉시 눌러도 포커스 전에 숫자키패드 속성을 먼저 적용합니다.
            doc.addEventListener('pointerdown', prepareFromEvent, true);
            doc.addEventListener('touchstart', prepareFromEvent, { capture: true, passive: true });
            doc.addEventListener('focusin', prepareFromEvent, true);

            observer = new MutationObserver(() => {
              bindInputs();
              restoreNextFocus();
            });
            observer.observe(doc.body, { childList: true, subtree: true });
          }

          timer = window.setInterval(() => {
            bindInputs();
            restoreNextFocus();
          }, 250);

          const cleanup = () => {
            try { if (observer) observer.disconnect(); } catch (e) {}
            try { if (timer) window.clearInterval(timer); } catch (e) {}
            try {
              if (doc && prepareFromEvent) {
                doc.removeEventListener('pointerdown', prepareFromEvent, true);
                doc.removeEventListener('touchstart', prepareFromEvent, true);
                doc.removeEventListener('focusin', prepareFromEvent, true);
              }
            } catch (e) {}
          };
          try { window.parent[RUNTIME_STORAGE_KEY] = { cleanup }; } catch (e) {}
          window.setTimeout(cleanup, 120000);
        })();
        </script>
    """
    rendered_script = (
        script
        .replace("__POWER_RULES_JSON__", rules_json)
        .replace("__POWER_NEXT_KEYS_JSON__", next_keys_json)
    )
    components.html(rendered_script, height=1)


# ==========================================
# 8-3. MY WORK LOG · 현장 기록 / 시설 이력 · V14
#      - 기존 전원 정밀점검 로직과 완전히 분리
#      - 텍스트/상태이력: Google Sheets
#      - 사진: Google Drive/Shared Drive (선택 설정)
#      - 원본 사진은 앱 서버에 영구 저장하지 않음
# ==========================================
WORK_LOG_SPREADSHEET_NAME = "Audit_Result_2026"
WORK_LOG_SHEET_NAME = "MY_WORK_LOG"
WORK_LOG_HISTORY_SHEET_NAME = "MY_WORK_LOG_HISTORY"
WORK_LOG_USER_SHEET_NAME = "MY_WORK_LOG_USERS"
WORK_LOG_DELETE_AUDIT_SHEET_NAME = "MY_WORK_LOG_DELETED"
WORK_LOG_MAX_PHOTOS = 10
WORK_LOG_IMAGE_MAX_SIDE = 1600
WORK_LOG_IMAGE_TARGET_BYTES = 450 * 1024
WORK_LOG_PIN_ITERATIONS = 210_000

# 기존 시트 열 순서를 절대 바꾸지 않고, 신규 권한 필드는 맨 뒤에 추가합니다.
WORK_LOG_HEADERS = [
    "저장일시", "기록ID", "작성자", "권역", "모국", "국소", "상태", "점검항목",
    "현상_특이사항", "조치내용", "후속조치", "비고", "사진수", "사진파일ID목록",
    "사진파일명목록", "최근수정일시", "작성자ID", "공개범위",
]
WORK_LOG_HISTORY_HEADERS = [
    "저장일시", "기록ID", "작성자", "상태", "변경구분", "조치내용", "후속조치", "비고", "작업자ID",
]
WORK_LOG_USER_HEADERS = [
    "사용자ID", "이름", "사번", "PIN_SALT", "PIN_HASH", "PIN변경필요", "활성", "최근로그인", "최근PIN변경일시",
]
WORK_LOG_DELETE_AUDIT_HEADERS = [
    "삭제일시", "기록ID", "작성자ID", "작성자", "삭제자ID", "삭제자", "공개범위", "사진수",
]
WORK_LOG_STATUS_OPTIONS = ["신규", "확인필요", "조치중", "재점검", "완료"]
WORK_LOG_ITEM_OPTIONS = ["전원", "축전지", "접지", "냉방", "출입", "안전", "기타"]
WORK_LOG_VISIBILITY_OPTIONS = ["공개", "비공개"]

# 최초 1회 로그인 공통 임시 PIN은 000000입니다.
# 평문 PIN 자체는 사용자 시트에 저장하지 않고, 사용자별 salt/PBKDF2-SHA256 해시만 저장합니다.
# 최초 로그인 후 반드시 본인만 아는 6자리 개인 PIN으로 변경해야 계속 사용할 수 있습니다.
WORK_LOG_INITIAL_PIN = "000000"
WORK_LOG_USER_BOOTSTRAP = [
    {"사용자ID": "U001", "이름": "정청운", "사번": "10001713", "PIN_SALT": "de28394a671befb76a8fd8ec1b904d72", "PIN_HASH": "8fb388bf36dd783aeda3bbfd362b5b035df14ef1d11b64bb078085e521a4f295"},
    {"사용자ID": "U002", "이름": "이학원", "사번": "10001612", "PIN_SALT": "2fcd8fad516f3d7b0719729fc9f60996", "PIN_HASH": "6d7041143724825e28f813004029a15caca9a580635878eb8c8a1e82217b7e19"},
    {"사용자ID": "U003", "이름": "이철순", "사번": "10002090", "PIN_SALT": "c5e1143a554e05ec9051b537281b9a1a", "PIN_HASH": "7aec4643550243e931bdaaa4ea931f3ed6ea0dcf4a53555971a0289d7ce15bae"},
    {"사용자ID": "U004", "이름": "소순고", "사번": "81000020", "PIN_SALT": "3d2a67162bfd799901747c7e5e716e5e", "PIN_HASH": "cee9feacb954b53343986b93c87b8d3fccab7d1fc6492a879549af01a5dab886"},
    {"사용자ID": "U005", "이름": "강만식", "사번": "10001009", "PIN_SALT": "7bfa69aaf3b2a6ece51c5fc5ac25d0d2", "PIN_HASH": "45c73984e01f3d857682124f4ce167feb4e4f2ed8178dda4fbeb05dc34eeec97"},
    {"사용자ID": "U006", "이름": "이민우", "사번": "10001522", "PIN_SALT": "18a1a33b692c2a2142e1eabe56cdd6ff", "PIN_HASH": "cc667aeef9fe18ad0d7772cef1ee26091ad8fb311c102cc8c8691f25a7e927f7"},
    {"사용자ID": "U007", "이름": "신진우", "사번": "10001405", "PIN_SALT": "00a1ec0138defc16e5845bb54ddc3c6d", "PIN_HASH": "4daf3b870e8f289b15a5b59a933fed897141858405724bd48efe84fd6de1772b"},
    {"사용자ID": "U008", "이름": "박동희", "사번": "10001280", "PIN_SALT": "27dd2fbd801a3cc0d8cc53b9436e11f6", "PIN_HASH": "9bf8757c5b114e3b051afece0bbb5fd41adc4acbab696d5f810552fd0525b4a7"},
    {"사용자ID": "U009", "이름": "김태수", "사번": "10001923", "PIN_SALT": "02bd5b1d4eedb7103955495269126453", "PIN_HASH": "6bf64d324a844f33de7b04203dbcfb3c42d6cb44785a0c6680eefc3b8462fceb"},
    {"사용자ID": "U010", "이름": "김수창", "사번": "10002211", "PIN_SALT": "41b790666a4e7c4654df4c3bb57d3680", "PIN_HASH": "c1ef17186b3dea51c524663a11fec23d7a65b4b2f898123b513afd4e35aed741"},
]
WORK_LOG_NAME_TO_USER_ID = {row["이름"]: row["사용자ID"] for row in WORK_LOG_USER_BOOTSTRAP}
WORK_LOG_EMPLOYEE_TO_NAME = {row["사번"]: row["이름"] for row in WORK_LOG_USER_BOOTSTRAP}


def _worklog_area_display(area: str) -> str:
    """WORK LOG에서 권역을 담당자 + 주요 지역이 함께 보이는 현장형 표기로 변환합니다."""
    area_value = str(area or "").strip()
    if area_value not in POWER_REGION_DATA:
        return "국사를 검색하면 자동 표시됩니다"

    region = POWER_REGION_DATA.get(area_value, {})
    inspectors = ", ".join(
        str(person).strip()
        for person in region.get("담당자", [])
        if str(person).strip()
    )

    # 예: "1권역 · 파주·문산·동두천 등" -> "1권역: 이철순, 김수창 (파주, 문산, 동두천 등)"
    if "·" in area_value:
        area_no, coverage = area_value.split("·", 1)
        coverage_text = ", ".join(
            part.strip() for part in coverage.split("·") if part.strip()
        )
    else:
        area_no = area_value
        coverage_text = ""

    if inspectors and coverage_text:
        return f"{area_no.strip()}: {inspectors} ({coverage_text})"
    if inspectors:
        return f"{area_no.strip()}: {inspectors}"
    return area_value


def _worklog_station_search_label(entry_id: str) -> str:
    """동일/유사 국사가 여러 곳일 때 WORK LOG 선택 후보를 알아보기 쉽게 표시합니다."""
    entry = POWER_STATION_SEARCH_BY_ID.get(str(entry_id or ""), {})
    if not entry:
        return "검색 결과 없음"
    area_display = _worklog_area_display(entry.get("area", ""))
    return (
        f"{entry.get('local', '')}  |  모국 {entry.get('mother', '')}  |  {area_display}"
    )


def _apply_worklog_station_search_entry(entry: dict) -> None:
    """선택한 국사의 권역·모국·국소를 WORK LOG 전용 상태에 자동 반영합니다."""
    if not entry:
        return

    selected_area = str(entry.get("area", "")).strip()
    st.session_state["worklog_area_key"] = selected_area
    st.session_state["worklog_mother"] = str(entry.get("mother", "")).strip()
    st.session_state["worklog_local"] = str(entry.get("local", "")).strip()
    st.session_state["worklog_station_search_applied"] = True
    st.session_state["worklog_station_search_candidates"] = []
    st.session_state["worklog_station_search_status"] = "applied"
    st.session_state["worklog_station_search_notice"] = (
        f"✅ {entry.get('local', '')} 국사 선택 완료 · "
        f"{_worklog_area_display(selected_area)} · 모국 {entry.get('mother', '')}"
    )


def _run_worklog_station_search() -> None:
    """정밀점검과 같은 국사 역색인을 사용해 WORK LOG 국사를 검색합니다."""
    query = str(st.session_state.get("worklog_station_search_query", "") or "").strip()
    st.session_state["worklog_station_search_notice"] = ""
    st.session_state["worklog_station_search_status"] = ""
    st.session_state["worklog_station_search_candidates"] = []
    st.session_state["worklog_station_search_choice"] = ""

    if not query:
        st.session_state["worklog_station_search_status"] = "empty"
        return

    normalized_query = _normalize_power_station_search(query)
    matches = _search_power_station_entries(query)
    exact_matches = [
        entry for entry in matches
        if _normalize_power_station_search(entry.get("local", "")) == normalized_query
    ]
    candidates = exact_matches if exact_matches else matches

    if len(candidates) == 1:
        _apply_worklog_station_search_entry(candidates[0])
        return

    if len(candidates) > 1:
        candidate_ids = [entry["id"] for entry in candidates]
        st.session_state["worklog_station_search_candidates"] = candidate_ids
        st.session_state["worklog_station_search_choice"] = candidate_ids[0]
        st.session_state["worklog_station_search_status"] = "multiple"
        st.session_state["worklog_station_search_applied"] = False
        return

    st.session_state["worklog_station_search_status"] = "none"
    st.session_state["worklog_station_search_applied"] = False


def _confirm_worklog_station_search_choice() -> None:
    """WORK LOG 검색 후보 중 사용자가 고른 국사를 확정합니다."""
    selected_id = str(st.session_state.get("worklog_station_search_choice", "") or "").strip()
    entry = POWER_STATION_SEARCH_BY_ID.get(selected_id)
    if not entry:
        st.session_state["worklog_station_search_status"] = "choice_required"
        return
    _apply_worklog_station_search_entry(entry)


def _worklog_secret_value(name: str, default=""):
    """WORK LOG 전용 Secrets를 평면/섹션 형식 모두에서 안전하게 읽습니다."""
    try:
        direct = st.secrets.get(name, default)
        if direct not in (None, ""):
            return direct
    except Exception:
        pass
    try:
        section = st.secrets.get("work_log", {})
        if hasattr(section, "get"):
            short_name = name.replace("work_log_", "")
            value = section.get(short_name, default)
            if value not in (None, ""):
                return value
    except Exception:
        pass
    return default



def _worklog_hash_pin(pin: str, salt_hex: str) -> str:
    """개인 PIN을 PBKDF2-SHA256으로 해시합니다. PIN 평문은 시트/코드에 저장하지 않습니다."""
    try:
        salt = bytes.fromhex(str(salt_hex or "").strip())
    except Exception:
        return ""
    if not salt:
        return ""
    return hashlib.pbkdf2_hmac(
        "sha256",
        str(pin or "").encode("utf-8"),
        salt,
        WORK_LOG_PIN_ITERATIONS,
    ).hex()


def _worklog_ensure_headers(worksheet, desired_headers: list[str]) -> list[str]:
    """기존 열 순서를 보존한 채 필요한 헤더만 맨 뒤에 추가합니다."""
    try:
        headers = [str(value).strip() for value in worksheet.row_values(1)]
    except Exception:
        headers = []

    if not headers:
        worksheet.append_row(desired_headers, value_input_option="USER_ENTERED")
        return list(desired_headers)

    missing = [header for header in desired_headers if header not in headers]
    if missing:
        try:
            needed_cols = len(headers) + len(missing)
            current_cols = int(getattr(worksheet, "col_count", 0) or 0)
            if current_cols < needed_cols:
                worksheet.add_cols(needed_cols - current_cols)
        except Exception:
            pass
        for header in missing:
            headers.append(header)
            worksheet.update_cell(1, len(headers), header)
    return headers


def _worklog_ensure_user_sheet(spreadsheet):
    """개인인증 사용자 시트를 생성하고, 최초 PIN 미변경 계정은 공통 임시 PIN 000000으로 안전하게 동기화합니다."""
    try:
        ws = spreadsheet.worksheet(WORK_LOG_USER_SHEET_NAME)
    except Exception:
        ws = spreadsheet.add_worksheet(
            title=WORK_LOG_USER_SHEET_NAME,
            rows=1000,
            cols=max(len(WORK_LOG_USER_HEADERS) + 2, 12),
        )
        ws.append_row(WORK_LOG_USER_HEADERS, value_input_option="USER_ENTERED")

    headers = _worklog_ensure_headers(ws, WORK_LOG_USER_HEADERS)
    values = ws.get_all_values()
    employee_index = headers.index("사번") if "사번" in headers else None
    existing_rows: dict[str, tuple[int, list[str]]] = {}
    if values and employee_index is not None:
        for row_no, row in enumerate(values[1:], start=2):
            value = str(row[employee_index] if employee_index < len(row) else "" or "").strip()
            if value:
                existing_rows[value] = (row_no, row)

    for bootstrap in WORK_LOG_USER_BOOTSTRAP:
        employee_no = bootstrap["사번"]
        existing = existing_rows.get(employee_no)
        if existing is None:
            row_map = {
                **bootstrap,
                "PIN변경필요": "Y",
                "활성": "Y",
                "최근로그인": "",
                "최근PIN변경일시": "",
            }
            ws.append_row(
                [row_map.get(header, "") for header in headers],
                value_input_option="USER_ENTERED",
            )
            continue

        row_no, row = existing
        row_map = {header: (row[idx] if idx < len(row) else "") for idx, header in enumerate(headers)}
        must_change = str(row_map.get("PIN변경필요", "") or "").strip().upper() in {"Y", "YES", "TRUE", "1"}

        # V12/V13에서 개인별 임시 PIN이 이미 배정되었더라도 아직 본인 PIN으로 바꾸지 않은 계정만
        # V14 공통 최초 PIN 000000으로 전환합니다. 이미 PIN변경필요=N인 사용자의 개인 PIN은 절대 건드리지 않습니다.
        if must_change:
            initial_updates = {
                "PIN_SALT": bootstrap["PIN_SALT"],
                "PIN_HASH": bootstrap["PIN_HASH"],
                "PIN변경필요": "Y",
            }
            for header, value in initial_updates.items():
                if header in headers and str(row_map.get(header, "") or "") != str(value):
                    ws.update_cell(row_no, headers.index(header) + 1, value)

    return ws

def _worklog_read_user_by_employee(employee_no: str) -> tuple[object | None, dict, int | None]:
    """사번으로 사용자 시트의 실제 행을 읽습니다."""
    employee_no = re.sub(r"\D", "", str(employee_no or ""))
    if not employee_no:
        return None, {}, None

    client = init_google_sheet_connection()
    if not client:
        return None, {}, None
    try:
        spreadsheet = client.open(WORK_LOG_SPREADSHEET_NAME)
        ws = _worklog_ensure_user_sheet(spreadsheet)
        values = ws.get_all_values()
        if not values:
            return ws, {}, None
        headers = [str(value).strip() for value in values[0]]
        if "사번" not in headers:
            return ws, {}, None
        employee_index = headers.index("사번")
        for row_no, row in enumerate(values[1:], start=2):
            current_no = str(row[employee_index] if employee_index < len(row) else "").strip()
            if current_no == employee_no:
                record = {
                    header: (row[index] if index < len(row) else "")
                    for index, header in enumerate(headers)
                }
                return ws, record, row_no
        return ws, {}, None
    except Exception:
        return None, {}, None


def _worklog_authenticate_user(employee_no: str, pin: str) -> tuple[bool, str, dict]:
    """사번 + 개인 PIN을 검증합니다."""
    employee_no = re.sub(r"\D", "", str(employee_no or ""))
    pin = re.sub(r"\D", "", str(pin or ""))
    if len(employee_no) < 6:
        return False, "사번을 정확히 입력해 주세요.", {}
    if len(pin) != 6:
        return False, "개인 PIN 6자리를 입력해 주세요.", {}

    now_ts = time.time()
    locked_until = float(st.session_state.get("worklog_login_locked_until", 0) or 0)
    if locked_until > now_ts:
        remaining = max(1, int(locked_until - now_ts))
        return False, f"로그인 실패가 반복되어 {remaining}초 후 다시 시도할 수 있습니다.", {}

    ws, record, row_no = _worklog_read_user_by_employee(employee_no)
    if ws is None or not record or row_no is None:
        return False, "등록된 사번을 확인하지 못했습니다.", {}
    if str(record.get("활성", "Y") or "Y").strip().upper() not in {"Y", "YES", "TRUE", "1", "활성"}:
        return False, "현재 사용이 중지된 계정입니다. 관리자에게 문의해 주세요.", {}

    expected_hash = str(record.get("PIN_HASH", "") or "").strip().lower()
    salt_hex = str(record.get("PIN_SALT", "") or "").strip()
    actual_hash = _worklog_hash_pin(pin, salt_hex).lower()
    if not expected_hash or not actual_hash or not hmac.compare_digest(expected_hash, actual_hash):
        failures = int(st.session_state.get("worklog_login_failures", 0) or 0) + 1
        st.session_state["worklog_login_failures"] = failures
        if failures >= 5:
            st.session_state["worklog_login_failures"] = 0
            st.session_state["worklog_login_locked_until"] = time.time() + 60
            return False, "PIN이 일치하지 않습니다. 보안을 위해 60초 동안 로그인을 잠급니다.", {}
        return False, f"사번 또는 PIN이 일치하지 않습니다. ({failures}/5)", {}

    st.session_state["worklog_login_failures"] = 0
    st.session_state["worklog_login_locked_until"] = 0
    user = {
        "user_id": str(record.get("사용자ID", "") or "").strip(),
        "name": str(record.get("이름", "") or "").strip(),
        "employee_no": employee_no,
    }
    if not user["user_id"] or not user["name"]:
        return False, "사용자 등록정보가 올바르지 않습니다.", {}

    try:
        headers = [str(value).strip() for value in ws.row_values(1)]
        if "최근로그인" in headers:
            ws.update_cell(
                row_no,
                headers.index("최근로그인") + 1,
                _korea_now().strftime("%Y-%m-%d %H:%M:%S"),
            )
    except Exception:
        pass

    must_change = str(record.get("PIN변경필요", "N") or "N").strip().upper() in {"Y", "YES", "TRUE", "1"}
    user["must_change_pin"] = must_change
    return True, f"{user['name']}님 인증되었습니다.", user


def _worklog_change_pin(employee_no: str, new_pin: str, confirm_pin: str) -> tuple[bool, str]:
    """현재 사용자의 개인 PIN을 새 6자리 PIN으로 변경합니다."""
    employee_no = re.sub(r"\D", "", str(employee_no or ""))
    new_pin = re.sub(r"\D", "", str(new_pin or ""))
    confirm_pin = re.sub(r"\D", "", str(confirm_pin or ""))
    if len(new_pin) != 6:
        return False, "새 PIN은 숫자 6자리로 입력해 주세요."
    if new_pin != confirm_pin:
        return False, "새 PIN과 확인 PIN이 일치하지 않습니다."
    if new_pin in {"000000", "111111", "123456", "654321", "121212", "777777"}:
        return False, "추측하기 쉬운 PIN은 사용할 수 없습니다."
    if employee_no and (new_pin in employee_no or employee_no.endswith(new_pin)):
        return False, "사번에 포함된 숫자를 그대로 PIN으로 사용하지 마세요."

    ws, record, row_no = _worklog_read_user_by_employee(employee_no)
    if ws is None or not record or row_no is None:
        return False, "사용자 계정을 찾지 못했습니다."

    salt_hex = os.urandom(16).hex()
    pin_hash = _worklog_hash_pin(new_pin, salt_hex)
    if not pin_hash:
        return False, "PIN 보안처리에 실패했습니다."

    try:
        headers = [str(value).strip() for value in ws.row_values(1)]
        updates = {
            "PIN_SALT": salt_hex,
            "PIN_HASH": pin_hash,
            "PIN변경필요": "N",
            "최근PIN변경일시": _korea_now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        for header, value in updates.items():
            if header in headers:
                ws.update_cell(row_no, headers.index(header) + 1, value)
        return True, "개인 PIN이 변경되었습니다."
    except Exception as error:
        return False, f"PIN 변경 실패: {error}"


def _worklog_current_user() -> dict:
    user = st.session_state.get("worklog_auth_user")
    return user if isinstance(user, dict) else {}


def _worklog_logout() -> None:
    """MY WORK LOG 개인 세션만 종료하고 다른 앱 기능에는 영향을 주지 않습니다."""
    for key in (
        "worklog_auth_user", "worklog_pin_change_required", "worklog_show_pin_change",
        "worklog_df", "worklog_loaded_at", "worklog_selected_id", "worklog_delete_pending_id",
        "worklog_search", "worklog_filter", "worklog_public_scope",
    ):
        st.session_state.pop(key, None)


def _worklog_user_id_from_name(name: str) -> str:
    return str(WORK_LOG_NAME_TO_USER_ID.get(str(name or "").strip(), "") or "")


def _worklog_record_owner_id(record) -> str:
    explicit = str(record.get("작성자ID", "") or "").strip()
    if explicit:
        return explicit
    return _worklog_user_id_from_name(record.get("작성자", ""))


def _worklog_record_visibility(record) -> str:
    value = str(record.get("공개범위", "") or "").strip()
    return "비공개" if value == "비공개" else "공개"


def _worklog_record_owned_by(record, auth_user: dict) -> bool:
    if not auth_user:
        return False
    return _worklog_record_owner_id(record) == str(auth_user.get("user_id", "") or "").strip()


def _worklog_filter_accessible_records(df: pd.DataFrame, auth_user: dict) -> pd.DataFrame:
    """공개 기록 + 로그인 사용자의 비공개 기록만 반환합니다. 비공개는 검색 전 단계에서 차단합니다."""
    if not isinstance(df, pd.DataFrame):
        return pd.DataFrame(columns=WORK_LOG_HEADERS)
    if df.empty:
        return df.copy()

    result = df.copy()
    if "작성자ID" not in result.columns:
        result["작성자ID"] = ""
    if "공개범위" not in result.columns:
        result["공개범위"] = ""

    result["작성자ID"] = result.apply(
        lambda row: str(row.get("작성자ID", "") or "").strip() or _worklog_user_id_from_name(row.get("작성자", "")),
        axis=1,
    )
    result["공개범위"] = result["공개범위"].apply(
        lambda value: "비공개" if str(value or "").strip() == "비공개" else "공개"
    )

    current_user_id = str(auth_user.get("user_id", "") or "").strip()
    if not current_user_id:
        return result.iloc[0:0].copy()

    allowed = (result["공개범위"] == "공개") | (
        (result["공개범위"] == "비공개") & (result["작성자ID"] == current_user_id)
    )
    return result[allowed].copy()


def _worklog_login_dialog_body() -> None:
    st.caption("팀 공유 기록과 개인 메모를 구분하기 위해 MY WORK LOG만 개인 인증을 사용합니다. 최초 로그인 공통 PIN은 000000입니다.")
    with st.form("worklog_personal_login_form", clear_on_submit=False):
        employee_no = st.text_input(
            "사번",
            placeholder="사번 입력",
            max_chars=10,
            key="worklog_login_employee_no",
        )
        pin = st.text_input(
            "개인 PIN",
            type="password",
            placeholder="최초 000000 / 이후 개인 PIN",
            max_chars=6,
            key="worklog_login_pin",
        )
        submitted = st.form_submit_button("🔐 로그인", type="primary", use_container_width=True)

    if submitted:
        ok, message, user = _worklog_authenticate_user(employee_no, pin)
        if ok:
            st.session_state["worklog_auth_user"] = {
                "user_id": user["user_id"],
                "name": user["name"],
                "employee_no": user["employee_no"],
            }
            st.session_state["worklog_pin_change_required"] = bool(user.get("must_change_pin"))
            st.session_state["worklog_df"] = None
            st.session_state["worklog_selected_id"] = ""
            st.session_state.pop("worklog_login_pin", None)
            st.success(message)
            time.sleep(0.35)
            st.rerun()
        else:
            st.error(message)


if hasattr(st, "dialog"):
    _worklog_login_dialog = st.dialog(
        "🔐 MY WORK LOG 개인 인증",
    )(_worklog_login_dialog_body)
else:
    _worklog_login_dialog = _worklog_login_dialog_body


def _worklog_drive_folder_id() -> str:
    """하위 호환용 Google Drive 폴더 ID입니다. Apps Script 방식에서는 Script Properties의 FOLDER_ID가 실제 저장 위치를 결정합니다."""
    return str(_worklog_secret_value("work_log_drive_folder_id", "") or "").strip()


def _worklog_photo_upload_url() -> str:
    """MY WORK LOG 사진 업로드용 Google Apps Script 웹 앱(/exec) URL을 반환합니다."""
    return str(_worklog_secret_value("work_log_photo_upload_url", "") or "").strip()


def _worklog_photo_upload_token() -> str:
    """Apps Script와 공유하는 사진 업로드 비밀 토큰을 반환합니다."""
    return str(_worklog_secret_value("work_log_upload_token", "") or "").strip()


WORK_LOG_PHOTO_ENGINE_VERSION = "V5-20260812-1945"


def _worklog_normalize_apps_script_url() -> tuple[str, str]:
    """Apps Script 웹 앱의 영구 /exec URL만 허용합니다.

    ContentService가 반환하는 script.googleusercontent.com 주소는 일회성 응답 URL이므로
    Secrets에 저장하면 이후 404가 발생할 수 있습니다.
    """
    raw_url = _worklog_photo_upload_url().strip().strip('"').strip("'")
    if not raw_url:
        return "", "Streamlit Secrets의 [work_log] photo_upload_url이 비어 있습니다."
    try:
        parsed = urlparse(raw_url)
    except Exception:
        return "", "photo_upload_url을 URL로 해석할 수 없습니다."

    host = str(parsed.netloc or "").lower().split(":")[0]
    path = str(parsed.path or "").rstrip("/")

    if parsed.scheme.lower() != "https":
        return "", "photo_upload_url은 https:// 주소여야 합니다."
    if host == "script.googleusercontent.com" or host.endswith(".script.googleusercontent.com"):
        return "", (
            "photo_upload_url에 Google의 일회성 리디렉션 주소(script.googleusercontent.com)가 들어 있습니다. "
            "Apps Script의 '배포 관리'에서 복사한 https://script.google.com/macros/s/.../exec 원본 주소를 넣어 주세요."
        )
    if host != "script.google.com":
        return "", (
            f"photo_upload_url 호스트가 {host or '확인 불가'}입니다. "
            "Apps Script 웹 앱의 원본 /exec 주소(https://script.google.com/macros/s/.../exec)를 사용해 주세요."
        )
    if path.endswith("/dev"):
        return "", "photo_upload_url이 /dev 개발용 주소입니다. 실제 배포용 /exec 주소를 사용해 주세요."
    if not re.fullmatch(r"/macros/s/[^/]+/exec", path):
        return "", (
            "photo_upload_url 형식이 Apps Script 배포용 /exec 주소와 일치하지 않습니다. "
            "배포 → 배포 관리에서 '웹 앱 URL'을 다시 복사해 주세요."
        )
    return f"https://script.google.com{path}", ""


def _worklog_photo_config_status() -> tuple[bool, list[str]]:
    """사진 업로드 Secrets 상태를 항목별로 진단합니다."""
    issues: list[str] = []

    raw_url = _worklog_photo_upload_url().strip()
    token = _worklog_photo_upload_token().strip()

    if not raw_url:
        issues.append("photo_upload_url 누락")
    else:
        _, url_error = _worklog_normalize_apps_script_url()
        if url_error:
            issues.append(f"photo_upload_url 오류: {url_error}")

    if not token:
        issues.append("upload_token 누락")

    return (len(issues) == 0), issues


def _worklog_photo_upload_ready() -> bool:
    ready, _ = _worklog_photo_config_status()
    return ready


def _worklog_follow_apps_script_response(first_response, timeout: int = 30):
    """Apps Script ContentService의 일회성 리디렉션을 즉시 따라가 최종 응답을 반환합니다."""
    response = first_response
    if first_response.status_code in {301, 302, 303, 307, 308}:
        redirect_url = str(first_response.headers.get("Location", "") or "").strip()
        if not redirect_url:
            return None, "Apps Script 리디렉션에 Location 주소가 없습니다."
        redirect_host = str(urlparse(redirect_url).netloc or "").lower()
        if "script.googleusercontent.com" not in redirect_host:
            return None, f"예상하지 않은 리디렉션 주소입니다: {redirect_host or '확인 불가'}"
        try:
            response = requests.get(
                redirect_url,
                timeout=timeout,
                allow_redirects=True,
                headers={
                    "Accept": "application/json,text/plain,*/*",
                    "User-Agent": "SMART-WORK-AI-AGENT/4.0",
                },
            )
        except Exception as error:
            return None, f"Apps Script 응답 리디렉션 처리 실패: {error}"
    return response, ""


def _worklog_apps_script_healthcheck() -> tuple[bool, str]:
    """Streamlit 서버에서 GET과 실제 POST 경로를 모두 점검합니다.

    브라우저 GET 성공은 로그인 쿠키 때문에 오판할 수 있으므로, 실제 사진 업로드와 같은
    서버 측 POST도 아주 작은 요청으로 확인합니다. POST probe는 사진을 생성하지 않습니다.
    현재 Apps Script doPost는 data가 없으면 '사진 데이터가 없습니다.' JSON을 반환하므로
    이 응답을 POST 경로 정상의 증거로 사용합니다.
    """
    upload_url, url_error = _worklog_normalize_apps_script_url()
    if url_error:
        return False, url_error
    upload_token = _worklog_photo_upload_token()
    if not upload_token:
        return False, "[work_log] upload_token이 설정되지 않았습니다."

    # 1) 익명 GET 확인
    try:
        get_first = requests.get(
            upload_url,
            timeout=20,
            allow_redirects=False,
            headers={
                "Accept": "application/json,text/plain,*/*",
                "User-Agent": "SMART-WORK-AI-AGENT/4.0",
            },
        )
        get_response, redirect_error = _worklog_follow_apps_script_response(get_first, timeout=20)
        if redirect_error:
            return False, f"GET 진단 실패: {redirect_error}"
        if get_response is None or get_response.status_code != 200:
            code = getattr(get_response, "status_code", get_first.status_code)
            return False, (
                f"GET 진단 실패 ({code}). Streamlit 서버가 Apps Script를 익명으로 열 수 없습니다. "
                "웹 앱 배포의 접근 권한이 로그인 없이 허용되는지 확인해야 합니다."
            )
        if get_response.text.lstrip().lower().startswith("<!doctype html"):
            return False, "GET 진단에서 JSON 대신 Google HTML이 반환되었습니다. 익명 접근 또는 배포 URL 문제입니다."
        try:
            get_json = get_response.json()
        except Exception:
            return False, f"GET 진단 응답이 JSON이 아닙니다: {get_response.text[:160]}"
        if not bool(get_json.get("ok")):
            return False, f"GET 진단 실패: {get_json}"
    except requests.Timeout:
        return False, "GET 진단 시간이 초과되었습니다."
    except Exception as error:
        return False, f"GET 진단 오류: {error}"

    # 2) 실제 사진 업로드와 같은 POST 경로 확인. data는 의도적으로 비워 파일을 만들지 않습니다.
    probe_payload = {
        "token": upload_token,
        "filename": "__worklog_probe__.jpg",
        "mimeType": "image/jpeg",
        "data": "",
    }
    probe_bytes = json.dumps(probe_payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
    try:
        post_first = requests.post(
            upload_url,
            data=probe_bytes,
            headers={
                "Content-Type": "text/plain; charset=utf-8",
                "Accept": "application/json,text/plain,*/*",
                "User-Agent": "SMART-WORK-AI-AGENT/4.0",
            },
            timeout=30,
            allow_redirects=False,
        )
        post_response, redirect_error = _worklog_follow_apps_script_response(post_first, timeout=30)
        if redirect_error:
            return False, f"POST 진단 실패: {redirect_error}"
        if post_response is None or post_response.status_code != 200:
            code = getattr(post_response, "status_code", post_first.status_code)
            body = getattr(post_response, "text", post_first.text)[:120]
            return False, (
                f"POST 진단 실패 ({code}). 브라우저 GET은 열려 있어도 Streamlit 서버의 POST가 차단된 상태입니다. "
                f"응답: {body}"
            )
        if post_response.text.lstrip().lower().startswith("<!doctype html"):
            return False, (
                "POST 진단에서 Google HTML 페이지가 반환되었습니다. 이 경우 폴더/토큰 문제가 아니라 "
                "웹 앱의 익명 POST 접근 또는 현재 /exec 배포 버전 문제입니다."
            )
        try:
            post_json = post_response.json()
        except Exception:
            return False, f"POST 진단 응답이 JSON이 아닙니다: {post_response.text[:160]}"

        # 현재 doPost에서 빈 data는 정상적으로 여기까지 도달하면 '사진 데이터가 없습니다.'를 반환합니다.
        if bool(post_json.get("ok")):
            return True, f"사진 연결 정상 · GET/POST 모두 통과 · PHOTO ENGINE {WORK_LOG_PHOTO_ENGINE_VERSION}"
        error_text = str(post_json.get("error", "") or "")
        if error_text == "사진 데이터가 없습니다.":
            return True, f"사진 연결 정상 · GET/POST 모두 통과 · PHOTO ENGINE {WORK_LOG_PHOTO_ENGINE_VERSION}"
        if error_text.lower() == "unauthorized":
            return False, "POST는 Apps Script에 도달했지만 UPLOAD_TOKEN이 일치하지 않습니다."
        return False, f"POST는 Apps Script에 도달했지만 doPost가 오류를 반환했습니다: {error_text or post_json}"
    except requests.Timeout:
        return False, "POST 진단 시간이 초과되었습니다."
    except Exception as error:
        return False, f"POST 진단 오류: {error}"


@st.cache_resource
def _worklog_google_credentials():
    """기존 Google 서비스 계정 정보를 재사용해 Drive API 인증 객체를 만듭니다."""
    if ServiceAccountCredentials is None:
        return None
    try:
        scope = [
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive",
        ]
        return ServiceAccountCredentials.from_json_keyfile_dict(
            st.secrets["gcp_service_account"], scope
        )
    except Exception:
        return None


def _worklog_drive_access_token() -> str:
    creds = _worklog_google_credentials()
    if creds is None:
        return ""
    try:
        token_info = creds.get_access_token()
        return str(getattr(token_info, "access_token", "") or "")
    except Exception:
        return ""


def _worklog_ensure_sheets(spreadsheet):
    """WORK LOG 본문/상태이력 시트를 생성하고 기존 열 순서를 보존한 채 신규 헤더를 보장합니다."""
    try:
        ws = spreadsheet.worksheet(WORK_LOG_SHEET_NAME)
    except Exception:
        ws = spreadsheet.add_worksheet(
            title=WORK_LOG_SHEET_NAME,
            rows=10000,
            cols=max(len(WORK_LOG_HEADERS) + 4, 24),
        )
        ws.append_row(WORK_LOG_HEADERS, value_input_option="USER_ENTERED")

    try:
        history_ws = spreadsheet.worksheet(WORK_LOG_HISTORY_SHEET_NAME)
    except Exception:
        history_ws = spreadsheet.add_worksheet(
            title=WORK_LOG_HISTORY_SHEET_NAME,
            rows=20000,
            cols=max(len(WORK_LOG_HISTORY_HEADERS) + 4, 16),
        )
        history_ws.append_row(WORK_LOG_HISTORY_HEADERS, value_input_option="USER_ENTERED")

    _worklog_ensure_headers(ws, WORK_LOG_HEADERS)
    _worklog_ensure_headers(history_ws, WORK_LOG_HISTORY_HEADERS)
    return ws, history_ws


def _worklog_make_id(now_dt: datetime.datetime | None = None) -> str:
    now_dt = now_dt or _korea_now()
    seed = f"{now_dt.isoformat()}|{time.time_ns()}"
    suffix = hashlib.sha256(seed.encode("utf-8")).hexdigest()[:6].upper()
    return f"WL-{now_dt.strftime('%Y%m%d-%H%M%S')}-{suffix}"


def _photo_capture_timestamp(uploaded_file, fallback_dt: datetime.datetime | None = None) -> str:
    """사진 EXIF 촬영시각을 우선 사용하고, 없으면 현재 한국시간을 파일명용 시각으로 반환합니다."""
    fallback_dt = fallback_dt or _korea_now()
    try:
        from io import BytesIO
        from PIL import Image

        raw = uploaded_file.getvalue() if uploaded_file is not None else b""
        if raw:
            image = Image.open(BytesIO(raw))
            exif = image.getexif()
            # DateTimeOriginal(36867) → DateTimeDigitized(36868) → DateTime(306) 순서
            for tag_id in (36867, 36868, 306):
                value = str(exif.get(tag_id, "") or "").strip()
                if not value:
                    continue
                for fmt in ("%Y:%m:%d %H:%M:%S", "%Y-%m-%d %H:%M:%S"):
                    try:
                        captured = datetime.datetime.strptime(value[:19], fmt)
                        return captured.strftime("%Y%m%d_%H%M%S")
                    except Exception:
                        continue
    except Exception:
        pass
    return fallback_dt.strftime("%Y%m%d_%H%M%S")


def _worklog_compress_image(uploaded_file) -> tuple[bytes | None, str, str, str]:
    """현장 사진을 방향보정하고 1600px/약 450KB 수준 JPEG로 최적화합니다."""
    if uploaded_file is None:
        return None, "", "", "사진이 없습니다."
    try:
        from io import BytesIO
        from PIL import Image, ImageOps

        raw = uploaded_file.getvalue()
        image = Image.open(BytesIO(raw))
        image = ImageOps.exif_transpose(image)
        if image.mode not in ("RGB", "L"):
            # 투명 PNG/WebP는 흰 배경으로 합성해 현장 문서 가독성을 유지합니다.
            if "A" in image.getbands():
                background = Image.new("RGB", image.size, "white")
                alpha = image.getchannel("A")
                background.paste(image.convert("RGB"), mask=alpha)
                image = background
            else:
                image = image.convert("RGB")
        elif image.mode == "L":
            image = image.convert("RGB")

        max_side = max(image.size)
        if max_side > WORK_LOG_IMAGE_MAX_SIDE:
            ratio = WORK_LOG_IMAGE_MAX_SIDE / float(max_side)
            image = image.resize(
                (max(1, int(image.width * ratio)), max(1, int(image.height * ratio))),
                Image.Resampling.LANCZOS,
            )

        best = None
        working = image
        for _resize_round in range(4):
            for quality in (84, 78, 72, 66, 60, 54, 48):
                buffer = BytesIO()
                working.save(
                    buffer,
                    format="JPEG",
                    quality=quality,
                    optimize=True,
                    progressive=True,
                )
                data = buffer.getvalue()
                best = data
                if len(data) <= WORK_LOG_IMAGE_TARGET_BYTES:
                    break
            if best is not None and len(best) <= WORK_LOG_IMAGE_TARGET_BYTES:
                break
            working = working.resize(
                (max(1, int(working.width * 0.86)), max(1, int(working.height * 0.86))),
                Image.Resampling.LANCZOS,
            )

        original_name = str(getattr(uploaded_file, "name", "field_photo") or "field_photo")
        safe_stem = re.sub(r"[^0-9A-Za-z가-힣_-]", "_", os.path.splitext(os.path.basename(original_name))[0])[:60] or "field_photo"
        return best, f"{safe_stem}.jpg", "image/jpeg", ""
    except ImportError:
        return None, "", "", "사진 자동 압축을 위해 Pillow 패키지가 필요합니다. requirements.txt에 Pillow를 추가해 주세요."
    except Exception as error:
        return None, "", "", f"사진 처리 실패: {error}"


def _worklog_upload_drive_image(image_bytes: bytes, file_name: str, mime_type: str) -> tuple[bool, dict, str]:
    """압축 사진 1장을 Apps Script로 저장합니다.

    JSON을 application/json으로 직접 보내지 않고 text/plain JSON으로 전송해 Apps Script 웹앱의
    요청 파싱/보안 프런트엔드와의 호환성을 높입니다. ContentService 리디렉션은 즉시 분리 처리합니다.
    """
    upload_url, url_error = _worklog_normalize_apps_script_url()
    upload_token = _worklog_photo_upload_token()
    if url_error:
        return False, {}, url_error
    if not upload_token:
        return False, {}, "Streamlit Secrets의 [work_log] upload_token이 설정되지 않았습니다."

    try:
        payload = {
            "token": upload_token,
            "filename": file_name,
            "mimeType": mime_type or "image/jpeg",
            "data": base64.b64encode(image_bytes).decode("ascii"),
        }
        payload_bytes = json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8")

        first = requests.post(
            upload_url,
            data=payload_bytes,
            headers={
                "Content-Type": "text/plain; charset=utf-8",
                "Accept": "application/json,text/plain,*/*",
                "User-Agent": "SMART-WORK-AI-AGENT/4.0",
                "Cache-Control": "no-cache",
            },
            timeout=60,
            allow_redirects=False,
        )
        response, redirect_error = _worklog_follow_apps_script_response(first, timeout=60)
        if redirect_error:
            return False, {}, redirect_error
        if response is None:
            return False, {}, "Apps Script 응답을 받지 못했습니다."
        if response.status_code != 200:
            return False, {}, (
                f"PHOTO ENGINE {WORK_LOG_PHOTO_ENGINE_VERSION} · Apps Script POST 실패 ({response.status_code}). "
                "이 오류는 Google Drive 폴더 문제가 아니라 /exec 배포 또는 익명 POST 접근 문제입니다."
            )
        if response.text.lstrip().lower().startswith("<!doctype html"):
            return False, {}, (
                f"PHOTO ENGINE {WORK_LOG_PHOTO_ENGINE_VERSION} · Apps Script가 JSON 대신 Google HTML을 반환했습니다. "
                "현재 /exec 배포의 익명 POST 접근이 허용되지 않았거나 운영 URL이 최신 배포가 아닙니다."
            )
        try:
            result = response.json()
        except Exception:
            return False, {}, f"Apps Script 응답이 JSON이 아닙니다: {response.text[:180]}"

        if not bool(result.get("ok")):
            error_text = str(result.get("error", "알 수 없는 오류") or "알 수 없는 오류")
            if error_text.lower() == "unauthorized":
                return False, {}, "Apps Script에는 도달했지만 UPLOAD_TOKEN이 일치하지 않습니다."
            return False, {}, f"Apps Script doPost 오류: {error_text}"

        file_id = str(result.get("fileId", "") or "").strip()
        if not file_id:
            return False, {}, "Apps Script 저장 응답에 fileId가 없습니다."
        return True, {
            "id": file_id,
            "name": str(result.get("fileName", file_name) or file_name),
            "webViewLink": str(result.get("fileUrl", "") or ""),
            "mimeType": mime_type or "image/jpeg",
        }, ""
    except requests.Timeout:
        return False, {}, "Apps Script 사진 저장 시간이 초과되었습니다."
    except Exception as error:
        return False, {}, f"Apps Script 사진 저장 오류: {error}"



def _worklog_apps_script_file_action(action: str, file_ids: list[str]) -> tuple[bool, dict, str]:
    """Apps Script 소유자 권한으로 Drive 파일을 휴지통 이동/복원합니다."""
    normalized_ids = [str(value or "").strip() for value in file_ids if str(value or "").strip()]
    if not normalized_ids:
        return True, {"processedIds": []}, ""

    upload_url, url_error = _worklog_normalize_apps_script_url()
    upload_token = _worklog_photo_upload_token()
    if url_error:
        return False, {}, url_error
    if not upload_token:
        return False, {}, "[work_log] upload_token이 설정되지 않았습니다."

    payload = {
        "token": upload_token,
        "action": action,
        "fileIds": normalized_ids,
    }
    try:
        first = requests.post(
            upload_url,
            data=json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8"),
            headers={
                "Content-Type": "text/plain; charset=utf-8",
                "Accept": "application/json,text/plain,*/*",
                "User-Agent": "SMART-WORK-AI-AGENT/5.0",
                "Cache-Control": "no-cache",
            },
            timeout=60,
            allow_redirects=False,
        )
        response, redirect_error = _worklog_follow_apps_script_response(first, timeout=60)
        if redirect_error:
            return False, {}, redirect_error
        if response is None or response.status_code != 200:
            code = getattr(response, "status_code", getattr(first, "status_code", ""))
            return False, {}, f"Apps Script 파일 처리 요청 실패 ({code})"
        if response.text.lstrip().lower().startswith("<!doctype html"):
            return False, {}, "Apps Script가 JSON 대신 Google HTML을 반환했습니다. 최신 /exec 배포를 확인해 주세요."
        try:
            result = response.json()
        except Exception:
            return False, {}, f"Apps Script 파일 처리 응답이 JSON이 아닙니다: {response.text[:180]}"

        if not bool(result.get("ok")):
            error_text = str(result.get("error", "") or "")
            if error_text.lower() == "unauthorized":
                return False, result, "UPLOAD_TOKEN이 일치하지 않습니다."
            if "사진 데이터가 없습니다" in error_text:
                return False, result, (
                    "현재 Apps Script 배포에는 사진 삭제 기능이 없습니다. "
                    "V12와 함께 제공된 Apps Script V2로 교체 후 배포를 업데이트해 주세요."
                )
            return False, result, error_text or "Drive 파일 처리에 실패했습니다."
        return True, result, ""
    except requests.Timeout:
        return False, {}, "Drive 파일 처리 시간이 초과되었습니다."
    except Exception as error:
        return False, {}, f"Drive 파일 처리 오류: {error}"


def _worklog_trash_drive_files(file_ids: list[str]) -> tuple[bool, list[str], str]:
    """사진 파일을 영구삭제하지 않고 사용자 Drive 휴지통으로 이동합니다."""
    ok, result, error = _worklog_apps_script_file_action("trash_files", file_ids)
    processed = [
        str(value or "").strip()
        for value in (result.get("processedIds", result.get("trashedIds", [])) or [])
        if str(value or "").strip()
    ]
    if ok:
        return True, processed, ""

    # 부분 처리된 경우 로그 삭제를 중단하고 이미 휴지통으로 간 파일은 되돌립니다.
    if processed:
        _worklog_apps_script_file_action("restore_files", processed)
    return False, [], error


def _worklog_restore_drive_files(file_ids: list[str]) -> None:
    normalized_ids = [str(value or "").strip() for value in file_ids if str(value or "").strip()]
    if normalized_ids:
        _worklog_apps_script_file_action("restore_files", normalized_ids)


@st.cache_data(ttl=300, show_spinner=False)
def _worklog_download_drive_image(file_id: str) -> bytes | None:
    """비공개 Drive 사진을 서비스 계정으로 읽어 최근 기록 썸네일에 사용합니다."""
    file_id = str(file_id or "").strip()
    if not file_id:
        return None
    token = _worklog_drive_access_token()
    if not token:
        return None
    try:
        response = requests.get(
            f"https://www.googleapis.com/drive/v3/files/{file_id}",
            params={"alt": "media", "supportsAllDrives": "true"},
            headers={"Authorization": f"Bearer {token}"},
            timeout=25,
        )
        if response.status_code == 200 and len(response.content) <= 8 * 1024 * 1024:
            return response.content
    except Exception:
        pass
    return None


def _render_power_photo_download(record, key_prefix: str) -> None:
    """정밀점검 1건에 연결된 비공개 Drive 사진을 전체보기/개별/ZIP 다운로드로 제공합니다."""
    photo_ids = [
        value.strip() for value in str(record.get("사진파일ID목록", "") or "").split("|") if value.strip()
    ]
    photo_names = [
        value.strip() for value in str(record.get("사진파일명목록", "") or "").split("|") if value.strip()
    ]
    if not photo_ids:
        st.info("이 정밀점검 기록에는 첨부된 사진이 없습니다.")
        return

    payloads = []
    failed_count = 0
    saved_stamp = re.sub(r"[^0-9]", "", str(record.get("저장일시", "") or ""))[:14] or _korea_now().strftime("%Y%m%d%H%M%S")
    safe_local = re.sub(r"[^0-9A-Za-z가-힣_-]", "_", str(record.get("국소", "") or "정밀점검"))[:40] or "정밀점검"

    for index, file_id in enumerate(photo_ids, 1):
        photo_bytes = _worklog_download_drive_image(file_id)
        if not photo_bytes:
            failed_count += 1
            continue
        stored_name = photo_names[index - 1] if index - 1 < len(photo_names) else f"정밀점검_{saved_stamp}_{index:02d}.jpg"
        download_name = stored_name if stored_name.lower().endswith(".jpg") else f"{stored_name}.jpg"
        payloads.append({"index": index, "bytes": photo_bytes, "name": download_name})

    if not payloads:
        st.warning("사진 정보는 있으나 현재 파일을 읽을 수 없습니다. Drive 읽기 권한을 확인해 주세요.")
        return

    st.caption(f"첨부사진 {len(photo_ids)}장 · Drive 폴더는 공개하지 않고 이 화면에서만 조회·다운로드합니다.")
    photo_cols = st.columns(2)
    for pos, payload in enumerate(payloads):
        with photo_cols[pos % 2]:
            st.image(payload["bytes"], caption=f"사진 {payload['index']} / {len(photo_ids)}", use_container_width=True)
            st.download_button(
                "📥 사진 다운로드",
                data=payload["bytes"],
                file_name=payload["name"],
                mime="image/jpeg",
                use_container_width=True,
                key=f"{key_prefix}_photo_{payload['index']}",
            )

    try:
        from io import BytesIO
        import zipfile

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as photo_zip:
            for payload in payloads:
                photo_zip.writestr(payload["name"], payload["bytes"])
        zip_name = f"{safe_local}_정밀점검_{saved_stamp}_사진{len(payloads)}장.zip"
        st.download_button(
            f"📦 사진 {len(payloads)}장 전체 ZIP 다운로드",
            data=zip_buffer.getvalue(),
            file_name=zip_name,
            mime="application/zip",
            use_container_width=True,
            key=f"{key_prefix}_zip",
        )
    except Exception as zip_error:
        st.warning(f"사진 ZIP 파일을 만들지 못했습니다. 개별 다운로드를 이용해 주세요. ({zip_error})")

    if failed_count:
        st.warning(f"첨부사진 중 {failed_count}장은 현재 읽을 수 없어 표시하지 못했습니다.")


def _worklog_collect_photos(camera_photo, uploaded_photos) -> list:
    """카메라 촬영 + 앨범 업로드를 중복 제거해 최대 10장으로 합칩니다."""
    candidates = []
    if camera_photo is not None:
        candidates.append(camera_photo)
    candidates.extend(list(uploaded_photos or []))

    unique = []
    seen = set()
    for file_obj in candidates:
        try:
            digest = hashlib.sha256(file_obj.getvalue()).hexdigest()
        except Exception:
            digest = f"{getattr(file_obj, 'name', '')}|{id(file_obj)}"
        if digest in seen:
            continue
        seen.add(digest)
        unique.append(file_obj)
        if len(unique) >= WORK_LOG_MAX_PHOTOS:
            break
    return unique



def _worklog_get_record_row(worksheet, record_id: str) -> tuple[dict, int | None, list[str]]:
    """기록ID로 시트의 실제 행/레코드를 찾습니다."""
    record_id = str(record_id or "").strip()
    try:
        values = worksheet.get_all_values()
    except Exception:
        return {}, None, []
    if not values:
        return {}, None, []
    headers = [str(value).strip() for value in values[0]]
    if "기록ID" not in headers:
        return {}, None, headers
    id_index = headers.index("기록ID")
    for row_no, row in enumerate(values[1:], start=2):
        current_id = str(row[id_index] if id_index < len(row) else "").strip()
        if current_id == record_id:
            record = {
                header: (row[index] if index < len(row) else "")
                for index, header in enumerate(headers)
            }
            return record, row_no, headers
    return {}, None, headers


def _worklog_can_access_record(record: dict, auth_user: dict) -> bool:
    if not auth_user:
        return False
    if _worklog_record_visibility(record) == "공개":
        return True
    return _worklog_record_owned_by(record, auth_user)


def save_work_log(record: dict, photos: list) -> tuple[bool, str, str]:
    """로그인 사용자 소유의 현장기록을 저장하고 사진은 비공개 Drive에 분리 저장합니다."""
    client = init_google_sheet_connection()
    if not client:
        return False, "Google Sheets 연결 실패: 기존 gcp_service_account 설정을 확인하세요.", ""

    writer = str(record.get("작성자", "")).strip()
    owner_id = str(record.get("작성자ID", "") or "").strip()
    visibility = str(record.get("공개범위", "공개") or "공개").strip()
    area = str(record.get("권역", "")).strip()
    mother = str(record.get("모국", "")).strip()
    local = str(record.get("국소", "")).strip()
    status = str(record.get("상태", "신규")).strip()
    items = record.get("점검항목", []) or []
    issue = str(record.get("현상_특이사항", "")).strip()
    action = str(record.get("조치내용", "")).strip()
    followup = str(record.get("후속조치", "")).strip()
    remark = str(record.get("비고", "")).strip()

    expected_owner = _worklog_user_id_from_name(writer)
    if not writer or not owner_id:
        return False, "MY WORK LOG 개인 인증 정보를 확인하지 못했습니다. 다시 로그인해 주세요.", ""
    if expected_owner and expected_owner != owner_id:
        return False, "로그인 사용자와 작성자 정보가 일치하지 않습니다. 다시 로그인해 주세요.", ""
    if visibility not in WORK_LOG_VISIBILITY_OPTIONS:
        return False, "공개범위 값이 올바르지 않습니다.", ""
    if area not in POWER_REGION_DATA:
        return False, "국사명을 검색하여 정확한 국사를 먼저 선택해 주세요.", ""
    area_map = POWER_REGION_DATA.get(area, {}).get("모국_국소", {})
    if mother not in area_map or local not in area_map.get(mother, []):
        return False, "선택한 국사의 권역·모국·국소 정보가 올바르지 않습니다. 국사를 다시 검색해 주세요.", ""
    if status not in WORK_LOG_STATUS_OPTIONS:
        return False, "상태이력 값이 올바르지 않습니다.", ""
    if not items:
        return False, "점검항목을 한 개 이상 선택해 주세요.", ""
    if not any([issue, action, followup, remark, photos]):
        return False, "현상·특이사항, 조치내용, 후속조치, 비고 또는 사진 중 한 가지 이상을 남겨 주세요.", ""
    if photos and not _worklog_photo_upload_ready():
        _, config_issues = _worklog_photo_config_status()
        issue_text = " / ".join(config_issues) if config_issues else "사진 업로드 설정 확인 필요"
        return False, (
            "사진 업로드 설정 오류: " + issue_text + ". "
            "Streamlit Secrets의 [work_log]에는 photo_upload_url과 upload_token이 필요합니다. "
            "photo_upload_url은 Apps Script 배포용 /exec 주소, upload_token은 Apps Script의 UPLOAD_TOKEN과 동일해야 합니다."
        ), ""

    if photos:
        preflight_ok, preflight_message = _worklog_apps_script_healthcheck()
        if not preflight_ok:
            return False, f"사진 업로드 사전진단 실패: {preflight_message}", ""

    now = _korea_now()
    record_id = _worklog_make_id(now)
    photo_ids = []
    photo_names = []
    used_worklog_photo_names: set[str] = set()

    for index, photo in enumerate(photos[:WORK_LOG_MAX_PHOTOS], 1):
        compressed, safe_name, mime_type, error = _worklog_compress_image(photo)
        if not compressed:
            return False, f"{index}번째 사진 처리 실패: {error}", ""
        capture_stamp = _photo_capture_timestamp(photo, fallback_dt=now)
        base_drive_name = f"WORK LOG_{capture_stamp}"
        drive_name = f"{base_drive_name}.jpg"
        duplicate_no = 2
        while drive_name in used_worklog_photo_names:
            drive_name = f"{base_drive_name}_{duplicate_no:02d}.jpg"
            duplicate_no += 1
        used_worklog_photo_names.add(drive_name)
        ok, drive_meta, upload_error = _worklog_upload_drive_image(compressed, drive_name, mime_type)
        if not ok:
            return False, f"{index}번째 사진 저장 실패: {upload_error}", ""
        photo_ids.append(str(drive_meta.get("id", "")))
        photo_names.append(str(drive_meta.get("name", drive_name)))

    try:
        spreadsheet = client.open(WORK_LOG_SPREADSHEET_NAME)
        ws, history_ws = _worklog_ensure_sheets(spreadsheet)
        saved_at = now.strftime("%Y-%m-%d %H:%M:%S")
        row_map = {
            "저장일시": saved_at,
            "기록ID": record_id,
            "작성자": writer,
            "권역": area,
            "모국": mother,
            "국소": local,
            "상태": status,
            "점검항목": ", ".join(str(v) for v in items),
            "현상_특이사항": issue,
            "조치내용": action,
            "후속조치": followup,
            "비고": remark,
            "사진수": len(photo_ids),
            "사진파일ID목록": "|".join(photo_ids),
            "사진파일명목록": "|".join(photo_names),
            "최근수정일시": saved_at,
            "작성자ID": owner_id,
            "공개범위": visibility,
        }
        actual_headers = [str(value).strip() for value in ws.row_values(1)]
        ws.append_row([row_map.get(header, "") for header in actual_headers], value_input_option="USER_ENTERED")

        history_headers = [str(value).strip() for value in history_ws.row_values(1)]
        history_map = {
            "저장일시": saved_at,
            "기록ID": record_id,
            "작성자": writer,
            "상태": status,
            "변경구분": "신규 등록",
            "조치내용": action,
            "후속조치": followup,
            "비고": remark,
            "작업자ID": owner_id,
        }
        history_ws.append_row(
            [history_map.get(header, "") for header in history_headers],
            value_input_option="USER_ENTERED",
        )
        scope_text = "팀 공유" if visibility == "공개" else "나만 보기"
        return True, f"MY WORK LOG가 저장되었습니다. · {scope_text} · 사진 {len(photo_ids)}장", record_id
    except Exception as error:
        return False, f"WORK LOG 저장 실패: {error}", ""


def load_work_logs(auth_user: dict | None = None) -> pd.DataFrame:
    """누적 WORK LOG를 읽은 뒤 공개 기록 + 로그인 사용자의 비공개 기록만 반환합니다."""
    auth_user = auth_user or _worklog_current_user()
    if not auth_user:
        return pd.DataFrame(columns=WORK_LOG_HEADERS)

    client = init_google_sheet_connection()
    if not client:
        raise RuntimeError("Google Sheets 연결 실패")
    spreadsheet = client.open(WORK_LOG_SPREADSHEET_NAME)
    try:
        ws = spreadsheet.worksheet(WORK_LOG_SHEET_NAME)
    except Exception:
        return pd.DataFrame(columns=WORK_LOG_HEADERS)

    values = ws.get_all_values()
    if not values:
        return pd.DataFrame(columns=WORK_LOG_HEADERS)

    headers = [str(v).strip() for v in values[0]]
    rows = [
        [row[i] if i < len(row) else "" for i in range(len(headers))]
        for row in values[1:]
        if any(str(cell or "").strip() for cell in row)
    ]
    df = pd.DataFrame(rows, columns=headers).fillna("")
    for required in WORK_LOG_HEADERS:
        if required not in df.columns:
            df[required] = ""

    df = _worklog_filter_accessible_records(df, auth_user)
    if not df.empty:
        df["_저장일시_dt"] = pd.to_datetime(df["저장일시"], errors="coerce")
        df = df.sort_values("_저장일시_dt", ascending=False, na_position="last")
    return df


def load_work_log_history(record_id: str, auth_user: dict | None = None) -> pd.DataFrame:
    """로그인 사용자가 볼 수 있는 기록에 대해서만 변경이력을 반환합니다."""
    auth_user = auth_user or _worklog_current_user()
    if not auth_user:
        return pd.DataFrame(columns=WORK_LOG_HISTORY_HEADERS)

    client = init_google_sheet_connection()
    if not client:
        return pd.DataFrame(columns=WORK_LOG_HISTORY_HEADERS)
    try:
        spreadsheet = client.open(WORK_LOG_SPREADSHEET_NAME)
        main_ws, history_ws = _worklog_ensure_sheets(spreadsheet)
        record, _, _ = _worklog_get_record_row(main_ws, record_id)
        if not record or not _worklog_can_access_record(record, auth_user):
            return pd.DataFrame(columns=WORK_LOG_HISTORY_HEADERS)

        records = history_ws.get_all_records()
        df = pd.DataFrame(records).fillna("") if records else pd.DataFrame(columns=WORK_LOG_HISTORY_HEADERS)
        if not df.empty and "기록ID" in df.columns:
            df = df[df["기록ID"].astype(str) == str(record_id)]
        return df
    except Exception:
        return pd.DataFrame(columns=WORK_LOG_HISTORY_HEADERS)


def update_work_log(
    record_id: str,
    writer: str,
    status: str,
    action: str,
    followup: str,
    remark: str,
    actor_user: dict | None = None,
) -> tuple[bool, str]:
    """공개 기록 또는 본인 비공개 기록만 갱신하고 실제 변경 사용자를 이력에 남깁니다."""
    record_id = str(record_id or "").strip()
    status = str(status or "").strip()
    actor_user = actor_user or _worklog_current_user()
    if not record_id:
        return False, "변경할 기록ID가 없습니다."
    if not actor_user:
        return False, "MY WORK LOG 개인 인증이 필요합니다."
    if status not in WORK_LOG_STATUS_OPTIONS:
        return False, "상태 값이 올바르지 않습니다."

    actor_name = str(actor_user.get("name", "") or "").strip()
    actor_id = str(actor_user.get("user_id", "") or "").strip()
    if not actor_name or not actor_id:
        return False, "로그인 사용자 정보를 확인하지 못했습니다."

    client = init_google_sheet_connection()
    if not client:
        return False, "Google Sheets 연결 실패"
    try:
        spreadsheet = client.open(WORK_LOG_SPREADSHEET_NAME)
        ws, history_ws = _worklog_ensure_sheets(spreadsheet)
        record, target_row, headers = _worklog_get_record_row(ws, record_id)
        if not record or target_row is None:
            return False, f"기록ID {record_id}를 찾지 못했습니다."
        if not _worklog_can_access_record(record, actor_user):
            return False, "이 기록에 접근할 권한이 없습니다."

        now_text = _korea_now().strftime("%Y-%m-%d %H:%M:%S")
        updates = {
            "상태": status,
            "조치내용": str(action or "").strip(),
            "후속조치": str(followup or "").strip(),
            "비고": str(remark or "").strip(),
            "최근수정일시": now_text,
        }
        for header, value in updates.items():
            if header in headers:
                ws.update_cell(target_row, headers.index(header) + 1, value)

        history_headers = [str(value).strip() for value in history_ws.row_values(1)]
        history_map = {
            "저장일시": now_text,
            "기록ID": record_id,
            "작성자": actor_name,
            "상태": status,
            "변경구분": "상태/조치 변경",
            "조치내용": updates["조치내용"],
            "후속조치": updates["후속조치"],
            "비고": updates["비고"],
            "작업자ID": actor_id,
        }
        history_ws.append_row(
            [history_map.get(header, "") for header in history_headers],
            value_input_option="USER_ENTERED",
        )
        return True, "상태이력과 조치내용이 업데이트되었습니다."
    except Exception as error:
        return False, f"WORK LOG 업데이트 실패: {error}"


def update_work_log_visibility(
    record_id: str,
    new_visibility: str,
    actor_user: dict | None = None,
) -> tuple[bool, str]:
    """작성자 본인만 자신의 과거/현재 WORK LOG 공개범위를 변경합니다.

    기존 자료에 작성자ID/공개범위가 비어 있어도 등록된 작성자 이름으로 소유권을
    복원한 뒤 명시값을 저장합니다. 공개범위 변경 사실은 MY_WORK_LOG_HISTORY에 남깁니다.
    """
    record_id = str(record_id or "").strip()
    new_visibility = str(new_visibility or "").strip()
    actor_user = actor_user or _worklog_current_user()

    if not record_id:
        return False, "변경할 기록ID가 없습니다."
    if not actor_user:
        return False, "MY WORK LOG 개인 인증이 필요합니다."
    if new_visibility not in WORK_LOG_VISIBILITY_OPTIONS:
        return False, "공개범위 값이 올바르지 않습니다."

    actor_name = str(actor_user.get("name", "") or "").strip()
    actor_id = str(actor_user.get("user_id", "") or "").strip()
    if not actor_name or not actor_id:
        return False, "로그인 사용자 정보를 확인하지 못했습니다."

    client = init_google_sheet_connection()
    if not client:
        return False, "Google Sheets 연결 실패"

    try:
        spreadsheet = client.open(WORK_LOG_SPREADSHEET_NAME)
        ws, history_ws = _worklog_ensure_sheets(spreadsheet)
        record, target_row, headers = _worklog_get_record_row(ws, record_id)
        if not record or target_row is None:
            return False, f"기록ID {record_id}를 찾지 못했습니다."
        if not _worklog_record_owned_by(record, actor_user):
            return False, "공개범위는 본인이 작성한 기록만 변경할 수 있습니다."

        old_visibility = _worklog_record_visibility(record)
        owner_id = _worklog_record_owner_id(record) or actor_id
        now_text = _korea_now().strftime("%Y-%m-%d %H:%M:%S")

        # 과거 자료의 빈 권한 필드도 이 시점에 명시값으로 보완합니다.
        updates = {
            "작성자ID": owner_id,
            "공개범위": new_visibility,
            "최근수정일시": now_text,
        }
        for header, value in updates.items():
            if header in headers:
                ws.update_cell(target_row, headers.index(header) + 1, value)

        if old_visibility != new_visibility:
            history_headers = [str(value).strip() for value in history_ws.row_values(1)]
            history_map = {
                "저장일시": now_text,
                "기록ID": record_id,
                "작성자": actor_name,
                "상태": str(record.get("상태", "") or "").strip(),
                "변경구분": f"공개범위 변경: {old_visibility} → {new_visibility}",
                "조치내용": "",
                "후속조치": "",
                "비고": "",
                "작업자ID": actor_id,
            }
            history_ws.append_row(
                [history_map.get(header, "") for header in history_headers],
                value_input_option="USER_ENTERED",
            )

        if old_visibility == new_visibility:
            return True, f"현재 공개범위가 이미 '{new_visibility}'입니다."
        scope_text = "🌐 공개 · 팀 공유" if new_visibility == "공개" else "🔒 비공개 · 나만 보기"
        return True, f"공개범위를 {scope_text}(으)로 변경했습니다. 사진 열람 권한도 같은 범위를 따릅니다."
    except Exception as error:
        return False, f"공개범위 변경 실패: {error}"


def _worklog_ensure_delete_audit_sheet(spreadsheet):
    try:
        ws = spreadsheet.worksheet(WORK_LOG_DELETE_AUDIT_SHEET_NAME)
    except Exception:
        ws = spreadsheet.add_worksheet(
            title=WORK_LOG_DELETE_AUDIT_SHEET_NAME,
            rows=5000,
            cols=max(len(WORK_LOG_DELETE_AUDIT_HEADERS) + 2, 12),
        )
        ws.append_row(WORK_LOG_DELETE_AUDIT_HEADERS, value_input_option="USER_ENTERED")
    _worklog_ensure_headers(ws, WORK_LOG_DELETE_AUDIT_HEADERS)
    return ws


def delete_work_log(record_id: str, auth_user: dict | None = None) -> tuple[bool, str]:
    """작성자 본인의 WORK LOG 1건과 연결 사진을 함께 삭제합니다. 사진은 Drive 휴지통으로 이동합니다."""
    auth_user = auth_user or _worklog_current_user()
    record_id = str(record_id or "").strip()
    if not auth_user:
        return False, "MY WORK LOG 개인 인증이 필요합니다."
    if not record_id:
        return False, "삭제할 기록ID가 없습니다."

    client = init_google_sheet_connection()
    if not client:
        return False, "Google Sheets 연결 실패"

    trashed_photo_ids: list[str] = []
    try:
        spreadsheet = client.open(WORK_LOG_SPREADSHEET_NAME)
        ws, history_ws = _worklog_ensure_sheets(spreadsheet)
        record, target_row, headers = _worklog_get_record_row(ws, record_id)
        if not record or target_row is None:
            return False, "삭제할 기록을 찾지 못했습니다."
        if not _worklog_record_owned_by(record, auth_user):
            return False, "본인이 작성한 기록만 삭제할 수 있습니다."

        photo_ids = [
            value.strip()
            for value in str(record.get("사진파일ID목록", "") or "").split("|")
            if value.strip()
        ]
        if photo_ids:
            photo_ok, trashed_photo_ids, photo_error = _worklog_trash_drive_files(photo_ids)
            if not photo_ok:
                return False, f"기록은 삭제하지 않았습니다. 연결 사진 정리에 실패했습니다: {photo_error}"

        # 본문을 지우기 전에 삭제 감사 메타데이터를 준비합니다. 개인 메모 내용 자체는 감사시트에 복제하지 않습니다.
        deleted_at = _korea_now().strftime("%Y-%m-%d %H:%M:%S")
        audit_map = {
            "삭제일시": deleted_at,
            "기록ID": record_id,
            "작성자ID": _worklog_record_owner_id(record),
            "작성자": str(record.get("작성자", "") or "").strip(),
            "삭제자ID": str(auth_user.get("user_id", "") or "").strip(),
            "삭제자": str(auth_user.get("name", "") or "").strip(),
            "공개범위": _worklog_record_visibility(record),
            "사진수": len(photo_ids),
        }

        # 이력은 아래 행부터 지워 행번호 변화를 방지합니다.
        history_values = history_ws.get_all_values()
        if history_values:
            history_headers = [str(value).strip() for value in history_values[0]]
            if "기록ID" in history_headers:
                id_index = history_headers.index("기록ID")
                delete_rows = [
                    row_no
                    for row_no, row in enumerate(history_values[1:], start=2)
                    if id_index < len(row) and str(row[id_index]).strip() == record_id
                ]
                for row_no in reversed(delete_rows):
                    history_ws.delete_rows(row_no)

        ws.delete_rows(target_row)

        audit_warning = ""
        try:
            audit_ws = _worklog_ensure_delete_audit_sheet(spreadsheet)
            audit_headers = [str(value).strip() for value in audit_ws.row_values(1)]
            audit_ws.append_row(
                [audit_map.get(header, "") for header in audit_headers],
                value_input_option="USER_ENTERED",
            )
        except Exception as audit_error:
            audit_warning = f" · 삭제 감사기록 저장 경고: {audit_error}"

        return True, f"WORK LOG 1건을 삭제했습니다. 연결 사진 {len(photo_ids)}장은 Drive 휴지통으로 이동했습니다.{audit_warning}"
    except Exception as error:
        # 본문 삭제가 실패했는데 사진만 휴지통으로 간 경우 자동 복원합니다.
        if trashed_photo_ids:
            _worklog_restore_drive_files(trashed_photo_ids)
        return False, f"WORK LOG 삭제 실패: {error}"


def _worklog_reset_entry_widgets() -> None:
    keys = [
        "worklog_writer", "worklog_area", "worklog_area_key", "worklog_mother", "worklog_local", "worklog_status",
        "worklog_items", "worklog_camera", "worklog_uploads", "worklog_issue", "worklog_action",
        "worklog_followup", "worklog_remark", "worklog_station_search_query",
        "worklog_station_search_candidates", "worklog_station_search_choice", "worklog_station_search_status",
        "worklog_station_search_notice", "worklog_station_search_applied",
        "worklog_items_confirmed_notice", "worklog_visibility",
    ]
    for key in keys:
        if key in st.session_state:
            del st.session_state[key]


# ==========================================
# 9. 메인 화면 및 탭 구성
# ==========================================
st.markdown("""
<div class="smart-work-brand">
  <div class="smart-work-brand-line"></div>
  <div class="smart-work-brand-title">
    <span class="smart-work-brand-icon">◆</span>
    <span>SMART WORK <b>AI AGENT</b></span>
  </div>
  <div class="smart-work-brand-subtitle">Integrated Field &amp; Business Assistant System</div>
</div>
""", unsafe_allow_html=True)

_now_kst = _korea_now()
CURRENT_YEAR = _now_kst.year
CURRENT_MONTH = _now_kst.month

campaign_info = {
    "key": f"{CURRENT_YEAR}-{CURRENT_MONTH:02d}",
    "title": _default_campaign_title(_now_kst),
    "sheet_name": f"{CURRENT_YEAR}_{CURRENT_MONTH:02d}_자율점검",
    "start_date": _now_kst.strftime("%Y.%m.%d"),
}

# ✅ 전 임직원 교육 집중 기간에는 앱 시작 시 Google Sheet를 읽지 않습니다.
# 기존 campaign_info 기본값만 사용하여 교육 화면 로딩 속도와 제출 안정성을 우선합니다.

# ✅ 상단 메뉴 카드형 디자인: 선택된 탭이 명확하게 보이도록 개선
st.markdown("""
<style>
/* 상단 브랜드: 기존 제목 문구는 유지하고 가독성/색상만 보강 */
.smart-work-brand {
    text-align: center;
    margin: 2px auto 18px auto;
    padding: 8px 12px 11px 12px;
}
.smart-work-brand-line {
    width: 78px;
    height: 4px;
    margin: 0 auto 9px auto;
    border-radius: 999px;
    background: linear-gradient(90deg, #D71920 0%, #FF5A5F 55%, #0F4C81 55%, #0F4C81 100%);
}
.smart-work-brand-title {
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 9px;
    color: #24364B;
    font-size: clamp(1.85rem, 4vw, 2.65rem);
    line-height: 1.08;
    font-weight: 950;
    letter-spacing: -0.035em;
}
.smart-work-brand-title b {
    color: #D71920;
    font-weight: 950;
}
.smart-work-brand-icon {
    color: #D71920;
    font-size: .72em;
    filter: drop-shadow(0 2px 3px rgba(215,25,32,.18));
}
.smart-work-brand-subtitle {
    margin-top: 6px;
    color: #64748B;
    font-size: .92rem;
    font-weight: 750;
}

/* 외부 스마트 내비 메뉴 1개: 기존 Streamlit 탭은 그대로 유지 */
.smart-navi-launch-wrap {
    display: flex;
    justify-content: flex-start;
    margin: 0 0 9px 0;
}
.smart-navi-launch {
    display: inline-flex;
    align-items: center;
    gap: 10px;
    min-height: 52px;
    padding: 8px 14px 8px 11px;
    border: 1.5px solid #D71920;
    border-left: 5px solid #D71920;
    border-radius: 15px;
    background: linear-gradient(135deg, #FFFFFF 0%, #FFF7F7 100%);
    box-shadow: 0 7px 18px rgba(15,23,42,.08);
    color: #24364B !important;
    text-decoration: none !important;
    transition: transform .16s ease, box-shadow .16s ease, border-color .16s ease;
}
.smart-navi-launch:hover {
    transform: translateY(-1px);
    box-shadow: 0 10px 23px rgba(215,25,32,.14);
    border-color: #B91218;
}
.smart-navi-launch-icon {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 36px;
    height: 36px;
    border-radius: 11px;
    background: #D71920;
    color: #FFFFFF;
    font-size: 20px;
    line-height: 1;
    box-shadow: 0 5px 12px rgba(215,25,32,.22);
}
.smart-navi-launch-copy {
    display: flex;
    flex-direction: column;
    line-height: 1.15;
}
.smart-navi-launch-title {
    color: #24364B;
    font-size: 1.02rem;
    font-weight: 950;
    white-space: nowrap;
}
.smart-navi-launch-sub {
    margin-top: 3px;
    color: #D71920;
    font-size: .72rem;
    font-weight: 850;
}
.smart-navi-launch-arrow {
    color: #D71920;
    font-size: 1.05rem;
    font-weight: 950;
    margin-left: 2px;
}

/* Streamlit 탭을 카드형 메뉴처럼 보이게 개선 */
div[data-testid="stTabs"] > div[role="tablist"] {
    gap: 10px !important;
    overflow-x: auto !important;
    flex-wrap: nowrap !important;
    scrollbar-width: thin !important;
    background: linear-gradient(135deg, #EEF4FF 0%, #F8FAFC 100%) !important;
    border: 1px solid #D8E3F2 !important;
    border-radius: 20px !important;
    padding: 10px !important;
    box-shadow: 0 8px 22px rgba(15, 23, 42, 0.08) !important;
}
div[data-testid="stTabs"] button[role="tab"] {
    min-height: 62px !important;
    padding: 12px 18px !important;
    border-radius: 16px !important;
    border: 1px solid #D8E3F2 !important;
    background: #FFFFFF !important;
    color: #334155 !important;
    font-weight: 900 !important;
    box-shadow: 0 5px 14px rgba(15, 23, 42, 0.06) !important;
    transition: all 0.18s ease-in-out !important;
}
div[data-testid="stTabs"] button[role="tab"] p {
    font-size: clamp(1.16rem, 2.2vw, 1.30rem) !important;
    font-weight: 950 !important;
    letter-spacing: -0.02em !important;
    line-height: 1.2 !important;
    margin: 0 !important;
}
div[data-testid="stTabs"] button[role="tab"]:hover {
    transform: translateY(-1px) !important;
    border-color: #60A5FA !important;
    box-shadow: 0 9px 20px rgba(37, 99, 235, 0.13) !important;
}
div[data-testid="stTabs"] button[role="tab"][aria-selected="true"] {
    background: linear-gradient(135deg, #1D4ED8 0%, #0EA5E9 100%) !important;
    color: #FFFFFF !important;
    border-color: #38BDF8 !important;
    box-shadow: 0 12px 28px rgba(37, 99, 235, 0.28) !important;
    transform: translateY(-2px) !important;
}
div[data-testid="stTabs"] button[role="tab"][aria-selected="true"] p,
div[data-testid="stTabs"] button[role="tab"][aria-selected="true"] * {
    color: #FFFFFF !important;
    -webkit-text-fill-color: #FFFFFF !important;
}
div[data-testid="stTabs"] button[role="tab"][aria-selected="false"] p,
div[data-testid="stTabs"] button[role="tab"][aria-selected="false"] * {
    color: #334155 !important;
    -webkit-text-fill-color: #334155 !important;
}
/* 선택된 탭 하단 기본 라인 숨김 */
div[data-testid="stTabs"] button[role="tab"]::after {
    display: none !important;
}
/* 상단 메뉴 전환 시 부드럽게 내려오는 느낌을 주되 기능 실행에는 관여하지 않습니다. */
div[data-testid="stTabs"] div[role="tabpanel"] {
    animation: smartWorkTabReveal .18s ease-out;
}
@keyframes smartWorkTabReveal {
    from { opacity:.72; transform:translateY(-5px); }
    to { opacity:1; transform:translateY(0); }
}

@media (max-width: 768px) {
    section.main .block-container { padding-left:.65rem !important; padding-right:.65rem !important; padding-top:.75rem !important; }
    div[data-testid="stTabs"] > div[role="tablist"] { padding:7px !important; gap:7px !important; border-radius:15px !important; }
    div[data-testid="stTabs"] button[role="tab"] { flex:0 0 auto !important; min-width:154px !important; min-height:56px !important; padding:9px 13px !important; }
    div[data-testid="stTabs"] button[role="tab"] p { font-size:1.08rem !important; }
    .smart-navi-launch-wrap { justify-content: stretch; }
    .smart-navi-launch { width: 100%; box-sizing: border-box; }
    .smart-work-brand { margin-bottom: 14px; }
}

</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="smart-navi-launch-wrap">
  <a class="smart-navi-launch" href="https://willowy-frangipane-e06d37.netlify.app/" target="_blank" rel="noopener noreferrer" aria-label="국사 스마트 내비게이션 새 창으로 열기">
    <span class="smart-navi-launch-icon" aria-hidden="true">📡</span>
    <span class="smart-navi-launch-copy">
      <span class="smart-navi-launch-title">국사 스마트 내비</span>
      <span class="smart-navi-launch-sub">SMART NAVIGATION · 새 창</span>
    </span>
    <span class="smart-navi-launch-arrow" aria-hidden="true">↗</span>
  </a>
</div>
""", unsafe_allow_html=True)

tab_worklog, tab_power, tab_law, tab_admin = st.tabs([
    "📝 MY WORK LOG",
    "🔋 국사 전원시설 정밀점검",
    "⚖️ LAW SEARCH",
    "🔒 관리자 모드",
])

# LAW SEARCH는 기존 법률 검토/AI 에이전트/스마트 요약 기능을 그대로 묶은 하위 메뉴입니다.
# 기존 기능 코드는 아래에서 각 컨테이너에 그대로 렌더링됩니다.
with tab_law:
    st.markdown(
        '<div class="law-search-hero"><b>⚖️ LAW SEARCH</b>'
        '<span>기존 감사·법률 지원 기능을 한 곳에서 선택해 사용합니다.</span></div>',
        unsafe_allow_html=True,
    )
    tab_doc, tab_chat, tab_summary = st.tabs([
        "📄 법률 검토", "💬 AI 에이전트", "📰 스마트 요약"
    ])

# ---------- (아이콘) 인라인 SVG: 애니메이션 모래시계 ----------
HOURGLASS_SVG = """
<svg width="18" height="18" viewBox="0 0 24 24" fill="none"
     xmlns="http://www.w3.org/2000/svg" aria-hidden="true">
  <path d="M6 2h12v5c0 2.2-1.4 4.2-3.5 5 2.1.8 3.5 2.8 3.5 5v5H6v-5c0-2.2 1.4-4.2 3.5-5C7.4 11.2 6 9.2 6 7V2Z"
        stroke="#0B5ED7" stroke-width="2" stroke-linejoin="round"/>
  <path d="M8 7h8M8 17h8" stroke="#0B5ED7" stroke-width="2" stroke-linecap="round"/>

  <rect x="9" y="8.2" width="6" height="3.0" rx="1.0" fill="#0B5ED7" opacity="0.95">
    <animate attributeName="height" values="3.0;0.3;3.0" dur="1.0s" repeatCount="indefinite" />
    <animate attributeName="y"      values="8.2;10.9;8.2" dur="1.0s" repeatCount="indefinite" />
  </rect>

  <rect x="9" y="15.8" width="6" height="0.3" rx="1.0" fill="#0B5ED7" opacity="0.95">
    <animate attributeName="height" values="0.3;3.0;0.3" dur="1.0s" repeatCount="indefinite" />
    <animate attributeName="y"      values="15.8;13.1;15.8" dur="1.0s" repeatCount="indefinite" />
  </rect>

  <circle cx="12" cy="12" r="0.8" fill="#0B5ED7" opacity="0.95">
    <animate attributeName="cy" values="11.2;14.2;11.2" dur="0.6s" repeatCount="indefinite"/>
    <animate attributeName="opacity" values="0.95;0.2;0.95" dur="0.6s" repeatCount="indefinite"/>
  </circle>
  <circle cx="11" cy="12" r="0.6" fill="#0B5ED7" opacity="0.80">
    <animate attributeName="cy" values="11.0;14.0;11.0" dur="0.7s" repeatCount="indefinite"/>
    <animate attributeName="opacity" values="0.8;0.15;0.8" dur="0.7s" repeatCount="indefinite"/>
  </circle>
  <circle cx="13" cy="12" r="0.6" fill="#0B5ED7" opacity="0.80">
    <animate attributeName="cy" values="11.4;14.4;11.4" dur="0.8s" repeatCount="indefinite"/>
    <animate attributeName="opacity" values="0.8;0.15;0.8" dur="0.8s" repeatCount="indefinite"/>
  </circle>
</svg>
"""

COUNTDOWN_SECONDS = 7  # ✅ 요청 확정: 7초

# =========================
# ✅ 체크 "순간" 감지 + 우측 카운트다운 렌더 유틸
# =========================
def _init_pledge_runtime(keys: list[str]) -> None:
    if "pledge_prev" not in st.session_state:
        st.session_state["pledge_prev"] = {k: False for k in keys}
    if "pledge_done" not in st.session_state:
        st.session_state["pledge_done"] = {k: False for k in keys}
    if "pledge_running" not in st.session_state:
        st.session_state["pledge_running"] = {k: False for k in keys}

def _order_enforce_cb(changed_key: str, prereq_keys: list[str], message: str) -> None:
    """체크 순서가 어긋나면 체크를 되돌리고, 경고 메시지를 세션에 기록합니다."""
    try:
        now_checked = bool(st.session_state.get(changed_key, False))
        prereq_ok = all(bool(st.session_state.get(k, False)) for k in prereq_keys)
        if now_checked and (not prereq_ok):
            st.session_state[changed_key] = False
            st.session_state["order_warning"] = message
    except Exception:
        pass

def _render_pledge_group(
    title: str,
    items: list[tuple[str, str]],
    all_keys: list[str],
    order_guard: dict | None = None,   # {"keys": [...], "prereq": [...], "message": "..."}
) -> None:
    st.markdown(f"### ■ {title}")

    guard_keys = set(order_guard.get("keys", [])) if isinstance(order_guard, dict) else set()
    prereq_keys = list(order_guard.get("prereq", [])) if isinstance(order_guard, dict) else []
    guard_msg = str(order_guard.get("message", "")) if isinstance(order_guard, dict) else ""

    for key, text in items:
        c1, c2, c3 = st.columns([0.06, 0.78, 0.16], vertical_alignment="center")

        with c1:
            cb_kwargs = dict(
                key=key,
                label_visibility="collapsed",
                disabled=bool(st.session_state["pledge_running"].get(key, False)),
            )

            # ✅ 관리자 서약을 임직원 서약보다 먼저 체크하려 하면: 체크를 되돌리고 토스트 경고
            if key in guard_keys:
                cb_kwargs.update(
                    dict(
                        on_change=_order_enforce_cb,
                        args=(key, prereq_keys, guard_msg),
                    )
                )

            st.checkbox("", **cb_kwargs)

        with c2:
            checked = bool(st.session_state.get(key, False))
            color = "#0B5ED7" if checked else "#2C3E50"
            weight = "900" if checked else "650"
            st.markdown(
                f"<div style='font-size:1.02rem; font-weight:{weight}; color:{color}; line-height:1.55;'>{text}</div>",
                unsafe_allow_html=True
            )

        with c3:
            ph = st.empty()
            now_checked = bool(st.session_state.get(key, False))
            prev_checked = bool(st.session_state["pledge_prev"].get(key, False))
            done = bool(st.session_state["pledge_done"].get(key, False))
            running = bool(st.session_state["pledge_running"].get(key, False))

            # ✅ 방금 체크된 순간에만 7초 카운트다운 실행
            if now_checked and (not prev_checked) and (not done) and (not running):
                st.session_state["pledge_running"][key] = True
                for sec in range(COUNTDOWN_SECONDS, 0, -1):
                    ph.markdown(
                        f"<div class='pledge-right'>{HOURGLASS_SVG}<span>{sec}s</span></div>",
                        unsafe_allow_html=True
                    )
                    time.sleep(1)
                st.session_state["pledge_running"][key] = False
                st.session_state["pledge_done"][key] = True
                ph.markdown(
                    "<div style='text-align:right; font-weight:900; color:#27AE60;'>✅ 완료</div>",
                    unsafe_allow_html=True
                )
            else:
                if running:
                    ph.markdown(
                        f"<div class='pledge-right'>{HOURGLASS_SVG}<span>...</span></div>",
                        unsafe_allow_html=True
                    )
                elif done and now_checked:
                    ph.markdown(
                        "<div style='text-align:right; font-weight:900; color:#27AE60;'>✅ 완료</div>",
                        unsafe_allow_html=True
                    )
                else:
                    ph.markdown("", unsafe_allow_html=True)


# --- [Top Tab: MY WORK LOG · 현장 기록 / 시설 이력] ---
with tab_worklog:
    # 조회를 닫았을 때 돌아올 MY WORK LOG 전용 상단 기준점입니다.
    st.markdown('<div id="worklog-top-anchor" style="height:1px;scroll-margin-top:18px;"></div>', unsafe_allow_html=True)
    st.markdown("""
    <style>
    .worklog-overview {
        display:grid;
        grid-template-columns:minmax(0,1fr) minmax(0,1fr);
        gap:12px;
        margin:8px 0 14px;
        align-items:stretch;
    }
    .worklog-hero {
        position:relative; overflow:hidden;
        background:linear-gradient(135deg,#FFFFFF 0%,#FFF7F7 48%,#F8FAFC 100%);
        border:1px solid #F1C7CA; border-left:8px solid #D71920;
        border-radius:18px; padding:15px 18px; margin:0;
        min-height:122px;
        box-shadow:0 8px 22px rgba(15,23,42,.07);
    }
    .worklog-hero:after {
        content:'📝'; position:absolute; right:18px; top:6px;
        font-size:64px; opacity:.08; transform:rotate(-6deg);
    }
    .worklog-hero .eyebrow { color:#D71920; font-size:.72rem; font-weight:950; letter-spacing:.12em; }
    .worklog-hero h2 { color:#B91218; margin:3px 0 2px; font-size:clamp(1.55rem,3vw,2.05rem); font-weight:950; letter-spacing:-.035em; }
    .worklog-hero .sub { color:#24364B; font-size:clamp(.96rem,2vw,1.12rem); font-weight:900; }
    .worklog-hero .desc { color:#64748B; margin-top:5px; font-size:.88rem; font-weight:750; line-height:1.45; padding-right:42px; }

    .worklog-section-title,
    .worklog-field-title {
        color:#24364B;
        font-size:1.08rem;
        font-weight:950;
        line-height:1.25;
        margin:8px 0 8px;
        letter-spacing:-.015em;
    }

    .worklog-dashboard {
        background:#FFFFFF;
        border:1px solid #DCE5F1;
        border-left:5px solid #0F4C81;
        border-radius:18px;
        padding:12px 13px;
        min-height:122px;
        display:flex;
        flex-direction:column;
        justify-content:center;
        box-shadow:0 8px 22px rgba(15,23,42,.06);
    }
    .worklog-dashboard-label {
        color:#64748B;
        font-size:.72rem;
        font-weight:900;
        letter-spacing:.08em;
        margin-bottom:8px;
    }
    .worklog-kpi-grid {
        display:grid;
        grid-template-columns:repeat(4,minmax(0,1fr));
        gap:6px;
        width:100%;
    }
    .worklog-kpi {
        min-width:0;
        display:flex;
        align-items:baseline;
        justify-content:center;
        gap:3px;
        padding:9px 3px;
        border-radius:11px;
        background:#F8FAFC;
        border:1px solid #E2E8F0;
        white-space:nowrap;
        overflow:hidden;
    }
    .worklog-kpi .label { color:#64748B; font-size:.66rem; font-weight:900; letter-spacing:-.025em; }
    .worklog-kpi .value { color:#0F3B66; font-size:1.08rem; font-weight:950; letter-spacing:-.04em; }
    .worklog-kpi .unit { color:#64748B; font-size:.62rem; font-weight:850; }
    .worklog-kpi.open { background:#FFF7ED; border-color:#FED7AA; }
    .worklog-kpi.open .value { color:#C2410C; }
    .worklog-kpi.doing { background:#FAF5FF; border-color:#E9D5FF; }
    .worklog-kpi.doing .value { color:#7E22CE; }
    .worklog-kpi.done { background:#F0FDF4; border-color:#BBF7D0; }
    .worklog-kpi.done .value { color:#15803D; }

    .worklog-storage-note {
        background:#EFF6FF; border:1px solid #BFDBFE; border-left:5px solid #2563EB;
        border-radius:14px; padding:9px 12px; color:#1E3A8A; font-size:.86rem; font-weight:750; line-height:1.45;
    }
    .worklog-storage-note.ready { background:#F0FDF4; border-color:#BBF7D0; border-left-color:#16A34A; color:#166534; }
    .worklog-storage-note.warn { background:#FFF7ED; border-color:#FED7AA; border-left-color:#F97316; color:#9A3412; }

    .worklog-card {
        background:#FFFFFF; border:1px solid #DCE5F1; border-radius:16px;
        padding:14px 15px; margin:0 0 10px; box-shadow:0 6px 16px rgba(15,23,42,.055);
    }
    .worklog-card-top { display:flex; align-items:flex-start; justify-content:space-between; gap:12px; }
    .worklog-place { color:#0F3B66; font-weight:950; font-size:1.08rem; }
    .worklog-meta { color:#64748B; font-size:.82rem; font-weight:750; margin-top:2px; }
    .worklog-body { color:#334155; font-weight:730; line-height:1.55; margin-top:8px; }
    .worklog-badge { display:inline-flex; padding:4px 9px; border-radius:999px; font-size:.78rem; font-weight:950; white-space:nowrap; }
    .worklog-badge.new { color:#B91C1C; background:#FEE2E2; }
    .worklog-badge.wait { color:#1D4ED8; background:#DBEAFE; }
    .worklog-badge.doing { color:#B45309; background:#FEF3C7; }
    .worklog-badge.recheck { color:#7E22CE; background:#F3E8FF; }
    .worklog-badge.done { color:#15803D; background:#DCFCE7; }

    .law-search-hero {
        display:flex; align-items:center; gap:12px; flex-wrap:wrap;
        background:linear-gradient(135deg,#EEF4FF,#F8FAFC); border:1px solid #C9D8EC;
        border-left:6px solid #24364B; border-radius:15px; padding:12px 15px; margin:6px 0 10px;
        color:#24364B;
    }
    .law-search-hero b { font-size:1.15rem; font-weight:950; }
    .law-search-hero span { color:#64748B; font-weight:750; }

    input[placeholder="예: 송포"]::placeholder {
        color:#94A3B8 !important;
        -webkit-text-fill-color:#94A3B8 !important;
        opacity:1 !important;
        font-weight:700 !important;
    }

    /* WORK LOG 조회 UI: PC는 대시보드 폭, 모바일은 공개 조회 시 조건+범위 / 검색어 / 불러오기 순으로 배치 */
    div[data-testid="stElementContainer"]:has(.worklog-search-row-marker),
    div[data-testid="stElementContainer"]:has(.worklog-recent-marker),
    div[data-testid="stElementContainer"]:has(.worklog-sticky-close-marker) {
        display:none !important;
    }

    @media (min-width:769px) {
        div[data-testid="stHorizontalBlock"]:has(.worklog-search-row-marker) {
            width:calc(50% - 6px) !important;
            margin-left:calc(50% + 6px) !important;
            gap:.45rem !important;
            align-items:flex-end !important;
        }
    }

    /* 조회 결과가 열려 있을 때만 생성되는 고정 닫기 버튼 */
    div[data-testid="stElementContainer"]:has(.worklog-sticky-close-marker)
      + div[data-testid="stElementContainer"],
    .st-key-worklog_results_close {
        position:fixed !important;
        left:50% !important;
        bottom:12px !important;
        transform:translateX(-50%) !important;
        width:min(410px, calc(100vw - 28px)) !important;
        z-index:9999 !important;
        padding:6px !important;
        background:rgba(255,255,255,.96) !important;
        border:1px solid #CBD5E1 !important;
        border-radius:14px !important;
        box-shadow:0 12px 30px rgba(15,23,42,.20) !important;
        backdrop-filter:blur(10px);
    }
    div[data-testid="stElementContainer"]:has(.worklog-sticky-close-marker)
      + div[data-testid="stElementContainer"] button,
    .st-key-worklog_results_close button {
        min-height:44px !important;
        background:#D71920 !important;
        color:#FFFFFF !important;
        font-weight:950 !important;
        border:none !important;
        border-radius:10px !important;
    }
    .worklog-close-safe-space { height:72px; }

    @media (max-width:768px) {
        .worklog-overview { grid-template-columns:1fr; gap:7px; margin:6px 0 10px; }
        .worklog-hero { padding:12px 13px; border-radius:15px; min-height:auto; }
        .worklog-hero:after { font-size:44px; right:7px; top:7px; }
        .worklog-hero .eyebrow { font-size:.64rem; }
        .worklog-hero h2 { font-size:1.42rem; }
        .worklog-hero .sub { font-size:.92rem; }
        .worklog-hero .desc { font-size:.75rem; line-height:1.36; padding-right:24px; }

        .worklog-dashboard { min-height:auto; padding:7px 6px; border-radius:14px; }
        .worklog-dashboard-label { display:none; }
        .worklog-kpi-grid { grid-template-columns:repeat(4,minmax(0,1fr)); gap:3px; }
        .worklog-kpi { padding:6px 1px; gap:2px; border-radius:8px; }
        .worklog-kpi .label { font-size:.54rem; letter-spacing:-.055em; }
        .worklog-kpi .value { font-size:.88rem; }
        .worklog-kpi .unit { font-size:.52rem; }

        .worklog-section-title,
        .worklog-field-title { font-size:1.02rem; font-weight:950; margin:8px 0 7px; }

        .worklog-card { padding:12px 11px; }

        input[placeholder="예: 송포"],
        input[placeholder="예: 정청운"],
        input[placeholder="필요한 추가 메모"],
        textarea[placeholder^="예: 축전지"],
        textarea[placeholder^="예: 단자"],
        textarea[placeholder^="예: 다음"] {
            font-size:16px !important;
        }

        /* 모바일: 검색 조건 + 검색어를 첫 줄, 불러오기를 둘째 줄 전체 폭으로 */
        div[data-testid="stHorizontalBlock"]:has(.worklog-search-row-marker) {
            display:grid !important;
            grid-template-columns:minmax(0,.82fr) minmax(0,1.38fr) !important;
            gap:6px !important;
            width:100% !important;
            margin:0 !important;
            align-items:end !important;
        }
        /* Streamlit 버전에 따라 컬럼 testid가 column / stColumn으로 달라질 수 있어 둘 다 대응 */
        div[data-testid="stHorizontalBlock"]:has(.worklog-search-row-marker)
          > div[data-testid="column"],
        div[data-testid="stHorizontalBlock"]:has(.worklog-search-row-marker)
          > div[data-testid="stColumn"] {
            width:100% !important;
            max-width:none !important;
            min-width:0 !important;
            flex:unset !important;
        }
        div[data-testid="stHorizontalBlock"]:has(.worklog-search-row-marker)
          > div[data-testid="column"]:nth-child(1),
        div[data-testid="stHorizontalBlock"]:has(.worklog-search-row-marker)
          > div[data-testid="stColumn"]:nth-child(1) { grid-column:1; grid-row:1; }
        div[data-testid="stHorizontalBlock"]:has(.worklog-search-row-marker)
          > div[data-testid="column"]:nth-child(2),
        div[data-testid="stHorizontalBlock"]:has(.worklog-search-row-marker)
          > div[data-testid="stColumn"]:nth-child(2) { grid-column:2; grid-row:1; }
        div[data-testid="stHorizontalBlock"]:has(.worklog-search-row-marker)
          > div[data-testid="column"]:nth-child(3),
        div[data-testid="stHorizontalBlock"]:has(.worklog-search-row-marker)
          > div[data-testid="stColumn"]:nth-child(3) {
            grid-column:1 / -1 !important;
            grid-row:2 !important;
            width:100% !important;
            max-width:none !important;
        }

        /* V15 공개 조회: 1행=검색조건+내/전체, 2행=검색어, 3행=불러오기 전체 폭 */
        div[data-testid="stHorizontalBlock"]:has(.worklog-public-scope-marker) {
            grid-template-columns:minmax(0,1fr) minmax(0,1fr) !important;
        }
        div[data-testid="stHorizontalBlock"]:has(.worklog-public-scope-marker)
          > div[data-testid="column"]:nth-child(1),
        div[data-testid="stHorizontalBlock"]:has(.worklog-public-scope-marker)
          > div[data-testid="stColumn"]:nth-child(1) { grid-column:1; grid-row:1; }
        div[data-testid="stHorizontalBlock"]:has(.worklog-public-scope-marker)
          > div[data-testid="column"]:nth-child(2),
        div[data-testid="stHorizontalBlock"]:has(.worklog-public-scope-marker)
          > div[data-testid="stColumn"]:nth-child(2) { grid-column:2; grid-row:1; }
        div[data-testid="stHorizontalBlock"]:has(.worklog-public-scope-marker)
          > div[data-testid="column"]:nth-child(3),
        div[data-testid="stHorizontalBlock"]:has(.worklog-public-scope-marker)
          > div[data-testid="stColumn"]:nth-child(3) {
            grid-column:1 / -1 !important; grid-row:2 !important; width:100% !important; max-width:none !important;
        }
        div[data-testid="stHorizontalBlock"]:has(.worklog-public-scope-marker)
          > div[data-testid="column"]:nth-child(4),
        div[data-testid="stHorizontalBlock"]:has(.worklog-public-scope-marker)
          > div[data-testid="stColumn"]:nth-child(4) {
            grid-column:1 / -1 !important; grid-row:3 !important; width:100% !important; max-width:none !important;
        }
        /* 불러오기 버튼은 모바일에서 검색영역 전체 폭 + 한 줄 고정 */
        .st-key-worklog_refresh {
            width:100% !important;
            max-width:none !important;
        }
        .st-key-worklog_refresh button {
            width:100% !important;
            max-width:none !important;
            min-height:46px !important;
            white-space:nowrap !important;
        }

        /* 모바일: 기존 2열의 순서만 뒤집어 불러오기 바로 아래에 최근 기록 표시 */
        div[data-testid="stHorizontalBlock"]:has(.worklog-recent-marker) {
            display:flex !important;
            flex-direction:column-reverse !important;
            gap:.7rem !important;
        }
        div[data-testid="stHorizontalBlock"]:has(.worklog-recent-marker)
          > div[data-testid="column"] {
            width:100% !important;
            min-width:0 !important;
            flex:1 1 100% !important;
        }

        div[data-testid="stElementContainer"]:has(.worklog-sticky-close-marker)
          + div[data-testid="stElementContainer"],
        .st-key-worklog_results_close {
            bottom:max(10px, env(safe-area-inset-bottom)) !important;
            width:calc(100vw - 22px) !important;
        }
        .worklog-close-safe-space { height:78px; }
    }

    .worklog-auth-card {
        background:linear-gradient(135deg,#FFFFFF 0%,#F8FAFC 100%);
        border:1px solid #CBD5E1;
        border-left:6px solid #D71920;
        border-radius:18px;
        padding:18px 20px;
        margin:8px 0 14px;
        box-shadow:0 8px 22px rgba(15,23,42,.07);
    }
    .worklog-auth-title { color:#24364B; font-size:1.28rem; font-weight:950; margin-bottom:5px; }
    .worklog-auth-desc { color:#64748B; font-size:.92rem; font-weight:760; line-height:1.55; }
    .worklog-userbar {
        display:flex; align-items:center; justify-content:space-between; gap:10px;
        background:#EEF6FF; border:1px solid #BFDBFE; border-radius:14px;
        padding:10px 13px; margin:4px 0 10px;
    }
    .worklog-userbar .name { color:#0F3B66; font-size:1rem; font-weight:950; }
    .worklog-userbar .meta { color:#64748B; font-size:.78rem; font-weight:760; margin-top:2px; }
    .worklog-privacy-guide {
        display:grid; grid-template-columns:1fr 1fr; gap:8px; margin:6px 0 8px;
    }
    .worklog-privacy-guide .public,
    .worklog-privacy-guide .private {
        border-radius:12px; padding:9px 11px; line-height:1.4;
        font-size:.82rem; font-weight:760;
    }
    .worklog-privacy-guide .public { background:#EFF6FF; border:1px solid #BFDBFE; color:#1E40AF; }
    .worklog-privacy-guide .private { background:#FFF7ED; border:1px solid #FED7AA; color:#9A3412; }
    .worklog-privacy-badge {
        display:inline-flex; align-items:center; padding:3px 8px; border-radius:999px;
        font-size:.72rem; font-weight:950; white-space:nowrap; margin-left:5px;
    }
    .worklog-privacy-badge.public { background:#DBEAFE; color:#1D4ED8; }
    .worklog-privacy-badge.private { background:#FFEDD5; color:#C2410C; }
    .worklog-visibility-manage {
        background:#F8FAFC; border:1px solid #CBD5E1; border-left:5px solid #2563EB;
        border-radius:14px; padding:11px 13px; margin:12px 0 8px;
        color:#334155; font-size:.86rem; font-weight:760; line-height:1.5;
    }
    .worklog-visibility-manage.private {
        background:#FFF7ED; border-color:#FED7AA; border-left-color:#F97316; color:#9A3412;
    }
    .worklog-delete-box {
        background:#FFF7F7; border:1px solid #FECACA; border-left:5px solid #DC2626;
        border-radius:14px; padding:11px 13px; margin-top:12px;
    }

    @media (max-width:768px) {
        .worklog-auth-card { padding:14px 13px; border-radius:15px; }
        .worklog-privacy-guide { grid-template-columns:1fr; gap:5px; }
        .worklog-userbar { align-items:flex-start; }
    }
    </style>
    """, unsafe_allow_html=True)


    if st.session_state.pop("worklog_scroll_to_top", False):
        components.html(
            """
            <script>
            (function () {
              let tries = 0;
              function goTop() {
                try {
                  const w = window.parent;
                  const d = w.document;
                  const anchor = d.getElementById('worklog-top-anchor');
                  if (anchor) {
                    anchor.scrollIntoView({ behavior: 'smooth', block: 'start' });
                    return;
                  }
                } catch (e) {}
                tries += 1;
                if (tries < 24) window.setTimeout(goTop, 70);
              }
              window.setTimeout(goTop, 90);
            })();
            </script>
            """,
            height=1,
        )

    worklog_auth_user = _worklog_current_user()

    if not worklog_auth_user:
        st.markdown(
            '<div class="worklog-auth-card">'
            '<div class="worklog-auth-title">🔐 MY WORK LOG 개인 인증</div>'
            '<div class="worklog-auth-desc">'
            'MY WORK LOG는 팀 공용 업무기록과 개인 업무메모를 함께 사용합니다. '
            '개인 인증 후 공개 기록은 팀과 공유하고, 비공개 기록은 작성한 본인만 검색·조회할 수 있습니다.<br><br>'
            '<b>🌐 공개 · 팀 공유</b> — 업무 공유가 필요한 기록<br>'
            '<b>🔒 비공개 · 나만 보기</b> — 기억용 메모·개인 업무기록'
            '</div></div>',
            unsafe_allow_html=True,
        )
        auth_button_col, auth_info_col = st.columns([1.5, 3.5], gap="small", vertical_alignment="center")
        with auth_button_col:
            if st.button(
                "🔐 개인 인증 시작",
                key="worklog_open_login",
                type="primary",
                use_container_width=True,
            ):
                _worklog_login_dialog()
        with auth_info_col:
            st.caption("사번 + 개인 PIN으로 로그인합니다. 최초 로그인 공통 PIN은 000000이며, 로그인 직후 본인 PIN으로 반드시 변경합니다.")

    elif bool(st.session_state.get("worklog_pin_change_required", False)):
        st.markdown(
            '<div class="worklog-auth-card">'
            '<div class="worklog-auth-title">🔑 최초 로그인 · 개인 PIN 설정</div>'
            '<div class="worklog-auth-desc">최초 로그인 공통 PIN 000000은 한 번만 사용합니다. '
            '앞으로 사용할 숫자 6자리 개인 PIN으로 변경해 주세요.</div></div>',
            unsafe_allow_html=True,
        )
        with st.form("worklog_first_pin_change_form", clear_on_submit=False):
            first_new_pin = st.text_input("새 개인 PIN", type="password", max_chars=6, key="worklog_first_new_pin")
            first_confirm_pin = st.text_input("새 PIN 확인", type="password", max_chars=6, key="worklog_first_confirm_pin")
            first_pin_submit = st.form_submit_button("🔐 개인 PIN 설정 완료", type="primary", use_container_width=True)
        if first_pin_submit:
            first_ok, first_message = _worklog_change_pin(
                str(worklog_auth_user.get("employee_no", "") or ""),
                first_new_pin,
                first_confirm_pin,
            )
            if first_ok:
                st.session_state["worklog_pin_change_required"] = False
                st.success(first_message)
                time.sleep(0.4)
                st.rerun()
            else:
                st.error(first_message)
        st.button("로그아웃", key="worklog_first_login_logout", on_click=_worklog_logout)

    else:
        auth_user = worklog_auth_user
        user_employee = str(auth_user.get("employee_no", "") or "")
        masked_employee = ("••••" + user_employee[-4:]) if len(user_employee) >= 4 else user_employee
        userbar_col, pin_col, logout_col = st.columns([5.2, 1.45, 1.25], gap="small", vertical_alignment="center")
        with userbar_col:
            st.markdown(
                f'<div class="worklog-userbar"><div><div class="name">👤 {html.escape(str(auth_user.get("name","")))}</div>'
                f'<div class="meta">MY WORK LOG 개인 인증 · 사번 {html.escape(masked_employee)}</div></div></div>',
                unsafe_allow_html=True,
            )
        with pin_col:
            if st.button("🔑 PIN 변경", key="worklog_pin_change_toggle", use_container_width=True):
                st.session_state["worklog_show_pin_change"] = not bool(st.session_state.get("worklog_show_pin_change", False))
        with logout_col:
            st.button("로그아웃", key="worklog_logout", use_container_width=True, on_click=_worklog_logout)

        if st.session_state.get("worklog_show_pin_change"):
            with st.container(border=True):
                st.markdown("#### 🔑 개인 PIN 변경")
                with st.form("worklog_regular_pin_change_form", clear_on_submit=True):
                    current_pin = st.text_input("현재 PIN", type="password", max_chars=6)
                    new_pin = st.text_input("새 PIN", type="password", max_chars=6)
                    new_pin_confirm = st.text_input("새 PIN 확인", type="password", max_chars=6)
                    pin_change_submit = st.form_submit_button("PIN 변경", type="primary", use_container_width=True)
                if pin_change_submit:
                    verify_ok, verify_message, _ = _worklog_authenticate_user(user_employee, current_pin)
                    if not verify_ok:
                        st.error(f"현재 PIN 확인 실패: {verify_message}")
                    else:
                        pin_ok, pin_message = _worklog_change_pin(user_employee, new_pin, new_pin_confirm)
                        if pin_ok:
                            st.session_state["worklog_show_pin_change"] = False
                            st.success(pin_message)
                            time.sleep(0.4)
                            st.rerun()
                        else:
                            st.error(pin_message)

        if "worklog_df" not in st.session_state:
            st.session_state["worklog_df"] = None
        if "worklog_loaded_at" not in st.session_state:
            st.session_state["worklog_loaded_at"] = ""
        if "worklog_selected_id" not in st.session_state:
            st.session_state["worklog_selected_id"] = ""

        def _worklog_close_loaded_results():
            """조회 결과만 닫고 새 현장기록 작성 중 입력값은 보존한 뒤 WORK LOG 상단으로 이동합니다."""
            st.session_state["worklog_df"] = None
            st.session_state["worklog_loaded_at"] = ""
            st.session_state["worklog_selected_id"] = ""
            st.session_state["worklog_search"] = ""
            st.session_state["worklog_filter"] = "전체"
            st.session_state["worklog_public_scope"] = "👤 내 기록"
            st.session_state["worklog_scroll_to_top"] = True
            for reset_key in (
                "worklog_update_writer", "worklog_update_status", "worklog_update_action",
                "worklog_update_followup", "worklog_update_remark",
            ):
                st.session_state.pop(reset_key, None)

        # PC/모바일 공용 상단: 왼쪽 MY WORK LOG 소개 + 오른쪽 미니 대시보드
        worklog_overview_slot = st.empty()

        photo_upload_ready, photo_config_issues = _worklog_photo_config_status()
        # 정상 운영 시 사진 저장 안내 문구는 표시하지 않습니다.
        # 설정 이상이 있을 때만 오류 원인을 보여 기존 사진 저장 안정성은 유지합니다.
        if not photo_upload_ready:
            issue_html = html.escape(" / ".join(photo_config_issues) if photo_config_issues else "사진 업로드 설정 확인 필요")
            st.markdown(
                f'<div class="worklog-storage-note warn">📷 사진 업로드 설정 확인 필요 · PHOTO ENGINE <b>{WORK_LOG_PHOTO_ENGINE_VERSION}</b><br><b>현재 진단:</b> {issue_html}<br>Streamlit Secrets의 <b>[work_log]</b>에 <b>photo_upload_url</b>과 <b>upload_token</b>을 확인해 주세요. 기존 내비·정밀점검 기능에는 영향이 없습니다.</div>',
                unsafe_allow_html=True,
            )

        # V15: 공개 조회는 "내 기록 / 전체 기록"을 선택할 수 있게 하고 검색어 폭을 줄입니다.
        public_filter_for_layout = str(st.session_state.get("worklog_filter", "전체") or "전체") == "🌐 공개"
        if public_filter_for_layout:
            search_condition_col, public_scope_col, search_text_col, search_load_col = st.columns(
                [0.24, 0.23, 0.31, 0.22], gap="small", vertical_alignment="bottom"
            )
        else:
            search_condition_col, search_text_col, search_load_col = st.columns(
                [0.30, 0.46, 0.24], gap="small", vertical_alignment="bottom"
            )
            public_scope_col = None

        with search_condition_col:
            st.markdown('<div class="worklog-search-row-marker"></div>', unsafe_allow_html=True)
            worklog_filter = st.selectbox(
                "검색 조건",
                ["전체", "🌐 공개", "🔒 내 비공개"] + WORK_LOG_STATUS_OPTIONS,
                key="worklog_filter",
            )

        worklog_public_scope = str(st.session_state.get("worklog_public_scope", "👤 내 기록") or "👤 내 기록")
        if worklog_filter == "🌐 공개":
            # Streamlit은 위젯 변경 시 즉시 rerun되므로 공개 선택 이후에는 4열 레이아웃으로 다시 그려집니다.
            if public_scope_col is not None:
                with public_scope_col:
                    st.markdown('<div class="worklog-public-scope-marker"></div>', unsafe_allow_html=True)
                    worklog_public_scope = st.selectbox(
                        "공개 기록 범위",
                        ["👤 내 기록", "👥 전체 기록"],
                        key="worklog_public_scope",
                        help="내 기록은 로그인한 본인이 작성한 공개 기록만, 전체 기록은 모든 사용자의 공개 기록을 조회합니다.",
                    )
            else:
                st.session_state["worklog_public_scope"] = "👤 내 기록"
                worklog_public_scope = "👤 내 기록"
        else:
            worklog_public_scope = "👤 내 기록"

        with search_text_col:
            worklog_search = st.text_input(
                "검색어",
                placeholder="국사 · 작성자 · 점검항목 등",
                key="worklog_search",
            ).strip()
        with search_load_col:
            refresh_worklog = st.button(
                "🔄 불러오기",
                use_container_width=True,
                type="primary",
                key="worklog_refresh",
            )

        if refresh_worklog:
            with st.spinner("Google Sheets에서 MY WORK LOG를 불러오는 중입니다..."):
                try:
                    st.session_state["worklog_df"] = load_work_logs(auth_user)
                    st.session_state["worklog_loaded_at"] = _korea_now().strftime("%Y-%m-%d %H:%M:%S")
                    st.session_state["worklog_selected_id"] = ""
                except Exception as error:
                    st.error(f"WORK LOG를 불러오지 못했습니다: {error}")

        loaded_df = st.session_state.get("worklog_df")
        if isinstance(loaded_df, pd.DataFrame):
            summary_source = loaded_df.copy()
            count_open = int(summary_source["상태"].isin(["신규", "확인필요"]).sum()) if not summary_source.empty else 0
            count_doing = int(summary_source["상태"].isin(["조치중", "재점검"]).sum()) if not summary_source.empty else 0
            count_done = int((summary_source["상태"] == "완료").sum()) if not summary_source.empty else 0
            count_all = int(len(summary_source))
        else:
            count_all = count_open = count_doing = count_done = 0

        count_all_text = f"{count_all:,}" if loaded_df is not None else "-"
        count_open_text = f"{count_open:,}" if loaded_df is not None else "-"
        count_doing_text = f"{count_doing:,}" if loaded_df is not None else "-"
        count_done_text = f"{count_done:,}" if loaded_df is not None else "-"

        worklog_overview_slot.markdown(
            f"""
            <div class="worklog-overview">
              <div class="worklog-hero">
                <div class="eyebrow">FIELD HISTORY &amp; RECORD</div>
                <h2>MY WORK LOG</h2>
                <div class="sub">현장 기록 · 시설 이력</div>
                <div class="desc">오늘의 현장 기록이 내일의 정확한 업무가 됩니다. 상태이력·점검항목·사진·조치사항을 국사별로 연결해 관리합니다.</div>
              </div>
              <div class="worklog-dashboard">
                <div class="worklog-dashboard-label">WORK LOG STATUS</div>
                <div class="worklog-kpi-grid">
                  <div class="worklog-kpi">
                    <span class="label">전체 기록</span><span class="value">{count_all_text}</span><span class="unit">건</span>
                  </div>
                  <div class="worklog-kpi open">
                    <span class="label">미처리</span><span class="value">{count_open_text}</span><span class="unit">건</span>
                  </div>
                  <div class="worklog-kpi doing">
                    <span class="label">진행·재점검</span><span class="value">{count_doing_text}</span><span class="unit">건</span>
                  </div>
                  <div class="worklog-kpi done">
                    <span class="label">완료</span><span class="value">{count_done_text}</span><span class="unit">건</span>
                  </div>
                </div>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        if st.session_state.get("worklog_loaded_at"):
            st.caption(f"최근 기록 조회시각: {st.session_state['worklog_loaded_at']} · 화면 진입만으로는 Google Sheets를 자동 조회하지 않습니다.")

        entry_col, recent_col = st.columns([0.94, 1.06], gap="large")

        with entry_col:
            with st.container(border=True):
                st.markdown('<div class="worklog-section-title">📝 새 현장기록</div>', unsafe_allow_html=True)
                st.caption("현장에서 10초 안에 남기고, 나중에 국사별 이력으로 다시 찾을 수 있도록 핵심 항목만 구성했습니다.")

                st.markdown('<div class="worklog-field-title">✍️ 작성자</div>', unsafe_allow_html=True)
                writer = str(auth_user.get("name", "") or "").strip()
                st.markdown(
                    f'<div class="worklog-userbar"><div><div class="name">👤 {html.escape(writer)}</div>'
                    f'<div class="meta">개인 인증된 사용자 · 작성자는 자동으로 고정됩니다.</div></div></div>',
                    unsafe_allow_html=True,
                )

                st.markdown('<div class="worklog-field-title">🔐 공개범위</div>', unsafe_allow_html=True)
                visibility_label = st.radio(
                    "공개범위",
                    ["🌐 공개 · 팀 공유", "🔒 비공개 · 나만 보기"],
                    horizontal=True,
                    key="worklog_visibility",
                    label_visibility="collapsed",
                )
                visibility = "비공개" if visibility_label.startswith("🔒") else "공개"
                st.markdown(
                    '<div class="worklog-privacy-guide">'
                    '<div class="public"><b>🌐 공개 · 팀 공유</b><br>업무 공유가 필요한 기록입니다. 인증된 팀원이 검색·조회할 수 있습니다.</div>'
                    '<div class="private"><b>🔒 비공개 · 나만 보기</b><br>개인 업무메모입니다. 작성한 본인만 검색·조회·사진 다운로드할 수 있습니다.</div>'
                    '</div>',
                    unsafe_allow_html=True,
                )

                st.markdown('<div class="worklog-field-title">📍 국사 검색 · 자동입력</div>', unsafe_allow_html=True)
                st.caption("정밀점검과 같은 국사 기준정보를 사용합니다. 국사명만 검색·선택하면 권역·모국·국소가 자동 반영됩니다.")

                with st.form(key="worklog_station_search_form", clear_on_submit=False):
                    station_search_col, station_search_btn_col = st.columns([4.2, 1.15], gap="small")
                    with station_search_col:
                        st.text_input(
                            "국사명",
                            key="worklog_station_search_query",
                            placeholder="예: 송포",
                            help="국사명을 입력한 뒤 확인을 누르세요. 키보드 Enter로도 검색할 수 있습니다.",
                            label_visibility="collapsed",
                        )
                    with station_search_btn_col:
                        worklog_station_search_submitted = st.form_submit_button("확인", use_container_width=True)

                if worklog_station_search_submitted:
                    _run_worklog_station_search()

                worklog_station_search_status = str(
                    st.session_state.get("worklog_station_search_status", "") or ""
                )
                worklog_station_candidate_ids = list(
                    st.session_state.get("worklog_station_search_candidates", []) or []
                )

                if worklog_station_search_status == "empty":
                    st.warning("국사명을 먼저 입력해 주세요. 예: 송포")
                elif worklog_station_search_status == "none":
                    st.warning("일치하는 국사를 찾지 못했습니다. 국사명을 다시 확인해 주세요.")
                elif worklog_station_search_status == "multiple" and worklog_station_candidate_ids:
                    st.info(
                        f"같은 이름 또는 유사한 국사가 {len(worklog_station_candidate_ids)}곳 있습니다. "
                        "아래에서 정확한 국사를 선택해 주세요."
                    )
                    with st.form(key="worklog_station_duplicate_form", clear_on_submit=False):
                        st.radio(
                            "국사 선택",
                            worklog_station_candidate_ids,
                            key="worklog_station_search_choice",
                            format_func=_worklog_station_search_label,
                        )
                        worklog_station_choice_submitted = st.form_submit_button(
                            "선택 확인", use_container_width=True
                        )
                    if worklog_station_choice_submitted:
                        _confirm_worklog_station_search_choice()

                worklog_station_notice = str(
                    st.session_state.get("worklog_station_search_notice", "") or ""
                )
                if worklog_station_notice:
                    st.success(worklog_station_notice)

                area = str(st.session_state.get("worklog_area_key", "") or "").strip()
                mother = str(st.session_state.get("worklog_mother", "") or "").strip()
                local = str(st.session_state.get("worklog_local", "") or "").strip()
                area_display = _worklog_area_display(area)

                st.markdown(
                    f"""<div style="display:grid;grid-template-columns:1fr 1fr;gap:8px;margin:6px 0 12px;">
                        <div style="grid-column:1/-1;background:#FAF5FF;border:1px solid #C4B5FD;border-radius:11px;padding:10px 12px;">
                            <div style="font-size:.78rem;font-weight:900;color:#6D28D9;margin-bottom:3px;">권역</div>
                            <div style="font-weight:950;color:#5B21B6;line-height:1.45;">{html.escape(area_display)}</div>
                        </div>
                        <div style="background:#F8FAFC;border:1px solid #CBD5E1;border-radius:11px;padding:10px 12px;">
                            <div style="font-size:.78rem;font-weight:900;color:#64748B;margin-bottom:3px;">모국</div>
                            <div style="font-weight:900;color:#24364B;">{html.escape(mother or '자동 표시')}</div>
                        </div>
                        <div style="background:#F8FAFC;border:1px solid #CBD5E1;border-radius:11px;padding:10px 12px;">
                            <div style="font-size:.78rem;font-weight:900;color:#64748B;margin-bottom:3px;">국소</div>
                            <div style="font-weight:900;color:#24364B;">{html.escape(local or '자동 표시')}</div>
                        </div>
                    </div>""",
                    unsafe_allow_html=True,
                )

                st.markdown('<div class="worklog-field-title">🔄 상태이력</div>', unsafe_allow_html=True)
                status = st.radio(
                    "상태이력",
                    WORK_LOG_STATUS_OPTIONS,
                    horizontal=True,
                    key="worklog_status",
                    label_visibility="collapsed",
                )

                st.markdown('<div class="worklog-field-title">🧰 점검항목</div>', unsafe_allow_html=True)
                item_select_col, item_ok_col = st.columns([4.35, 1.0], gap="small", vertical_alignment="center")
                with item_select_col:
                    items = st.multiselect(
                        "점검항목",
                        WORK_LOG_ITEM_OPTIONS,
                        placeholder="전원 · 축전지 · 접지 · 냉방 · 출입 · 안전 · 기타",
                        key="worklog_items",
                        label_visibility="collapsed",
                    )
                with item_ok_col:
                    worklog_items_ok = st.button(
                        "확인",
                        key="worklog_items_ok",
                        use_container_width=True,
                    )
                if worklog_items_ok:
                    selected_item_text = ", ".join(items) if items else "선택 없음"
                    st.session_state["worklog_items_confirmed_notice"] = f"선택 완료: {selected_item_text}"
                if st.session_state.get("worklog_items_confirmed_notice"):
                    st.caption(st.session_state["worklog_items_confirmed_notice"])

                st.markdown('<div class="worklog-field-title">📷 현장사진</div>', unsafe_allow_html=True)
                photo_c1, photo_c2 = st.columns(2)
                with photo_c1:
                    camera_photo = st.camera_input("현장에서 바로 촬영", key="worklog_camera")
                with photo_c2:
                    uploaded_photos = st.file_uploader(
                        "앨범/파일에서 선택",
                        type=["jpg", "jpeg", "png", "webp"],
                        accept_multiple_files=True,
                        key="worklog_uploads",
                    ) or []
                photos = _worklog_collect_photos(camera_photo, uploaded_photos)
                st.caption(f"선택 사진 {len(photos)}장 / 최대 {WORK_LOG_MAX_PHOTOS}장 · 저장 시 자동 압축(최대 변 1600px, 목표 약 450KB/장)")

                st.markdown('<div class="worklog-field-title">📝 현상·특이사항</div>', unsafe_allow_html=True)
                issue = st.text_area(
                    "현상·특이사항",
                    placeholder="예: 축전지 1조 7번 셀 전압이 다른 셀보다 낮게 측정됨",
                    height=105,
                    key="worklog_issue",
                    label_visibility="collapsed",
                )
                st.markdown('<div class="worklog-field-title">🛠️ 조치내용</div>', unsafe_allow_html=True)
                action = st.text_area(
                    "조치내용",
                    placeholder="예: 단자 상태 확인 및 재측정",
                    height=85,
                    key="worklog_action",
                    label_visibility="collapsed",
                )
                st.markdown('<div class="worklog-field-title">🔁 후속조치</div>', unsafe_allow_html=True)
                followup = st.text_area(
                    "후속조치",
                    placeholder="예: 다음 방문 시 1조 7번 셀 재확인",
                    height=85,
                    key="worklog_followup",
                    label_visibility="collapsed",
                )
                st.markdown('<div class="worklog-field-title">📌 비고</div>', unsafe_allow_html=True)
                remark = st.text_input(
                    "비고",
                    placeholder="필요한 추가 메모",
                    key="worklog_remark",
                    label_visibility="collapsed",
                )

                save_log = st.button(
                    "💾 MY WORK LOG 저장",
                    use_container_width=True,
                    type="primary",
                    key="worklog_save",
                )
                if save_log:
                    record = {
                        "작성자": writer,
                        "작성자ID": str(auth_user.get("user_id", "") or "").strip(),
                        "공개범위": visibility,
                        "권역": area,
                        "모국": mother,
                        "국소": local,
                        "상태": status,
                        "점검항목": items,
                        "현상_특이사항": issue,
                        "조치내용": action,
                        "후속조치": followup,
                        "비고": remark,
                    }
                    with st.spinner("현장 기록과 사진을 안전하게 저장하고 있습니다..."):
                        ok, message, record_id = save_work_log(record, photos)
                    if ok:
                        st.success(f"✅ {message} · 기록ID: {record_id}")
                        try:
                            st.session_state["worklog_df"] = load_work_logs(auth_user)
                            st.session_state["worklog_loaded_at"] = _korea_now().strftime("%Y-%m-%d %H:%M:%S")
                        except Exception:
                            pass
                        _worklog_reset_entry_widgets()
                        time.sleep(0.8)
                        st.rerun()
                    else:
                        st.error(f"❌ {message}")

        def _render_worklog_inline_detail(selected, selected_id: str) -> None:
            """선택한 최근 기록의 기존 상세·조치 화면을 해당 카드 바로 아래에 렌더링합니다."""
            st.markdown("---")
            st.markdown("### 📷 첨부 현장사진")
            st.caption(
                f"기록ID: {selected_id} · {selected.get('모국','')} / {selected.get('국소','')} · "
                "Drive 폴더는 공개하지 않고 이 화면에서만 조회·다운로드합니다."
            )

            selected_photo_ids = [
                v.strip() for v in str(selected.get("사진파일ID목록", "") or "").split("|") if v.strip()
            ]
            selected_photo_names = [
                v.strip() for v in str(selected.get("사진파일명목록", "") or "").split("|") if v.strip()
            ]

            photo_payloads = []
            failed_photo_count = 0
            local_for_name = str(selected.get("국소", "") or selected.get("모국", "") or "WORK_LOG").strip()
            writer_for_name = str(selected.get("작성자", "") or "현장").strip()
            saved_for_name = re.sub(r"[^0-9]", "", str(selected.get("저장일시", "") or ""))[:8] or _korea_now().strftime("%Y%m%d")
            item_for_name = str(selected.get("점검항목", "") or "현장사진").split(",")[0].strip() or "현장사진"

            safe_local = re.sub(r"[^0-9A-Za-z가-힣_-]", "_", local_for_name)[:40] or "WORK_LOG"
            safe_writer = re.sub(r"[^0-9A-Za-z가-힣_-]", "_", writer_for_name)[:30] or "현장"
            safe_item = re.sub(r"[^0-9A-Za-z가-힣_-]", "_", item_for_name)[:30] or "현장사진"

            for photo_index, file_id in enumerate(selected_photo_ids, 1):
                photo_bytes = _worklog_download_drive_image(file_id)
                if not photo_bytes:
                    failed_photo_count += 1
                    continue
                download_name = f"{safe_local}_{safe_item}_{safe_writer}_{saved_for_name}_{photo_index:02d}.jpg"
                stored_name = selected_photo_names[photo_index - 1] if photo_index - 1 < len(selected_photo_names) else download_name
                photo_payloads.append({
                    "index": photo_index,
                    "id": file_id,
                    "bytes": photo_bytes,
                    "download_name": download_name,
                    "stored_name": stored_name,
                })

            if photo_payloads:
                st.caption(f"첨부사진 {len(selected_photo_ids)}장 · 전체 보기 및 개별/일괄 다운로드")
                photo_cols = st.columns(2)
                for payload_pos, payload in enumerate(photo_payloads):
                    with photo_cols[payload_pos % 2]:
                        st.image(
                            payload["bytes"],
                            caption=f"사진 {payload['index']} / {len(selected_photo_ids)}",
                            use_container_width=True,
                        )
                        st.download_button(
                            "📥 사진 다운로드",
                            data=payload["bytes"],
                            file_name=payload["download_name"],
                            mime="image/jpeg",
                            use_container_width=True,
                            key=f"worklog_photo_download_{selected_id}_{payload['index']}",
                        )

                try:
                    from io import BytesIO
                    import zipfile

                    zip_buffer = BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as photo_zip:
                        for payload in photo_payloads:
                            photo_zip.writestr(payload["download_name"], payload["bytes"])
                    zip_file_name = f"{safe_local}_WORK_LOG_{saved_for_name}_사진{len(photo_payloads)}장.zip"
                    st.download_button(
                        f"📦 사진 {len(photo_payloads)}장 전체 ZIP 다운로드",
                        data=zip_buffer.getvalue(),
                        file_name=zip_file_name,
                        mime="application/zip",
                        use_container_width=True,
                        type="primary",
                        key=f"worklog_photo_zip_{selected_id}",
                    )
                except Exception as zip_error:
                    st.warning(f"사진 ZIP 파일을 만들지 못했습니다. 개별 다운로드를 이용해 주세요. ({zip_error})")

                if failed_photo_count:
                    st.warning(f"첨부사진 중 {failed_photo_count}장은 현재 읽을 수 없어 표시하지 못했습니다.")
            elif selected_photo_ids:
                st.warning("첨부사진 정보는 있으나 현재 사진 파일을 읽을 수 없습니다. Drive 읽기 권한을 확인해 주세요.")
            else:
                st.info("이 기록에는 첨부된 현장사진이 없습니다.")

            # V13: 기록 소유자는 과거/현재 자료 모두 공개 ↔ 비공개를 언제든 변경할 수 있습니다.
            is_owner = _worklog_record_owned_by(selected, auth_user)
            current_visibility = _worklog_record_visibility(selected)
            current_visibility_label = (
                "🌐 공개 · 팀 공유" if current_visibility == "공개" else "🔒 비공개 · 나만 보기"
            )
            visibility_box_class = "private" if current_visibility == "비공개" else "public"
            st.markdown("### 🔐 공개범위 관리")
            st.markdown(
                f'<div class="worklog-visibility-manage {visibility_box_class}">'
                f'<b>현재 공개범위: {html.escape(current_visibility_label)}</b><br>'
                '공개 기록은 인증된 팀원이 검색·조회할 수 있고, 비공개 기록은 작성한 본인만 검색·조회·사진 다운로드할 수 있습니다. '
                '공개범위를 바꾸면 연결 사진도 같은 권한을 따르며, 다른 사용자의 화면에는 다음 조회부터 반영됩니다.</div>',
                unsafe_allow_html=True,
            )

            if is_owner:
                visibility_options = ["🌐 공개 · 팀 공유", "🔒 비공개 · 나만 보기"]
                visibility_manage_key = f"worklog_manage_visibility_{selected_id}"
                if visibility_manage_key not in st.session_state:
                    st.session_state[visibility_manage_key] = current_visibility_label
                managed_visibility_label = st.radio(
                    "내 기록 공개범위 변경",
                    visibility_options,
                    horizontal=True,
                    key=visibility_manage_key,
                    label_visibility="collapsed",
                )
                new_visibility = "비공개" if managed_visibility_label.startswith("🔒") else "공개"
                change_needed = new_visibility != current_visibility
                if st.button(
                    "💾 공개범위 변경 저장" if change_needed else "현재 공개범위 유지",
                    type="primary" if change_needed else "secondary",
                    use_container_width=True,
                    disabled=not change_needed,
                    key=f"worklog_visibility_save_{selected_id}",
                ):
                    with st.spinner("공개범위를 변경하고 권한을 반영하고 있습니다..."):
                        visibility_ok, visibility_message = update_work_log_visibility(
                            selected_id, new_visibility, actor_user=auth_user
                        )
                    if visibility_ok:
                        st.session_state["worklog_df"] = load_work_logs(auth_user)
                        st.session_state["worklog_selected_id"] = selected_id
                        st.success(visibility_message)
                        time.sleep(0.5)
                        st.rerun()
                    else:
                        st.error(visibility_message)
            else:
                st.caption("※ 공개범위 변경 권한은 이 기록을 작성한 본인에게만 있습니다.")

            st.markdown("### 🔄 상태이력 · 조치 업데이트")
            st.caption(f"기록ID: {selected_id} · {selected.get('모국','')} / {selected.get('국소','')}")
            u1, u2 = st.columns(2)
            with u1:
                update_writer = str(auth_user.get("name", "") or "").strip()
                st.text_input(
                    "변경 작성자",
                    value=update_writer,
                    disabled=True,
                    key="worklog_update_writer",
                    help="상태·조치 변경 이력에는 현재 로그인 사용자가 자동 기록됩니다.",
                )
                current_status = str(selected.get("상태", "신규"))
                status_index = WORK_LOG_STATUS_OPTIONS.index(current_status) if current_status in WORK_LOG_STATUS_OPTIONS else 0
                update_status = st.selectbox("변경 상태", WORK_LOG_STATUS_OPTIONS, index=status_index, key="worklog_update_status")
                update_action = st.text_area("조치내용", value=str(selected.get("조치내용", "")), height=100, key="worklog_update_action")
            with u2:
                update_followup = st.text_area("후속조치", value=str(selected.get("후속조치", "")), height=100, key="worklog_update_followup")
                update_remark = st.text_area("비고", value=str(selected.get("비고", "")), height=100, key="worklog_update_remark")

            uc1, uc2 = st.columns(2)
            with uc1:
                if st.button("💾 상태·조치 저장", type="primary", use_container_width=True, key="worklog_update_save"):
                    ok, msg = update_work_log(
                        selected_id,
                        update_writer,
                        update_status,
                        update_action,
                        update_followup,
                        update_remark,
                        actor_user=auth_user,
                    )
                    if ok:
                        st.session_state["worklog_df"] = load_work_logs(auth_user)
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
            with uc2:
                if st.button("닫기", use_container_width=True, key="worklog_update_close"):
                    st.session_state["worklog_selected_id"] = ""
                    st.session_state.pop("worklog_delete_pending_id", None)
                    for detail_key in (
                        "worklog_update_writer", "worklog_update_status", "worklog_update_action",
                        "worklog_update_followup", "worklog_update_remark",
                    ):
                        st.session_state.pop(detail_key, None)
                    st.rerun()

            if is_owner:
                delete_photo_count = len([
                    value for value in str(selected.get("사진파일ID목록", "") or "").split("|") if value.strip()
                ])
                st.markdown(
                    f'<div class="worklog-delete-box"><b>🗑️ 내 기록 삭제</b><br>'
                    f'본인이 작성한 기록만 삭제할 수 있습니다. 삭제 시 연결 사진 {delete_photo_count}장은 '
                    'Google Drive에서 영구삭제하지 않고 휴지통으로 이동합니다.</div>',
                    unsafe_allow_html=True,
                )
                pending_delete_id = str(st.session_state.get("worklog_delete_pending_id", "") or "").strip()
                if pending_delete_id != selected_id:
                    if st.button(
                        "🗑️ 이 기록 삭제",
                        key=f"worklog_delete_request_{selected_id}",
                        use_container_width=True,
                    ):
                        st.session_state["worklog_delete_pending_id"] = selected_id
                        st.rerun()
                else:
                    st.warning(
                        f"정말 삭제하시겠습니까? · {selected.get('모국','')} / {selected.get('국소','')} · "
                        f"{selected.get('저장일시','')} · 사진 {delete_photo_count}장"
                    )
                    delete_yes, delete_no = st.columns(2)
                    with delete_yes:
                        if st.button(
                            "삭제 확정",
                            type="primary",
                            use_container_width=True,
                            key=f"worklog_delete_confirm_{selected_id}",
                        ):
                            with st.spinner("내 기록과 연결 사진을 안전하게 정리하고 있습니다..."):
                                delete_ok, delete_message = delete_work_log(selected_id, auth_user)
                            if delete_ok:
                                st.session_state["worklog_df"] = load_work_logs(auth_user)
                                st.session_state["worklog_selected_id"] = ""
                                st.session_state.pop("worklog_delete_pending_id", None)
                                st.success(delete_message)
                                time.sleep(0.6)
                                st.rerun()
                            else:
                                st.error(delete_message)
                    with delete_no:
                        if st.button(
                            "취소",
                            use_container_width=True,
                            key=f"worklog_delete_cancel_{selected_id}",
                        ):
                            st.session_state.pop("worklog_delete_pending_id", None)
                            st.rerun()
            else:
                st.caption("※ 삭제 권한은 이 기록을 작성한 본인에게만 있습니다.")

            history_df = load_work_log_history(selected_id, auth_user)
            if not history_df.empty:
                st.markdown("#### 🕘 변경 이력")
                history_display = history_df[[c for c in WORK_LOG_HISTORY_HEADERS if c in history_df.columns]].copy()
                st.dataframe(history_display, use_container_width=True, hide_index=True)


        with recent_col:
            st.markdown('<div class="worklog-recent-marker"></div>', unsafe_allow_html=True)
            with st.container(border=True):
                st.markdown('<div class="worklog-section-title">🕘 최근 기록</div>', unsafe_allow_html=True)

                if not isinstance(loaded_df, pd.DataFrame):
                    st.info("검색 조건과 검색어를 정한 뒤 ‘불러오기’를 누르면 최근 현장이력이 표시됩니다.")
                else:
                    display_logs = loaded_df.copy()
                    if worklog_filter == "🌐 공개":
                        display_logs = display_logs[display_logs["공개범위"].astype(str) == "공개"]
                        if worklog_public_scope == "👤 내 기록":
                            current_user_id = str(auth_user.get("user_id", "") or "").strip()
                            display_logs = display_logs[
                                display_logs["작성자ID"].astype(str).str.strip() == current_user_id
                            ]
                    elif worklog_filter == "🔒 내 비공개":
                        display_logs = display_logs[display_logs["공개범위"].astype(str) == "비공개"]
                    elif worklog_filter in WORK_LOG_STATUS_OPTIONS:
                        display_logs = display_logs[display_logs["상태"].astype(str) == worklog_filter]
                    if worklog_search:
                        search_cols = [
                            "작성자", "권역", "모국", "국소", "상태", "점검항목",
                            "현상_특이사항", "조치내용", "후속조치", "비고",
                        ]
                        mask = display_logs[search_cols].apply(
                            lambda row: row.astype(str).str.contains(worklog_search, case=False, na=False).any(),
                            axis=1,
                        )
                        display_logs = display_logs[mask]

                    if display_logs.empty:
                        st.warning("조건에 맞는 WORK LOG가 없습니다.")
                    else:
                        for _, log in display_logs.head(12).iterrows():
                            record_id = str(log.get("기록ID", "")).strip()
                            status_value = str(log.get("상태", "신규")).strip()
                            badge_class = {
                                "신규": "new", "확인필요": "wait", "조치중": "doing",
                                "재점검": "recheck", "완료": "done",
                            }.get(status_value, "wait")
                            place_text = f"{str(log.get('모국','')).strip()} · {str(log.get('국소','')).strip()}"
                            main_note = str(log.get("현상_특이사항", "")).strip() or str(log.get("조치내용", "")).strip() or "기록 내용 없음"
                            saved_text = str(log.get("저장일시", "")).strip()
                            items_text = str(log.get("점검항목", "")).strip()
                            privacy_value = _worklog_record_visibility(log)
                            privacy_class = "private" if privacy_value == "비공개" else "public"
                            privacy_text = "🔒 나만 보기" if privacy_value == "비공개" else "🌐 팀 공유"
                            photo_ids = [v for v in str(log.get("사진파일ID목록", "")).split("|") if v.strip()]

                            st.markdown(
                                f'<div class="worklog-card">'
                                f'<div class="worklog-card-top"><div>'
                                f'<div class="worklog-place">📍 {html.escape(place_text)}'
                                f'<span class="worklog-privacy-badge {privacy_class}">{privacy_text}</span></div>'
                                f'<div class="worklog-meta">{html.escape(items_text)} · {html.escape(saved_text)} · {html.escape(str(log.get("작성자","")))}</div>'
                                f'</div><span class="worklog-badge {badge_class}">{html.escape(status_value)}</span></div>'
                                f'<div class="worklog-body">{html.escape(main_note)}</div>'
                                f'</div>',
                                unsafe_allow_html=True,
                            )

                            if photo_ids:
                                thumbnail_ids = photo_ids[:4]
                                thumb_cols = st.columns(len(thumbnail_ids))
                                for photo_index, file_id in enumerate(thumbnail_ids):
                                    with thumb_cols[photo_index]:
                                        photo_bytes = _worklog_download_drive_image(file_id)
                                        if photo_bytes:
                                            st.image(photo_bytes, use_container_width=True)
                                if len(photo_ids) > 4:
                                    st.caption(f"📷 사진 {len(photo_ids)}장 · 화면에는 처음 4장만 미리보기")

                            action_c1, action_c2 = st.columns(2)
                            with action_c1:
                                if st.button("상세·조치", key=f"worklog_detail_{record_id}", use_container_width=True):
                                    previous_detail_id = str(st.session_state.get("worklog_selected_id", "") or "").strip()
                                    if previous_detail_id != record_id:
                                        for detail_key in (
                                            "worklog_update_writer", "worklog_update_status", "worklog_update_action",
                                            "worklog_update_followup", "worklog_update_remark",
                                        ):
                                            st.session_state.pop(detail_key, None)
                                        st.session_state.pop("worklog_delete_pending_id", None)
                                    st.session_state["worklog_selected_id"] = record_id
                                    st.rerun()
                            with action_c2:
                                if status_value != "완료":
                                    if st.button("✅ 완료", key=f"worklog_done_{record_id}", use_container_width=True):
                                        ok, msg = update_work_log(
                                            record_id,
                                            str(auth_user.get("name", "") or "현장 사용자"),
                                            "완료",
                                            str(log.get("조치내용", "")),
                                            str(log.get("후속조치", "")),
                                            str(log.get("비고", "")),
                                            actor_user=auth_user,
                                        )
                                        if ok:
                                            st.session_state["worklog_df"] = load_work_logs(auth_user)
                                            st.success(msg)
                                            st.rerun()
                                        else:
                                            st.error(msg)
                                else:
                                    st.button("완료됨", disabled=True, key=f"worklog_done_disabled_{record_id}", use_container_width=True)

                            # V15: 상세·조치 편집은 선택한 카드 바로 아래에 삽입하고, 다음 기록은 그 아래에 이어집니다.
                            active_detail_id = str(st.session_state.get("worklog_selected_id", "") or "").strip()
                            if active_detail_id == record_id:
                                _render_worklog_inline_detail(log, record_id)

        if isinstance(st.session_state.get("worklog_df"), pd.DataFrame):
            st.markdown('<div class="worklog-close-safe-space"></div>', unsafe_allow_html=True)
            st.markdown('<div class="worklog-sticky-close-marker"></div>', unsafe_allow_html=True)
            st.button(
                "✕ 조회 닫기",
                key="worklog_results_close",
                use_container_width=True,
                on_click=_worklog_close_loaded_results,
            )



# --- [Tab 1: 국사 전원시설 정밀점검] ---
with tab_power:
    if st.session_state.get("power_current_theme") not in POWER_THEME_ORDER:
        st.session_state["power_current_theme"] = POWER_THEME_ORDER[0]
    if _power_get("power_phase_type", "삼상") not in {"삼상", "단상"}:
        _power_set("power_phase_type", "삼상")
    if _power_get("power_battery_set", "1조 셀 측정") not in {"1조 셀 측정", "2조 셀 측정"}:
        _power_set("power_battery_set", "1조 셀 측정")
    if "power_unlocked_theme_index" not in st.session_state:
        st.session_state["power_unlocked_theme_index"] = len(POWER_THEME_ORDER) - 1
    if "power_theme_confirmations" not in st.session_state:
        st.session_state["power_theme_confirmations"] = {}
    if "power_panel_nonce" not in st.session_state:
        st.session_state["power_panel_nonce"] = 0
    unlocked_index = _power_unlocked_theme_index()
    _power_draft()
    selected_worker_state = str(st.session_state.get("power_worker", "담당자 선택")).strip()
    expected_area = _major_area_for_worker_value(selected_worker_state)
    if expected_area in POWER_REGION_DATA:
        # 과거 세션에 개별 담당자가 남아 있어도 현장 화면에서는 권역 담당자 2명을 한 묶음으로 통일합니다.
        selected_worker_state = _automatic_inspector_display(expected_area)
        st.session_state["power_worker"] = selected_worker_state
    else:
        st.session_state["power_worker"] = "담당자 선택"
        selected_worker_state = "담당자 선택"
        expected_area = "권역 선택"
    if st.session_state.get("power_major_area", "권역 선택") != expected_area:
        st.session_state["power_major_area"] = expected_area
    expected_group = _inspector_group_for_area(expected_area) if expected_area in POWER_REGION_DATA else ""
    if st.session_state.get("power_inspector_group", "") != expected_group:
        st.session_state["power_inspector_group"] = expected_group

    st.markdown("### 🔋 국사 전원시설 정밀점검")
    st.caption("국사명을 검색해 선택하면 담당자 2명·주요 점검권역·모국·국소가 자동 입력됩니다. 기존 수동 선택 방식도 그대로 사용할 수 있습니다.")

    st.markdown("""
    <style>
    .power-mobile-hero {
        background: linear-gradient(135deg, #EAF4FF 0%, #F8FAFC 54%, #ECFDF5 100%);
        border: 1px solid #D5E3F3;
        border-left: 8px solid #0B5CAB;
        border-radius: 20px;
        padding: 17px 19px;
        margin: 8px 0 12px 0;
        box-shadow: 0 8px 24px rgba(15, 23, 42, 0.08);
    }
    .power-mobile-hero h3 { margin:0 0 7px 0; color:#0F172A; font-weight:950; font-size:clamp(1.28rem,3.5vw,1.60rem); letter-spacing:-0.02em; }
    .power-mobile-hero p { margin:0; color:#334155; line-height:1.65; font-weight:720; font-size:clamp(.94rem,2.6vw,1.04rem); }
    .power-basic-card {
        background:#FFFFFF; border:1px solid #D8E3F2; border-radius:17px;
        padding:14px 15px 6px; margin:10px 0 10px; box-shadow:0 6px 18px rgba(15,23,42,.06);
    }
    .power-basic-title {font-size:clamp(1.16rem,3.2vw,1.34rem); font-weight:950; color:#0B5CAB; margin-bottom:9px; letter-spacing:-.01em;}
    .power-sticky-card {
        position: sticky; top: 0.35rem; z-index: 990;
        background: rgba(15, 23, 42, 0.96); color: #FFFFFF;
        border: 1px solid rgba(255,255,255,0.18); border-radius: 15px;
        padding: 10px 14px; margin: 8px 0 11px 0;
        box-shadow: 0 10px 24px rgba(15, 23, 42, 0.22); backdrop-filter: blur(10px);
    }
    .power-sticky-title { font-size:.82rem; opacity:.82; font-weight:850; margin-bottom:3px; }
    .power-sticky-main { font-size:clamp(1.00rem,3vw,1.16rem); font-weight:950; line-height:1.35; word-break:keep-all; }
    .power-sticky-area { font-size:clamp(.88rem,2.7vw,1.02rem); font-weight:850; color:#BFDBFE; margin-top:2px; }
    .power-sticky-sub { font-size:.88rem; opacity:.92; margin-top:3px; line-height:1.4; }
    .power-theme-heading {
        font-size:1.18rem; font-weight:950; color:#0B5CAB;
        margin:12px 0 9px 0; padding-bottom:8px; border-bottom:2px solid #DCEAF7;
    }
    .power-unit-guide {
        background:#EFF6FF; border:1px solid #BFDBFE; border-radius:12px;
        padding:10px 12px; color:#1E3A8A; font-weight:800; margin:8px 0 12px;
    }
    .power-missing-box {
        background:#FFF7ED; border:1px solid #FDBA74; border-left:7px solid #F97316;
        border-radius:15px; padding:13px 15px; margin:10px 0;
        color:#7C2D12; line-height:1.55; font-weight:750;
    }
    .power-complete-box {
        background:#ECFDF5; border:1px solid #86EFAC; border-radius:15px;
        padding:13px 15px; color:#166534; font-weight:800;
    }
    .power-battery-flow {
        display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); gap:7px;
        margin:8px 0 11px;
    }
    .power-battery-flow-step {
        min-height:66px; display:flex; flex-direction:column; align-items:center; justify-content:center;
        padding:8px 6px; border:1.5px solid #CBD5E1; border-radius:13px;
        background:#F8FAFC; color:#64748B; text-align:center; line-height:1.25;
        font-size:clamp(.76rem,2.3vw,.91rem); font-weight:900;
    }
    .power-battery-flow-step .step-no {
        display:inline-flex; align-items:center; justify-content:center; width:25px; height:25px;
        margin-bottom:4px; border-radius:999px; background:#E2E8F0; color:#475569; font-weight:950;
    }
    .power-battery-flow-step.active {
        border-color:#F59E0B; background:linear-gradient(135deg,#FFF7ED,#FEF3C7);
        color:#92400E; box-shadow:0 0 0 2px rgba(245,158,11,.18);
    }
    .power-battery-flow-step.active .step-no { background:#F59E0B; color:#FFFFFF; }
    .power-battery-flow-step.done {
        border-color:#34D399; background:linear-gradient(135deg,#ECFDF5,#DCFCE7); color:#166534;
    }
    .power-battery-flow-step.done .step-no { background:#10B981; color:#FFFFFF; }
    .power-battery-path-guide {
        margin:7px 0 12px; padding:10px 12px; border-radius:12px;
        background:#EFF6FF; border:1px solid #93C5FD; color:#1E3A8A;
        font-size:clamp(.82rem,2.5vw,.96rem); font-weight:800; line-height:1.55;
    }
    .power-notes-guide {
        margin:4px 0 9px; padding:13px 15px; border-radius:14px;
        background:linear-gradient(135deg,#FFF7ED,#FFFBEB); border:1px solid #FDBA74;
        border-left:7px solid #F97316; color:#7C2D12;
    }
    .power-notes-guide .title { font-size:clamp(1.10rem,3.4vw,1.30rem); font-weight:950; margin-bottom:4px; }
    .power-notes-guide .desc { font-size:clamp(.88rem,2.6vw,1rem); font-weight:750; line-height:1.45; }
    div.st-key-_ui_power_notes textarea {
        min-height:145px !important; border:2px solid #F97316 !important;
        background:#FFFEF7 !important; box-shadow:0 0 0 3px rgba(249,115,22,.10) !important;
        font-size:clamp(1.04rem,3vw,1.16rem) !important; font-weight:800 !important; line-height:1.55 !important;
    }
    div.st-key-_ui_power_notes textarea::placeholder {
        color:#94A3B8 !important; -webkit-text-fill-color:#94A3B8 !important; font-weight:600 !important; opacity:1 !important;
    }
    div.st-key-_ui_power_notes label p {
        color:#9A3412 !important; font-size:clamp(1rem,3vw,1.14rem) !important; font-weight:950 !important;
    }
    .power-loaded-box {
        background:#F0FDF4; border:1px solid #86EFAC; border-left:6px solid #16A34A;
        border-radius:13px; padding:11px 13px; color:#166534; font-weight:800; margin:8px 0;
    }
    .power-sheet-note {
        background:#F8FAFC; border:1px solid #CBD5E1; border-radius:14px;
        padding:13px 15px; color:#334155; font-weight:700; line-height:1.55;
    }
    .power-ground-heading {
        border-radius:13px; padding:11px 14px; margin:9px 0 9px;
        font-size:clamp(1.05rem,3vw,1.22rem); font-weight:950; color:#0F172A;
        box-shadow:0 5px 14px rgba(15,23,42,.07);
    }
    .power-ground-heading span { font-weight:850; opacity:.78; }
    .power-ground-heading.security {
        background:linear-gradient(135deg,#DBEAFE,#E0F2FE);
        border:1px solid #60A5FA; border-left:7px solid #2563EB;
    }
    .power-ground-heading.telecom {
        background:linear-gradient(135deg,#DCFCE7,#ECFDF5);
        border:1px solid #4ADE80; border-left:7px solid #16A34A;
    }
    .power-ground-heading.lightning {
        background:linear-gradient(135deg,#FEF3C7,#FFF7ED);
        border:1px solid #FBBF24; border-left:7px solid #F97316;
    }
    .power-menu-legend {
        display:flex; flex-wrap:wrap; gap:7px; margin:7px 0 10px;
        font-size:.80rem; font-weight:850; color:#334155;
    }
    .power-menu-legend span {
        background:#FFFFFF; border:1px solid #D8E3F2; border-radius:999px;
        padding:5px 9px; box-shadow:0 3px 9px rgba(15,23,42,.05);
    }
    @keyframes powerBlindDown {
        0% { opacity:0; transform:scaleY(0.08) translateY(-10px); clip-path:inset(0 0 92% 0); }
        70% { opacity:1; transform:scaleY(1.015) translateY(0); clip-path:inset(0 0 0 0); }
        100% { opacity:1; transform:scaleY(1) translateY(0); clip-path:inset(0 0 0 0); }
    }
    details:has(.power-panel-marker) [data-testid="stExpanderDetails"] {
        transform-origin:top center;
        animation:powerBlindDown .42s cubic-bezier(.2,.82,.25,1) both;
    }
    details:has(.power-panel-marker) > summary {
        background:linear-gradient(135deg,#EFF6FF,#F8FAFC) !important;
        border-radius:14px !important;
        border:1px solid #BFDBFE !important;
        padding:0.65rem 0.85rem !important;
    }
    .power-menu-status {
        background:#F8FAFC; border:1px solid #D8E3F2; border-radius:12px;
        padding:9px 12px; margin:8px 0 10px; color:#475569; font-weight:750;
    }
    div[class*="st-key-power_theme_menu_"] button {
        min-height: 56px !important; padding: 0.50rem 0.35rem !important;
        border-radius: 14px !important; font-size: 0.94rem !important; line-height: 1.15 !important;
        font-weight:950 !important; transition:transform .18s ease, box-shadow .18s ease !important;
    }
    div[class*="st-key-power_theme_menu_"] button:hover:not(:disabled) {
        transform:translateY(-2px) !important;
    }
    div[class*="st-key-_ui_power_battery_"] input {
        text-align:center !important; padding-left:0.25rem !important; padding-right:0.25rem !important;
        min-height:44px !important; font-weight:850 !important;
    }
    div[class*="st-key-_ui_power_battery_"] label p {
        text-align:center !important; font-size:0.86rem !important; font-weight:950 !important;
    }
    .power-measurement-menu-title {
        font-size:clamp(1.26rem,3.7vw,1.52rem); font-weight:950; color:#0F3B66;
        margin:14px 0 9px; padding:10px 13px; border-radius:14px;
        background:linear-gradient(135deg,#DBEAFE,#ECFEFF); border:1px solid #93C5FD;
        border-left:7px solid #0284C7; box-shadow:0 5px 14px rgba(2,132,199,.10);
    }
    details:has(.power-panel-marker) > summary p,
    details:has(.power-panel-marker) > summary span {
        font-size:clamp(1.10rem,3.4vw,1.30rem) !important; font-weight:950 !important; color:#0F3B66 !important;
    }
    div[class*="st-key-_ui_power_"] label p {
        font-size:clamp(.96rem,2.8vw,1.08rem) !important; font-weight:950 !important; color:#172033 !important; line-height:1.35 !important;
    }
    div[class*="st-key-_ui_power_"] input,
    div[class*="st-key-_ui_power_"] textarea {
        font-size:clamp(1.04rem,3vw,1.18rem) !important; font-weight:900 !important; color:#0F172A !important;
        -webkit-text-fill-color:#0F172A !important; min-height:48px !important; border:1.5px solid #94A3B8 !important;
    }
    div.st-key-power_worker label p, div.st-key-power_mother label p, div.st-key-power_local label p,
    div.st-key-power_station_search_query label p {
        font-size:clamp(.98rem,2.9vw,1.10rem) !important; font-weight:950 !important; color:#1E293B !important;
    }
    div.st-key-power_worker [data-baseweb="select"], div.st-key-power_mother [data-baseweb="select"], div.st-key-power_local [data-baseweb="select"] {
        width:100% !important; height:52px !important; min-height:52px !important; max-height:52px !important;
        box-sizing:border-box !important; border:2px solid #38BDF8 !important; border-radius:12px !important;
        background:#F0F9FF !important; box-shadow:0 5px 14px rgba(14,165,233,.12) !important;
        font-size:clamp(1rem,3vw,1.12rem) !important; font-weight:850 !important; opacity:1 !important;
    }
    div.st-key-power_worker [data-baseweb="select"] > div,
    div.st-key-power_mother [data-baseweb="select"] > div,
    div.st-key-power_local [data-baseweb="select"] > div {
        height:48px !important; min-height:48px !important; border:0 !important; border-radius:10px !important;
        background:transparent !important; box-shadow:none !important;
    }
    div.st-key-power_station_search_query input {
        min-height:54px !important; border:2.5px solid #7C3AED !important; border-radius:13px !important;
        background:#FFFFFF !important; box-shadow:0 6px 16px rgba(124,58,237,.16) !important;
        font-size:clamp(1.03rem,3vw,1.16rem) !important; font-weight:950 !important; color:#4C1D95 !important;
        -webkit-text-fill-color:#4C1D95 !important;
    }
    .power-station-search-guide {
        margin:2px 0 10px; padding:13px 14px; border-radius:15px;
        background:linear-gradient(135deg,#FFF7F7 0%,#FAF5FF 48%,#EFF6FF 100%);
        border:2px solid #7C3AED; border-left:7px solid #D71920;
        box-shadow:0 8px 20px rgba(76,29,149,.13); color:#3B0764; line-height:1.5;
    }
    .power-station-search-guide .title { font-size:clamp(1.08rem,3.2vw,1.26rem); font-weight:950; color:#6D28D9; }
    .power-station-search-guide .desc { margin-top:3px; font-size:.88rem; font-weight:800; color:#4B5563; }
    .power-station-duplicate-guide {
        margin:8px 0 7px; padding:9px 11px; border-radius:11px; background:#FFF7ED;
        border:1px solid #FDBA74; color:#9A3412; font-size:.88rem; font-weight:900; line-height:1.45;
    }
    .power-station-applied-note {
        margin:8px 0 10px; padding:10px 12px; border-radius:12px;
        background:linear-gradient(135deg,#F5F3FF,#FFF1F2); border:1.5px solid #A78BFA;
        color:#6D28D9; font-weight:950; line-height:1.5; box-shadow:0 4px 12px rgba(109,40,217,.09);
    }
    div.st-key-power_station_search_form button,
    div.st-key-power_station_duplicate_form button {
        min-height:54px !important; border-radius:13px !important; border:0 !important;
        background:linear-gradient(135deg,#D71920 0%,#7C3AED 100%) !important;
        color:#FFFFFF !important; font-weight:950 !important; box-shadow:0 7px 17px rgba(124,58,237,.20) !important;
    }
    div.st-key-power_station_search_form button *,
    div.st-key-power_station_duplicate_form button * { color:#FFFFFF !important; -webkit-text-fill-color:#FFFFFF !important; }
    div.st-key-power_station_search_choice [role="radiogroup"] {
        padding:4px 2px 2px;
    }
    div.st-key-power_station_search_choice label {
        margin:4px 0 !important; padding:8px 10px !important; border-radius:10px !important;
        background:#FAF5FF !important; border:1px solid #DDD6FE !important;
    }
    .power-basic-auto-label {
        font-size:clamp(.98rem,2.9vw,1.10rem); font-weight:950; color:#1E293B; margin:0 0 6px 1px;
    }
    .power-basic-auto-card {
        width:100%; height:52px; min-height:52px; max-height:52px;
        display:flex; align-items:center; justify-content:center; box-sizing:border-box;
        padding:0 12px; border:2px solid #38BDF8; border-radius:12px;
        background:linear-gradient(135deg,#EFF6FF 0%,#F0F9FF 55%,#ECFEFF 100%);
        color:#0F3C5D; font-size:clamp(1rem,3vw,1.14rem); font-weight:950;
        letter-spacing:-0.02em; text-align:center; box-shadow:0 5px 14px rgba(14,165,233,.12);
        overflow:hidden;
    }
    .power-basic-auto-card.is-empty { color:#64748B; border-color:#CBD5E1; background:#F8FAFC; box-shadow:none; }
    .power-basic-auto-card.search-applied {
        color:#7C3AED; border-color:#8B5CF6; background:linear-gradient(135deg,#FAF5FF,#FFF1F2);
        box-shadow:0 5px 14px rgba(124,58,237,.14);
    }
    .power-basic-history-title {
        margin:14px 0 8px; padding-top:12px; border-top:1px solid #D8E3F2;
        color:#0B5CAB; font-size:clamp(1.02rem,3vw,1.16rem); font-weight:950;
    }
    .power-auto-worker-label { font-size:clamp(.98rem,2.9vw,1.10rem); font-weight:950; color:#1E293B; margin:0 0 6px 1px; }
    .power-auto-worker-card {
        min-height:48px; display:flex; align-items:center; justify-content:center;
        padding:9px 12px; border:2px solid #38BDF8; border-radius:12px;
        background:linear-gradient(135deg,#E0F2FE 0%,#F0F9FF 55%,#ECFEFF 100%);
        color:#0F3C5D; font-size:clamp(1.02rem,3.2vw,1.18rem); font-weight:950;
        letter-spacing:-0.02em; text-align:center; box-shadow:0 5px 14px rgba(14,165,233,.13);
    }
    .power-auto-worker-card.is-empty { color:#64748B; border-color:#CBD5E1; background:#F8FAFC; }
    div[class*="st-key-power_theme_menu_"] button p { font-size:clamp(.95rem,2.8vw,1.08rem) !important; font-weight:950 !important; }
    @media (max-width: 768px) {
        .power-mobile-hero { padding:14px 13px; border-radius:15px; }
        .power-sticky-card { top:0.2rem; border-radius:13px; padding:9px 11px; }
        .power-sticky-main { font-size:1rem; }
        div[class*="st-key-_ui_power_"] input,
        div[class*="st-key-_ui_power_"] textarea { font-size:16px !important; min-height:47px !important; }
        div[class*="st-key-_ui_power_"] label p { font-size:.94rem !important; }
        div[class*="st-key-_ui_power_battery_"] input { padding:.35rem .08rem !important; text-align:center !important; }
        div[class*="st-key-_ui_power_battery_"] label p { font-size:.78rem !important; }
        div[class*="st-key-power_theme_menu_"] button { min-height:60px !important; }
        .power-battery-flow { grid-template-columns:repeat(2,minmax(0,1fr)); gap:6px; }
        .power-battery-flow-step { min-height:60px; }
    }
    @media (max-width: 430px) {
        .power-mobile-hero { padding:13px 11px; }
        .power-basic-card { padding:12px 11px 5px; }
        .power-sticky-card { margin-left:-.1rem; margin-right:-.1rem; }
        .power-menu-legend { gap:5px; font-size:.75rem; }
        .power-menu-legend span { padding:4px 7px; }
        div[class*="st-key-_ui_power_battery_"] label p { font-size:.72rem !important; }
        div[class*="st-key-_ui_power_battery_"] input { font-size:15px !important; min-height:44px !important; }
    }
    </style>
    <div class="power-mobile-hero">
      <h3>새로운 전원 정밀점검 전용 공간</h3>
      <p>국사명을 검색해 선택하면 담당자 2명·주요 점검권역·모국·국소가 자동 입력됩니다. 기존 수동 선택과 과거 측정값 불러오기도 그대로 사용할 수 있습니다.</p>
    </div>
    """, unsafe_allow_html=True)

    # 기본정보와 과거 측정값 불러오기를 하나의 블록으로 묶어 표시합니다.
    with st.container(border=True):
        st.markdown('<div class="power-basic-title">👤 기본정보</div>', unsafe_allow_html=True)

        # 국사 검색은 기존 기본정보보다 먼저, 별도 카드처럼 눈에 띄게 배치합니다.
        with st.container(border=True):
            st.markdown(
                '<div class="power-station-search-guide">'
                '<div class="title">📍 국사 검색 · 빠른 자동입력</div>'
                '<div class="desc">예: <b>송포</b> 입력 → 오른쪽 <b>확인</b>. 한 곳이면 즉시 자동입력되고, 같은 이름이 여러 곳이면 아래 후보 중 하나만 선택하면 됩니다.</div>'
                '</div>',
                unsafe_allow_html=True,
            )

            with st.form(key="power_station_search_form", clear_on_submit=False):
                search_input_col, search_button_col = st.columns([4.2, 1.15], gap="small")
                with search_input_col:
                    st.text_input(
                        "국사명",
                        key="power_station_search_query",
                        placeholder="예: 송포",
                        help="국사명을 입력한 뒤 오른쪽 확인 버튼을 누르세요. 키보드 Enter로도 확인할 수 있습니다.",
                    )
                with search_button_col:
                    st.markdown("<div style='height:1.70rem'></div>", unsafe_allow_html=True)
                    search_submitted = st.form_submit_button("확인", use_container_width=True)

            if search_submitted:
                _run_power_station_search()

            search_status = str(st.session_state.get("power_station_search_status", "") or "")
            candidate_ids = list(st.session_state.get("power_station_search_candidates", []) or [])

            if search_status == "empty":
                st.warning("국사명을 먼저 입력해 주세요. 예: 송포")
            elif search_status == "none":
                st.warning("일치하는 국사를 찾지 못했습니다. 국사명을 다시 확인해 주세요.")
            elif search_status == "multiple" and candidate_ids:
                st.markdown(
                    f'<div class="power-station-duplicate-guide">⚠️ 같은 이름 또는 유사한 국사가 {len(candidate_ids)}곳 있습니다. 아래에서 정확한 국사를 선택한 뒤 <b>선택 확인</b>을 눌러 주세요.</div>',
                    unsafe_allow_html=True,
                )
                with st.form(key="power_station_duplicate_form", clear_on_submit=False):
                    st.radio(
                        "국사 선택",
                        candidate_ids,
                        key="power_station_search_choice",
                        format_func=_power_station_search_label,
                    )
                    duplicate_submitted = st.form_submit_button("선택 확인", use_container_width=True)
                if duplicate_submitted:
                    _confirm_power_station_search_choice()

            station_search_notice = str(st.session_state.get("power_station_search_notice", "") or "")
            if station_search_notice:
                st.markdown(
                    f'<div class="power-station-applied-note">{html.escape(station_search_notice)}</div>',
                    unsafe_allow_html=True,
                )

        # 검색으로 자동 반영된 경우 아래 기본정보의 선택값을 자주색으로 강조합니다.
        if bool(st.session_state.get("power_station_search_applied", False)):
            st.markdown(
                """
                <style>
                div.st-key-power_worker [data-baseweb="select"],
                div.st-key-power_mother [data-baseweb="select"],
                div.st-key-power_local [data-baseweb="select"] {
                    border-color:#8B5CF6 !important; background:#FAF5FF !important;
                    box-shadow:0 5px 14px rgba(124,58,237,.14) !important;
                }
                div.st-key-power_worker [role="combobox"] *,
                div.st-key-power_mother [role="combobox"] *,
                div.st-key-power_local [role="combobox"] * {
                    color:#7C3AED !important; -webkit-text-fill-color:#7C3AED !important; font-weight:950 !important;
                }
                </style>
                """,
                unsafe_allow_html=True,
            )

        worker_options = ["담당자 선택"] + POWER_INSPECTOR_GROUP_OPTIONS
        if st.session_state.get("power_worker", "담당자 선택") not in worker_options:
            legacy_area = _major_area_for_worker_value(st.session_state.get("power_worker", ""))
            st.session_state["power_worker"] = (
                _automatic_inspector_display(legacy_area)
                if legacy_area in POWER_REGION_DATA else "담당자 선택"
            )

        basic_row1 = st.columns(2, gap="small")
        with basic_row1[0]:
            selected_worker = st.selectbox(
                "담당자 *",
                worker_options,
                key="power_worker",
                on_change=_on_power_worker_change,
            )

        selected_area = _major_area_for_worker_value(selected_worker)
        if st.session_state.get("power_major_area", "권역 선택") != selected_area:
            st.session_state["power_major_area"] = selected_area
        with basic_row1[1]:
            if selected_area in POWER_REGION_DATA:
                area_class = "power-basic-auto-card search-applied" if st.session_state.get("power_station_search_applied", False) else "power-basic-auto-card"
            else:
                area_class = "power-basic-auto-card is-empty"
            area_text = html.escape(selected_area if selected_area in POWER_REGION_DATA else "담당자를 선택하면 자동 표시됩니다")
            st.markdown(
                f'<div class="power-basic-auto-label">주요 점검권역 *</div>'
                f'<div class="{area_class}">{area_text}</div>',
                unsafe_allow_html=True,
            )

        area_station_map = _power_area_station_map(selected_area)
        basic_row2 = st.columns(2, gap="small")
        mother_options = ["모국 선택"] + list(area_station_map.keys())
        if st.session_state.get("power_mother", "모국 선택") not in mother_options:
            st.session_state["power_mother"] = "모국 선택"
        with basic_row2[0]:
            selected_mother = st.selectbox(
                "모국 *",
                mother_options,
                key="power_mother",
                disabled=selected_area not in POWER_REGION_DATA,
                on_change=_on_power_mother_change,
            )

        local_options = ["국소 선택"]
        if selected_mother in area_station_map:
            local_options += area_station_map[selected_mother]
        if st.session_state.get("power_local", "국소 선택") not in local_options:
            st.session_state["power_local"] = "국소 선택"
        with basic_row2[1]:
            st.selectbox(
                "국소 *",
                local_options,
                key="power_local",
                disabled=selected_mother not in area_station_map,
                on_change=_clear_power_measurements_after_station_change,
            )

        selected_local = st.session_state.get("power_local", "국소 선택")
        can_load = selected_mother in area_station_map and selected_local in area_station_map.get(selected_mother, [])

        st.markdown('<div class="power-basic-history-title">↩️ 과거 측정값 불러오기</div>', unsafe_allow_html=True)
        history_periods = {
            "최근 6개월": 183,
            "최근 1년": 365,
            "최근 2년": 730,
        }
        history_col1, history_col2 = st.columns([0.82, 1.18], gap="small")
        with history_col1:
            history_period_label = st.selectbox(
                "조회 기간",
                list(history_periods.keys()),
                key="power_history_period_label",
                disabled=not can_load,
            )
        with history_col2:
            history_search_clicked = st.button(
                "🔎 과거 측정기록 조회",
                key="power_search_history",
                use_container_width=True,
                disabled=not can_load,
            )

        if history_search_clicked:
            within_days = history_periods.get(history_period_label, 183)
            with st.spinner("동일 국소의 과거 측정기록을 확인하고 있습니다..."):
                ok, message, records = list_power_inspection_history(
                    selected_mother,
                    selected_local,
                    within_days=within_days,
                    max_records=100,
                )
            if ok:
                st.session_state["power_history_records"] = records
                st.session_state["power_history_message"] = message
                st.session_state["power_history_station"] = f"{selected_mother}|{selected_local}"
                st.session_state["power_history_selected_index"] = 0
                st.rerun()
            else:
                st.session_state.pop("power_history_records", None)
                st.warning(message)

        history_records = list(st.session_state.get("power_history_records", []))
        expected_history_station = f"{selected_mother}|{selected_local}" if can_load else ""
        if st.session_state.get("power_history_station", "") != expected_history_station:
            history_records = []
            st.session_state.pop("power_history_records", None)

        if history_records:
            st.success(st.session_state.get("power_history_message", f"과거 측정기록 {len(history_records)}건을 조회했습니다."))

            def _power_history_label(index: int) -> str:
                record = history_records[index]
                saved_at = str(record.get("저장일시", "일시 미상")).strip() or "일시 미상"
                worker = str(record.get("점검자", "점검자 미상")).strip() or "점검자 미상"
                phase = str(record.get("전원구분", "-")).strip() or "-"
                completion = str(record.get("입력완료율(%)", "-")).strip() or "-"
                return f"{saved_at} · {worker} · {phase} · 완료율 {completion}%"

            selected_history_index = st.selectbox(
                "불러올 측정기록",
                options=list(range(len(history_records))),
                format_func=_power_history_label,
                key="power_history_selected_index",
            )
            selected_history_record = history_records[int(selected_history_index)]
            selected_history_photo_count = len([
                value for value in str(selected_history_record.get("사진파일ID목록", "") or "").split("|") if value.strip()
            ])
            if selected_history_photo_count:
                with st.expander(f"📷 선택 기록 현장사진 {selected_history_photo_count}장 · 보기/다운로드", expanded=False):
                    _render_power_photo_download(
                        selected_history_record,
                        key_prefix=f"power_history_{selected_history_record.get('점검ID', selected_history_index)}",
                    )
            if st.button(
                "↩️ 선택한 측정값 불러오기",
                key="power_load_selected_history",
                use_container_width=True,
                type="primary",
            ):
                selected_record = history_records[int(selected_history_index)]
                _set_power_state_from_record(selected_record)
                st.session_state["power_loaded_message"] = (
                    f"선택한 측정값을 불러왔습니다. "
                    f"({selected_record.get('저장일시', '')})"
                )
                st.rerun()

        if st.session_state.pop("power_loaded_notice", False):
            st.toast("선택한 과거 측정값을 입력폼에 불러왔습니다. 변경된 값만 수정해 주세요.", icon="↩️")
        if st.session_state.get("power_loaded_source_id"):
            st.markdown(
                f'<div class="power-loaded-box">↩️ 기존 측정값 사용 중 · '
                f'원본 저장일시: {html.escape(str(st.session_state.get("power_loaded_source_saved_at", "-")))} · '
                f'원본 점검ID: {html.escape(str(st.session_state.get("power_loaded_source_id", "-")))}</div>',
                unsafe_allow_html=True,
            )

    if st.session_state.pop("power_basic_changed_notice", False):
        st.info("기본정보가 변경되었습니다. 기존 측정값은 삭제하지 않고 그대로 유지했습니다. 현재 국소의 측정값인지 확인해 주세요.")

    basic_missing = _power_basic_missing()
    if basic_missing:
        st.caption("※ 담당자를 선택하면 주요 점검권역이 자동 표시됩니다. 모국·국소까지 선택해야 최종 전송할 수 있습니다.")

    worker_value = str(st.session_state.get("power_worker", "")).strip()
    major_area_value = str(st.session_state.get("power_major_area", "권역 선택")).strip()
    worker_summary = html.escape(worker_value if _power_worker_matches_area(worker_value, major_area_value) else "담당자 미선택")
    major_area_summary = html.escape(major_area_value or "권역 미선택")
    draft_saved_at = html.escape(str(st.session_state.get("power_draft_saved_at", "")).strip() or "-")
    mother_summary = str(st.session_state.get("power_mother", "모국 선택"))
    local_summary = str(st.session_state.get("power_local", "국소 선택"))
    station_summary = "미선택"
    if mother_summary in POWER_STATION_MAP and local_summary in POWER_STATION_MAP.get(mother_summary, []):
        station_summary = f"{mother_summary} / {local_summary}"
    current_theme = st.session_state.get("power_current_theme", POWER_THEME_ORDER[0])
    st.markdown(
        f"""<div class="power-sticky-card">
          <div class="power-sticky-title">현재 점검정보 · {html.escape(current_theme)}</div>
          <div class="power-sticky-main">👤 {worker_summary}</div>
          <div class="power-sticky-area">🗺️ {major_area_summary}</div>
          <div class="power-sticky-sub">📍 {html.escape(station_summary)} · 임시저장 {draft_saved_at}</div>
        </div>""",
        unsafe_allow_html=True,
    )

    st.markdown('<div class="power-measurement-menu-title">📋 측정 메뉴</div>', unsafe_allow_html=True)
    confirmations = dict(st.session_state.get("power_theme_confirmations", {}))
    current_theme = st.session_state.get("power_current_theme", POWER_THEME_ORDER[0])

    # 측정 순서는 자유이며, 색상으로 현재·완료·미측정 상태를 구분합니다.
    state_css_rules: list[str] = []
    for theme_index, theme in enumerate(POWER_THEME_ORDER):
        selector = f'div[class*="st-key-power_theme_menu_{theme_index}"] button'
        is_completed = theme in confirmations
        is_current = theme == current_theme
        if is_current:
            style = (
                "background:linear-gradient(135deg,#FFB703,#FB8500)!important;"
                "color:#FFFFFF!important;border:2px solid #FFD166!important;"
                "box-shadow:0 10px 24px rgba(251,133,0,.38)!important;"
            )
        elif is_completed:
            style = (
                "background:linear-gradient(135deg,#34D399,#16A34A)!important;"
                "color:#FFFFFF!important;border:2px solid #86EFAC!important;"
                "box-shadow:0 8px 20px rgba(22,163,74,.28)!important;"
            )
        else:
            style = (
                "background:linear-gradient(135deg,#38BDF8,#2563EB)!important;"
                "color:#FFFFFF!important;border:2px solid #93C5FD!important;"
                "box-shadow:0 8px 20px rgba(37,99,235,.25)!important;"
            )
        state_css_rules.append(f"{selector}{{{style}}}{selector} *{{color:inherit!important;}}")

    st.markdown("<style>" + "".join(state_css_rules) + "</style>", unsafe_allow_html=True)
    st.markdown(
        '<div class="power-menu-legend">'
        '<span>🟠 현재 측정</span><span>🟢 확인·임시저장 완료</span>'
        '<span>🔵 미측정 또는 이동 가능</span>'
        '</div>',
        unsafe_allow_html=True,
    )

    for row_start in range(0, len(POWER_THEME_ORDER), 2):
        menu_columns = st.columns(2, gap="small")
        for offset, column in enumerate(menu_columns):
            theme_index = row_start + offset
            if theme_index >= len(POWER_THEME_ORDER):
                continue
            theme = POWER_THEME_ORDER[theme_index]
            is_completed = theme in confirmations
            is_current = theme == current_theme
            if is_completed:
                status_mark = " ✅"
            elif _power_theme_started(theme):
                status_mark = " ◐"
            else:
                status_mark = ""
            prefix = "▼ " if is_current else ""
            with column:
                st.button(
                    f"{prefix}{POWER_THEME_ICON[theme]} {theme}{status_mark}",
                    key=f"power_theme_menu_{theme_index}",
                    type="primary" if is_current else "secondary",
                    use_container_width=True,
                    on_click=_activate_power_theme,
                    args=(theme,),
                )

    pending_target = st.session_state.get("power_pending_theme_switch")
    pending_from = st.session_state.get("power_pending_from_theme")
    if pending_target in POWER_THEME_ORDER and pending_from in POWER_THEME_ORDER[:-1]:
        detected_missing = _power_theme_missing(pending_from)
        missing_preview = ", ".join(detected_missing[:8]) if detected_missing else "시스템상 확인된 공란 없음"
        more_text = f" 외 {len(detected_missing) - 8}개" if len(detected_missing) > 8 else ""
        st.markdown(
            f'<div class="power-missing-box"><b>현재 입력값을 임시저장하고 이동하시겠습니까?</b><br>'
            f'현재 메뉴: {html.escape(pending_from)}<br>'
            f'이동할 메뉴: {html.escape(pending_target)}<br>'
            f'시스템 확인 공란: {html.escape(missing_preview)}{more_text}<br><br>'
            '이 이동은 측정 완료 처리와 별개입니다. 완료 판단은 각 테마 하단의 ‘측정 완료’ 버튼으로 확정합니다.</div>',
            unsafe_allow_html=True,
        )
        if st.session_state.get("power_navigation_error"):
            st.error(st.session_state["power_navigation_error"])
        move_col, stay_col = st.columns(2, gap="small")
        with move_col:
            st.button(
                "현재값 저장·이동",
                key="power_confirm_theme_switch",
                type="primary",
                use_container_width=True,
                on_click=_confirm_power_theme_switch,
            )
        with stay_col:
            st.button(
                "계속 입력",
                key="power_cancel_theme_switch",
                use_container_width=True,
                on_click=_cancel_power_theme_switch,
            )

    current_missing = _power_theme_missing(current_theme) if current_theme != "최종 확인·전송" else []
    current_status_text = (
        f"현재 ‘{current_theme}’ 입력 중 · 시스템 확인 공란 {len(current_missing)}개 · 모든 메뉴는 자유롭게 이동할 수 있습니다. 테마를 마치면 하단의 ‘측정 완료’를 눌러 주세요."
        if current_theme != "최종 확인·전송"
        else "모든 측정단계를 마쳤습니다. 최종 내용을 확인한 뒤 전송해 주세요."
    )
    st.markdown(f'<div class="power-menu-status">{html.escape(current_status_text)}</div>', unsafe_allow_html=True)

    if st.session_state.pop("power_temp_saved_notice", False):
        st.toast("현재 측정값을 임시 저장했습니다.", icon="✅")

    current_theme = st.session_state.get("power_current_theme", POWER_THEME_ORDER[0])
    _hydrate_power_theme_from_draft(current_theme)
    panel_nonce = int(st.session_state.get("power_panel_nonce", 0) or 0)

    # 측정 입력칸이 렌더링되기 전에 모바일 숫자키패드/자동 소수점 감시기를 먼저 설치합니다.
    # 특히 1조→2조 화면 전환 직후 새로 생성되는 셀 입력에도 동일 규칙이 즉시 적용됩니다.
    _render_power_auto_decimal_script()

    with st.expander(f"{POWER_THEME_ICON[current_theme]} {current_theme}", expanded=True):
        st.markdown(f'<div class="power-panel-marker" data-panel="{panel_nonce}"></div>', unsafe_allow_html=True)
        if current_theme == "전압·전류 측정":
            _hydrate_power_widget("power_phase_type", "삼상")
            st.radio(
                "전원 방식 선택",
                ["삼상", "단상"],
                format_func=lambda value: "삼상 전류" if value == "삼상" else "단상 전류",
                horizontal=True,
                key=_power_widget_key("power_phase_type"),
                on_change=_on_power_phase_change,
            )
            phase_type = _power_get("power_phase_type", "삼상")
            if phase_type == "삼상":
                voltage_fields = [
                    ("R-S 전압 (V)", "power_three_voltage_rs"),
                    ("S-T 전압 (V)", "power_three_voltage_st"),
                    ("T-R 전압 (V)", "power_three_voltage_tr"),
                    ("R-N 전압 (V)", "power_three_voltage_rn"),
                ]
                for start in range(0, len(voltage_fields), 2):
                    row1 = st.columns(2, gap="small")
                    for column, (label, key) in zip(row1, voltage_fields[start:start + 2]):
                        with column:
                            _power_text_input(label, key=key)
                current_fields = [
                    ("R상 전류 (A)", "power_three_current_r"),
                    ("S상 전류 (A)", "power_three_current_s"),
                    ("T상 전류 (A)", "power_three_current_t"),
                    ("N상 전류 (A)", "power_three_current_n"),
                ]
                for start in range(0, len(current_fields), 2):
                    current_row = st.columns(2, gap="small")
                    for column, (label, key) in zip(current_row, current_fields[start:start + 2]):
                        with column:
                            _power_text_input(label, key=key)
            else:
                row = st.columns(2, gap="small")
                with row[0]:
                    _power_text_input("단상 전압 (V)", key="power_single_voltage")
                with row[1]:
                    _power_text_input("단상 전류 (A)", key="power_single_current")

        elif current_theme == "축전지 측정":
            _hydrate_power_widget("power_battery_set", "1조 셀 측정")
            battery_exit_stage = st.session_state.get("power_battery_exit_stage")
            battery2_enabled = _power_battery2_enabled()
            current_battery_set = _power_get("power_battery_set", "1조 셀 측정")
            selected_group = 2 if current_battery_set == "2조 셀 측정" and battery2_enabled else 1

            battery_options = ["1조 셀 측정"]
            if battery2_enabled:
                battery_options.append("2조 셀 측정")
            battery_ui_key = _power_widget_key("power_battery_set")
            if st.session_state.get(battery_ui_key, "1조 셀 측정") not in battery_options:
                st.session_state[battery_ui_key] = "1조 셀 측정"
                _power_set("power_battery_set", "1조 셀 측정")

            st.radio(
                "현재 입력할 축전지",
                battery_options,
                horizontal=True,
                key=battery_ui_key,
                on_change=_on_power_battery_set_change,
            )
            selected_group = 1 if _power_get("power_battery_set", "1조 셀 측정") == "1조 셀 측정" else 2
            group_caption = (
                "1조의 실제 설치 셀 수만 입력한 뒤 아래 ‘1조 셀 측정 완료’를 눌러 주세요."
                if selected_group == 1
                else "2조의 실제 설치 셀 수만 입력한 뒤 아래 ‘2조 셀 측정 완료’를 눌러 주세요."
            )
            st.caption(group_caption)

            # 측정값 화면을 먼저 보여 준 뒤 진행상태와 2조 측정 여부를 아래에 배치합니다.
            _render_power_battery_summary(selected_group)

            if battery_exit_stage == "ask_group2_complete":
                step_classes = ["done", "done", "active", ""]
            elif selected_group == 2:
                step_classes = ["done", "done", "done", "active"]
            else:
                step_classes = ["active", "", "", ""]

            step_labels = [
                "1조 셀 입력",
                "1조 측정 완료",
                "2조 측정 여부",
                "2조 측정 또는 테마 완료",
            ]
            step_html = ''.join(
                f'<div class="power-battery-flow-step {step_classes[index]}">'
                f'<span class="step-no">{index + 1}</span>{html.escape(label)}</div>'
                for index, label in enumerate(step_labels)
            )
            st.markdown(f'<div class="power-battery-flow">{step_html}</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="power-battery-path-guide">'
                '<b>1조만 측정:</b> 1조 입력 → 1조 측정 완료 → 2조 ‘아니오’ → 다음 메뉴<br>'
                '<b>2조까지 측정:</b> 1조 입력 → 1조 측정 완료 → 2조 ‘예’ → 2조 입력 → 2조 측정 완료 → 다음 메뉴'
                '</div>',
                unsafe_allow_html=True,
            )

            if battery_exit_stage == "ask_group2_complete":
                st.markdown(
                    '<div class="power-missing-box"><b>1조 셀 입력값을 완료했습니다. 2조 셀도 측정하시겠습니까?</b><br>'
                    '예를 선택하면 1조 값은 그대로 저장되고 2조 셀 입력 화면이 열립니다.<br>'
                    '아니오를 선택하면 축전지 측정이 완료되고 다음 미완료 측정 메뉴로 이동합니다.</div>',
                    unsafe_allow_html=True,
                )
                group2_yes, group2_no = st.columns(2, gap="small")
                with group2_yes:
                    st.button(
                        "예 · 2조 셀 측정",
                        key="power_complete_measure_group2",
                        type="primary",
                        use_container_width=True,
                        on_click=_finish_battery_completion,
                        args=(True,),
                    )
                with group2_no:
                    st.button(
                        "아니오 · 축전지 측정 완료 후 다음 메뉴",
                        key="power_complete_skip_group2",
                        use_container_width=True,
                        on_click=_finish_battery_completion,
                        args=(False,),
                    )

        elif current_theme == "접지저항 측정":
            st.markdown(
                '<div class="power-ground-heading security">🛡️ 보안접지 <span>(Ω)</span></div>',
                unsafe_allow_html=True,
            )
            st.caption("보안접지는 보안 1종·2종·3종으로 구분합니다.")
            security_columns = st.columns(3, gap="small")
            with security_columns[0]:
                _power_text_input("보안 1종", key="power_security_ground_1")
            with security_columns[1]:
                _power_text_input("보안 2종", key="power_security_ground_2")
            with security_columns[2]:
                _power_text_input("보안 3종", key="power_security_ground_3")

            ground_columns = st.columns(2, gap="small")
            with ground_columns[0]:
                st.markdown(
                    '<div class="power-ground-heading telecom">📡 통신접지 <span>(Ω)</span></div>',
                    unsafe_allow_html=True,
                )
                _power_text_input("통신접지(메인)", key="power_telecom_ground")
            with ground_columns[1]:
                st.markdown(
                    '<div class="power-ground-heading lightning">⚡ 피뢰침접지 <span>(Ω)</span></div>',
                    unsafe_allow_html=True,
                )
                _power_text_input("피뢰침접지", key="power_lightning_ground")

        elif current_theme == "최종 확인·전송":
            st.markdown(
                '<div class="power-notes-guide">'
                '<div class="title">📝 특이사항 입력</div>'
                '<div class="desc">현장 특이사항이나 미측정 사유, 후속 조치가 필요한 경우 아래 메모 입력란에 작성해 주세요. 특이사항이 없으면 비워 두어도 됩니다.</div>'
                '</div>',
                unsafe_allow_html=True,
            )
            _power_text_area(
                "특이사항 입력란 (선택)",
                key="power_notes",
                height=145,
                placeholder="특이사항이 있을 때 여기에 입력하세요.",
            )
            st.caption("작성 예: 축전지 2조 미측정 사유 · 접지선 보완 필요 · 다음 점검 시 확인사항")

            st.markdown("**📷 정밀점검 현장사진 (선택)**")
            power_photo_c1, power_photo_c2 = st.columns(2, gap="small")
            with power_photo_c1:
                power_camera_photo = st.camera_input(
                    "현장에서 바로 촬영",
                    key="power_camera_photo",
                )
            with power_photo_c2:
                power_uploaded_photos = st.file_uploader(
                    "앨범/파일에서 선택",
                    type=["jpg", "jpeg", "png", "webp"],
                    accept_multiple_files=True,
                    key="power_uploaded_photos",
                ) or []
            power_photos = _worklog_collect_photos(power_camera_photo, power_uploaded_photos)
            st.caption(
                f"선택 사진 {len(power_photos)}장 / 최대 {WORK_LOG_MAX_PHOTOS}장 · "
                "저장 시 자동 방향보정·압축 후 비공개 Drive에 저장됩니다."
            )

            payload_preview = _build_power_payload_from_state(final_confirmed=True)
            all_missing = _power_payload_missing_items(payload_preview)
            expected_count = max(_power_expected_item_count(payload_preview), 1)
            completion_rate = round(((expected_count - len(all_missing)) / expected_count) * 100, 1)

            metric1, metric2, metric3 = st.columns(3, gap="small")
            metric1.metric("입력 완료율", f"{completion_rate}%")
            metric2.metric("누락 측정항목", f"{len(all_missing)}개")
            cell_summary = f"1조 {payload_preview.get('battery1_cell_count', 0)}셀"
            if payload_preview.get("battery_group_count") == 2:
                cell_summary += f" · 2조 {payload_preview.get('battery2_cell_count', 0)}셀"
            metric3.metric("축전지 측정", cell_summary)

            summary_df = pd.DataFrame([
                {"구분": "점검자", "내용": payload_preview.get("worker") or "미입력"},
                {"구분": "주요 점검권역", "내용": payload_preview.get("major_area") or "미선택"},
                {"구분": "점검 국사", "내용": f"{payload_preview.get('mother')} / {payload_preview.get('local')}"},
                {"구분": "전원 구분", "내용": payload_preview.get("phase_type")},
                {"구분": "입력 방식", "내용": "기존값 불러오기 후 수정" if payload_preview.get("source_inspection_id") else "신규 입력"},
                {"구분": "특이사항", "내용": payload_preview.get("notes") or "없음"},
            ])
            st.dataframe(summary_df, use_container_width=True, hide_index=True)

            basic_missing = _power_basic_missing()
            if basic_missing:
                st.error("저장 필수 기본정보가 누락되었습니다: " + ", ".join(basic_missing))

            if all_missing:
                preview = ", ".join(all_missing[:12])
                more = f" 외 {len(all_missing) - 12}개" if len(all_missing) > 12 else ""
                st.markdown(
                    f'<div class="power-missing-box"><b>입력하지 않은 측정값이 있습니다.</b><br>'
                    f'{html.escape(preview)}{more}<br><br>'
                    '누락값은 Google Sheets에 공란으로 저장되며 누락항목도 함께 기록됩니다.</div>',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown('<div class="power-complete-box">✅ 모든 예정 측정항목이 입력되었습니다.</div>', unsafe_allow_html=True)

            final_confirmed = st.checkbox(
                "입력한 측정값과 누락항목을 모두 확인했으며, 현재 내용으로 최종 전송합니다.",
                key="power_final_confirmed",
            )
            st.markdown("**모든 측정값이 정확하게 입력되었는지 최종 확인한 후 ‘전송’을 눌러 주세요.**")
            submit_power = st.button(
                "📤 전송",
                key="power_final_submit",
                type="primary",
                use_container_width=True,
                disabled=bool(basic_missing) or not final_confirmed,
            )

            if submit_power:
                payload = _build_power_payload_from_state(final_confirmed=final_confirmed)
                photo_signature = []
                for photo in power_photos:
                    try:
                        photo_signature.append(hashlib.sha256(photo.getvalue()).hexdigest())
                    except Exception:
                        photo_signature.append(str(getattr(photo, "name", "photo")))
                payload_signature = hashlib.sha256(
                    json.dumps(
                        {"payload": payload, "photos": photo_signature},
                        ensure_ascii=False, sort_keys=True, default=str,
                    ).encode("utf-8")
                ).hexdigest()
                now_ts = time.time()
                last_signature = st.session_state.get("power_last_signature", "")
                last_saved_ts = float(st.session_state.get("power_last_saved_ts", 0) or 0)
                if payload_signature == last_signature and (now_ts - last_saved_ts) < 30:
                    st.warning("같은 측정값이 방금 저장되었습니다. 중복 전송을 방지했습니다.")
                else:
                    save_spinner_text = (
                        "Google Sheets와 Drive에 측정값·사진을 저장하고 있습니다..."
                        if power_photos else "Google Sheets에 측정값을 저장하고 있습니다..."
                    )
                    with st.spinner(save_spinner_text):
                        ok, message, inspection_id = save_power_inspection_result(payload, power_photos)
                    if ok:
                        st.session_state["power_last_signature"] = payload_signature
                        st.session_state["power_last_saved_ts"] = now_ts
                        st.success(f"✅ {message} · 점검ID: {inspection_id}")
                        time.sleep(2)
                        _reset_power_inspection()
                        st.rerun()
                    else:
                        st.error(f"❌ 저장 실패: {message}")

            st.markdown(
                f'<div class="power-sheet-note"><b>📊 Google Sheets 저장 구조</b><br>'
                f'스프레드시트: <b>{POWER_INSPECTION_SPREADSHEET_NAME}</b><br>'
                f'워크시트: <b>{POWER_INSPECTION_SHEET_NAME}</b><br>'
                '점검 1건은 새 행으로 추가됩니다. 사진은 비공개 Drive에 분리 저장되고 사진 ID/파일명만 시트에 연결 기록됩니다. 최근 측정값을 불러온 경우 원본 점검ID와 원본 저장일시도 함께 기록됩니다.</div>',
                unsafe_allow_html=True,
            )

        show_theme_complete_button = not (
            current_theme == "축전지 측정"
            and st.session_state.get("power_battery_exit_stage") == "ask_group2_complete"
        )
        if current_theme in POWER_THEME_ORDER[:-1] and show_theme_complete_button:
            theme_missing = _power_theme_missing(current_theme)
            st.markdown("---")
            if current_theme == "축전지 측정":
                selected_battery_group = 1 if _power_get("power_battery_set", "1조 셀 측정") == "1조 셀 측정" else 2
                if selected_battery_group == 1:
                    complete_title = "1조 셀 입력을 마쳤습니까?"
                    complete_description = "‘1조 셀 측정 완료’를 누르면 현재 값이 저장되고 2조 셀 측정 여부를 확인합니다."
                    complete_button_label = "✅ 1조 셀 측정 완료"
                else:
                    complete_title = "2조 셀 입력을 마쳤습니까?"
                    complete_description = "‘2조 셀 측정 완료’를 누르면 축전지 테마가 완료되고 다음 미완료 측정 메뉴로 이동합니다."
                    complete_button_label = "✅ 2조 셀 측정 완료"
            else:
                complete_title = "현재 테마 측정을 마쳤습니까?"
                complete_description = "‘측정 완료’를 누르면 현재 값이 임시저장되고 완료 상태로 표시됩니다."
                complete_button_label = "✅ 측정 완료"

            st.markdown(
                f'<div class="power-complete-box"><b>{html.escape(complete_title)}</b><br>'
                f'{html.escape(complete_description)} '
                f'시스템 확인 공란: {len(theme_missing)}개</div>',
                unsafe_allow_html=True,
            )
            st.button(
                complete_button_label,
                key=f"power_measurement_complete_{POWER_THEME_ORDER.index(current_theme)}",
                type="primary",
                use_container_width=True,
                on_click=_complete_current_power_theme,
            )

    st.info("측정값은 화면과 분리된 임시저장소에 즉시 보존됩니다. 기본정보·특이사항을 제외한 측정 입력은 모바일 숫자키패드를 사용하며, Enter/확인을 누르면 자동 소수점 적용 후 다음 입력칸으로 이동합니다.")


# --- [Tab 2: 법률 리스크/규정/계약 검토 & 감사보고서 작성] ---
with tab_doc:
    st.markdown("### 📄 법률 검토 · 감사보고서 작성/검증")

    if "api_key" not in st.session_state:
        st.warning("🔒 로그인 후 이용 가능합니다.")
    else:
        st.markdown("""
        <div class="audit-message-v2">
            <h4>🧭 AI 검토 품질 업그레이드 적용</h4>
            <p>최신 Gemini 모델 우선 선택, 검색 보강 옵션, 조항별 리스크 표, 수정문안, 감사보고서 품질검증 구조를 적용했습니다. 기존 법률 검토 기능은 그대로 유지됩니다.</p>
        </div>
        """, unsafe_allow_html=True)

        cur1, cur2 = st.tabs(["⚖️ 법률 리스크 심층 검토", "🔍 감사보고서 작성·검증"])

        with cur1:
            st.markdown("#### ⚖️ 법률 리스크 정밀 검토")
            st.caption("계약서·규정·공문·검토자료를 업로드하면 쟁점, 리스크, 근거 방향, 수정문안을 구조적으로 정리합니다.")

            uploaded_file = st.file_uploader("파일 업로드 (PDF, Word, TXT)", type=["txt", "pdf", "docx"], key="cur1_file")

            col_a, col_b = st.columns(2)
            with col_a:
                doc_type = st.selectbox(
                    "문서 유형",
                    ["계약서", "약관/일반조건", "사내 규정", "공문/통지문", "감사·조사 자료", "기타"],
                    index=0,
                    key="cur1_doc_type"
                )
                analysis_depth = st.selectbox(
                    "분석 수준",
                    ["핵심 요약", "리스크 식별(중점)", "조항/근거 중심(심층)", "수정문안 중심"],
                    index=2,
                    key="cur1_depth"
                )
            with col_b:
                company_position = st.selectbox(
                    "검토 관점",
                    ["우리 회사 입장", "갑/발주자 입장", "을/수급자 입장", "중립 검토"],
                    index=0,
                    key="cur1_position"
                )
                focus_area = st.selectbox(
                    "중점 분야",
                    ["전체 리스크", "계약대금/지급조건", "손해배상/면책", "하도급/공정거래", "개인정보/정보보호", "노무/안전보건", "부패방지/이해충돌"],
                    index=0,
                    key="cur1_focus"
                )

            use_search = st.toggle(
                "🔎 최신 법령·판례·가이드 검색 보강 사용",
                value=True,
                key="cur1_search",
                help="Gemini API의 Google Search Grounding을 우선 시도합니다. 패키지/계정에서 지원되지 않으면 일반 분석으로 대체됩니다."
            )

            if st.button("🚀 법률 리스크 분석 시작", use_container_width=True, key="cur1_run"):
                if not uploaded_file:
                    st.warning("⚠️ 먼저 파일을 업로드해주세요.")
                else:
                    content = read_file(uploaded_file)
                    if not content:
                        st.error("❌ 파일에서 텍스트를 추출하지 못했습니다. 스캔 PDF인 경우 OCR 처리 후 다시 업로드해 주세요.")
                    else:
                        with st.spinner("🧠 법률·컴플라이언스 관점에서 심층 분석 중입니다..."):
                            try:
                                prompt = build_legal_review_prompt(content, analysis_depth, doc_type, focus_area, company_position)
                                response, grounded, warning = generate_ai_response(prompt, task="legal", use_search=use_search, temperature=0.15)
                                st.success("✅ 분석 완료")
                                render_ai_response(response, grounded=grounded, warning=warning)
                            except Exception as e:
                                st.error(f"오류: {e}")

        with cur2:
            st.markdown("#### 🔍 감사보고서 작성·검증")
            st.caption("면담자료, 증거자료, 회사 규정, 기존 보고서를 바탕으로 감사보고서 초안 작성 또는 품질검증을 수행합니다.")

            mode = st.radio(
                "작업 모드",
                ["🧾 감사보고서 초안 생성", "✅ 감사보고서 검증·교정(오탈자/논리/형식)"],
                horizontal=True,
                key="cur2_mode"
            )
            is_draft_mode = "초안" in mode

            with st.expander("🔐 보안·주의사항", expanded=False):
                st.markdown(
                    "- 민감정보는 업로드 전 내부 보안 기준에 따라 비식별 처리하는 것이 안전합니다.\n"
                    "- 본 기능은 감사 판단을 보조하는 도구이며, 최종 판단·결재 책임은 감사실에 있습니다.\n"
                    "- 자료에 없는 사실은 생성하지 않도록 프롬프트에 제한을 두었습니다."
                )

            interview_audio = None
            interview_transcript = None
            evidence_files = []
            draft_text = ""
            draft_file = None

            if is_draft_mode:
                st.markdown("### ① 감사 자료 입력")
                cL, cR = st.columns(2)
                with cL:
                    interview_audio = st.file_uploader("🎧 면담 음성(mp3/wav/mp4) — 선택", type=["mp3", "wav", "mp4"], key="cur2_audio")
                    interview_transcript = st.file_uploader("📝 면담 녹취/메모(PDF/DOCX/TXT) — 권장", type=["txt", "pdf", "docx"], key="cur2_transcript")
                with cR:
                    evidence_files = st.file_uploader(
                        "📂 조사·증거/확인 자료 — 복수 업로드 가능",
                        type=["pdf", "png", "jpg", "jpeg", "xlsx", "xls", "csv", "txt", "docx"],
                        accept_multiple_files=True,
                        key="cur2_evidence"
                    ) or []
            else:
                st.markdown("### ① 검증 대상 보고서 입력")
                cL, cR = st.columns(2)
                with cL:
                    draft_text = st.text_area("검증할 감사보고서 붙여넣기", height=230, key="cur2_draft")
                with cR:
                    draft_file = st.file_uploader("또는 파일 업로드(PDF/DOCX/TXT)", type=["pdf", "docx", "txt"], key="cur2_draft_file")

            st.markdown("### ② 회사 규정/판단 기준 · ③ 표준 보고서 형식")
            left, right = st.columns(2)
            with left:
                regulations = st.file_uploader(
                    "📘 회사 규정/기준(인사규정·징계기준·윤리지침 등)",
                    type=["pdf", "docx", "txt"],
                    accept_multiple_files=True,
                    key="cur2_regs"
                ) or []
            with right:
                reference_reports = st.file_uploader(
                    "📑 참고 보고서 형식 — 선택",
                    type=["pdf", "docx", "txt"],
                    accept_multiple_files=True,
                    key="cur2_refs"
                ) or []

            st.markdown("### ④ 사건 개요 및 작성 옵션")
            row1, row2 = st.columns(2)
            with row1:
                case_title = st.text_input("사건명/건명", placeholder="예: 법인카드 사적 사용 의혹 조사", key="cur2_title")
            with row2:
                report_tone = st.selectbox(
                    "문서 톤",
                    ["감사보고서(공식·중립)", "보고서(간결·결정 중심)", "상신용(결재/조치 권고 중심)"],
                    index=0,
                    key="cur2_tone"
                )
            case_scope = st.text_area("사건 개요 요약 — 무엇을/언제/누가/어떤 경위로", height=120, key="cur2_scope")

            if st.button("🧠 감사보고서 AI 실행", use_container_width=True, key="cur2_run"):
                materials = []

                if is_draft_mode:
                    if interview_transcript:
                        transcript_text = read_file(interview_transcript)
                        if transcript_text:
                            materials.append(f"[면담 녹취/메모]\n{transcript_text}")
                    if evidence_files:
                        materials.append("[조사·증거 자료]\n" + read_multiple_files(evidence_files))
                    if interview_audio:
                        audio_file = process_media_file(interview_audio)
                        if audio_file:
                            materials.append("[면담 음성 파일]\nAI 업로드 파일이 함께 전달됩니다.")
                else:
                    if draft_text.strip():
                        materials.append("[검증 대상 보고서 - 붙여넣기]\n" + draft_text.strip())
                    if draft_file:
                        extracted = read_file(draft_file)
                        if extracted:
                            materials.append("[검증 대상 보고서 - 파일]\n" + extracted)

                regulations_text = read_multiple_files(regulations) if regulations else ""
                refs_text = read_multiple_files(reference_reports) if reference_reports else ""
                materials_text = "\n\n".join(materials).strip()

                if not case_title.strip():
                    st.warning("⚠️ 사건명/건명을 입력해 주세요.")
                elif not case_scope.strip() and not materials_text:
                    st.warning("⚠️ 사건 개요 또는 감사 자료 중 하나 이상은 입력해야 합니다.")
                else:
                    with st.spinner("📑 감사보고서 품질 기준에 맞춰 처리 중입니다..."):
                        try:
                            prompt = build_audit_report_prompt(mode, case_title, case_scope, report_tone, materials_text, regulations_text, refs_text)
                            if is_draft_mode and interview_audio:
                                # 음성 파일이 있는 경우 멀티모달 입력을 시도합니다.
                                audio_file = process_media_file(interview_audio)
                                if audio_file:
                                    response, grounded, warning = generate_ai_response([prompt, audio_file], task="report", use_search=False, temperature=0.12)
                                else:
                                    response, grounded, warning = generate_ai_response(prompt, task="report", use_search=False, temperature=0.12)
                            else:
                                response, grounded, warning = generate_ai_response(prompt, task="report", use_search=False, temperature=0.12)
                            st.success("✅ 처리 완료")
                            render_ai_response(response, grounded=grounded, warning=warning)
                        except Exception as e:
                            st.error(f"오류: {e}")

# --- [Tab 3: AI 에이전트] ---
with tab_chat:
    st.markdown("### 💬 AI 에이전트(챗봇)")
    if "api_key" not in st.session_state:
        st.warning("🔒 로그인 후 이용 가능합니다.")
    else:
        st.markdown("""
        <div class="audit-message-v2">
            <h4>🤝 감사·법률·컴플라이언스 전용 챗봇</h4>
            <p>질문을 그대로 전달하지 않고, 감사실 업무 기준에 맞춘 역할·답변 구조·한계 고지를 적용합니다. 최신 이슈는 검색 보강 모드를 사용할 수 있습니다.</p>
        </div>
        """, unsafe_allow_html=True)

        if "messages" not in st.session_state:
            st.session_state.messages = []

        chat_mode = st.selectbox(
            "상담 모드",
            ["감사·컴플라이언스 일반 상담", "최신 법령·판례·뉴스 검색 보강", "문서/보고서 문안 개선"],
            index=0,
            key="chat_mode"
        )
        use_chat_search = "최신" in chat_mode

        c1, c2 = st.columns([0.78, 0.22])
        with c1:
            with st.form(key="chat_input_form", clear_on_submit=True):
                user_input = st.text_input("질문 입력", placeholder="예: 계약서 지급조건 100일 조항의 리스크를 검토해줘")
                send_btn = st.form_submit_button("전송 📤", use_container_width=True)
        with c2:
            st.markdown("<div style='height:29px'></div>", unsafe_allow_html=True)
            if st.button("대화 초기화", use_container_width=True, key="chat_clear"):
                st.session_state.messages = []
                st.rerun()

        if send_btn and user_input:
            history_before = st.session_state.messages.copy()
            st.session_state.messages.append({"role": "user", "content": user_input})
            with st.spinner("답변 생성 중..."):
                try:
                    prompt = build_chat_prompt(user_input, history_before, chat_mode)
                    response, grounded, warning = generate_ai_response(prompt, task="chat", use_search=use_chat_search, temperature=0.2)
                    answer = _extract_response_text(response) or "응답을 생성하지 못했습니다."
                    sources = _extract_grounding_sources(response)
                    if grounded and sources:
                        src_text = "\n\n---\n**참고 출처**\n" + "\n".join([f"- [{s['title']}]({s['uri']})" for s in sources])
                        answer += src_text
                    elif warning:
                        answer += "\n\nℹ️ 검색 보강 호출이 실패하여 일반 AI 답변으로 대체되었습니다."
                    st.session_state.messages.append({"role": "assistant", "content": answer})
                except Exception as e:
                    st.error(f"오류: {e}")

        for msg in reversed(st.session_state.messages):
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

# --- [Tab 4: 스마트 요약] ---
with tab_summary:
    st.markdown("### 📰 스마트 요약")
    if "api_key" not in st.session_state:
        st.warning("🔒 로그인 후 이용 가능합니다.")
    else:
        st.markdown("""
        <div class="audit-message-v2">
            <h4>🧩 출처·리스크 중심 스마트 요약</h4>
            <p>뉴스, 웹페이지, 유튜브, 회의 음성, 텍스트를 업무 브리핑 형식으로 요약합니다. 필요 시 Google Search Grounding을 사용해 최신성도 보강합니다.</p>
        </div>
        """, unsafe_allow_html=True)

        st_type = st.radio("입력 방식", ["URL (유튜브/웹)", "미디어 파일", "텍스트"], horizontal=True)
        summary_mode = st.selectbox(
            "요약 모드",
            ["뉴스·보도자료 브리핑", "유튜브·교육영상 요약", "회의·면담 녹취 요약", "감사자료·증거자료 요약", "일반 요약"],
            index=0,
            key="summary_mode"
        )
        output_style = st.selectbox(
            "출력 방식",
            ["임원 보고용", "실무 체크리스트형", "교육자료 재가공용", "상세 분석형"],
            index=1,
            key="summary_style"
        )
        use_summary_search = st.toggle(
            "🔎 최신 검색 보강 사용(URL/뉴스 요약 권장)",
            value=("URL" in st_type),
            key="summary_search"
        )

        final_input = None
        is_multimodal = False
        source_hint = "직접 입력"

        if "URL" in st_type:
            url = st.text_input("URL 입력", placeholder="https:// 또는 유튜브 링크")
            source_hint = url or "URL"
            if url and "youtu" in url:
                with st.spinner("자막 추출 중..."):
                    final_input = get_youtube_transcript(url)
                    if not final_input:
                        st.info("자막을 찾지 못해 음성 분석을 시도합니다. 영상 길이와 권한에 따라 시간이 걸릴 수 있습니다.")
                        final_input = download_and_upload_youtube_audio(url)
                        is_multimodal = True if final_input else False
            elif url:
                with st.spinner("웹페이지 본문을 추출 중..."):
                    final_input = get_web_content(url)
                    if not final_input:
                        st.warning("웹페이지 본문을 직접 추출하지 못했습니다. 검색 보강 모드를 켠 상태로 URL 중심 요약을 시도할 수 있습니다.")
                        final_input = f"다음 URL의 공개 정보를 검색해 업무 브리핑 형식으로 요약하세요: {url}"

        elif "미디어" in st_type:
            mf = st.file_uploader("파일 업로드", type=["mp3", "wav", "mp4"])
            source_hint = getattr(mf, "name", "미디어 파일") if mf else "미디어 파일"
            if mf:
                final_input = process_media_file(mf)
                is_multimodal = True
        else:
            final_input = st.text_area("텍스트 입력", height=230)
            source_hint = "붙여넣은 텍스트"

        if st.button("⚡ 요약 실행", use_container_width=True):
            if final_input:
                with st.spinner("요약 중..."):
                    try:
                        if is_multimodal:
                            prompt = build_summary_prompt(summary_mode, output_style, source_hint, "첨부된 미디어 파일의 내용을 분석하세요.")
                            response, grounded, warning = generate_ai_response([prompt, final_input], task="summary", use_search=False, temperature=0.18)
                        else:
                            prompt = build_summary_prompt(summary_mode, output_style, source_hint, str(final_input))
                            response, grounded, warning = generate_ai_response(prompt, task="summary", use_search=use_summary_search, temperature=0.18)
                        st.success("✅ 요약 완료")
                        render_ai_response(response, grounded=grounded, warning=warning)
                    except Exception as e:
                        st.error(f"오류: {e}")
            else:
                st.warning("⚠️ 요약할 URL, 파일 또는 텍스트를 입력해 주세요.")

# --- [Tab 5: 관리자 대시보드 - 수동 로딩형 현황판] ---
with tab_admin:
    st.markdown("### 🔋 전원 정밀점검 데이터 관리")
    st.caption(
        "Google Sheets의 최신 누적 점검자료를 필요한 시점에 불러와 지역·국소·기간별로 확인하고 "
        "CSV 또는 Excel 파일로 다운로드합니다. 이 화면은 조회 버튼을 누를 때만 Google Sheets를 읽습니다."
    )

    st.markdown("""
    <div style="background:linear-gradient(135deg,#E0F2FE 0%,#ECFDF5 100%); border:1px solid #7DD3FC;
                border-radius:18px; padding:18px 20px; margin:8px 0 16px; box-shadow:0 8px 22px rgba(14,116,144,0.10);">
      <div style="font-size:1.08rem; font-weight:950; color:#0F172A; margin-bottom:6px;">지역 사용자 이용방법</div>
      <div style="color:#334155; font-weight:750; line-height:1.65;">
        ① 최신 누적 데이터를 불러옵니다. ② 담당 모국·국소와 기간을 선택합니다.<br>
        ③ 화면에서 측정값을 확인하거나 CSV·Excel로 내려받아 현장 및 사무실 업무에 활용합니다.
      </div>
    </div>
    """, unsafe_allow_html=True)

    if "power_admin_df" not in st.session_state:
        st.session_state["power_admin_df"] = None
    if "power_admin_loaded_at" not in st.session_state:
        st.session_state["power_admin_loaded_at"] = ""
    if "power_admin_error" not in st.session_state:
        st.session_state["power_admin_error"] = ""

    def _load_power_inspection_admin_df() -> pd.DataFrame:
        """사용자가 조회 버튼을 누른 시점의 최신 전원 정밀점검 누적자료를 읽습니다."""
        client = init_google_sheet_connection()
        if not client:
            raise RuntimeError("Google Sheets 연결 실패: Streamlit Secrets와 서비스 계정 권한을 확인하세요.")

        spreadsheet = client.open(POWER_INSPECTION_SPREADSHEET_NAME)
        try:
            worksheet = spreadsheet.worksheet(POWER_INSPECTION_SHEET_NAME)
        except Exception:
            return pd.DataFrame(columns=POWER_INSPECTION_HEADERS)

        values = worksheet.get_all_values()
        if not values:
            return pd.DataFrame(columns=POWER_INSPECTION_HEADERS)

        headers = [str(value).strip() for value in values[0]]
        rows = values[1:]
        normalized_rows = [
            [row[index] if index < len(row) else "" for index in range(len(headers))]
            for row in rows
        ]
        return pd.DataFrame(normalized_rows, columns=headers).fillna("")

    load_col, reset_col, status_col = st.columns([0.27, 0.20, 0.53], vertical_alignment="center")
    with load_col:
        load_power_data = st.button(
            "🔄 최신 누적 데이터 불러오기",
            type="primary",
            use_container_width=True,
            key="power_admin_load_latest",
        )
    with reset_col:
        clear_power_data = st.button(
            "🧹 조회화면 초기화",
            use_container_width=True,
            key="power_admin_clear_data",
        )
    with status_col:
        if st.session_state.get("power_admin_loaded_at"):
            st.caption(
                f"마지막 조회: {st.session_state['power_admin_loaded_at']} · "
                "최신 자료가 필요하면 다시 불러오기를 누르세요."
            )
        else:
            st.caption("아직 Google Sheets 데이터를 불러오지 않았습니다.")

    if clear_power_data:
        st.session_state["power_admin_df"] = None
        st.session_state["power_admin_loaded_at"] = ""
        st.session_state["power_admin_error"] = ""
        st.rerun()

    if load_power_data:
        with st.spinner("Google Sheets에서 전원 정밀점검 최신 누적자료를 불러오는 중입니다..."):
            try:
                st.session_state["power_admin_df"] = _load_power_inspection_admin_df()
                st.session_state["power_admin_loaded_at"] = _korea_now().strftime("%Y-%m-%d %H:%M:%S")
                st.session_state["power_admin_error"] = ""
            except Exception as error:
                st.session_state["power_admin_df"] = None
                st.session_state["power_admin_error"] = str(error)

    if st.session_state.get("power_admin_error"):
        st.error(f"전원 정밀점검 데이터를 불러오지 못했습니다: {st.session_state['power_admin_error']}")

    power_admin_df = st.session_state.get("power_admin_df")

    if power_admin_df is None:
        st.info("업무에 최신 측정자료가 필요할 때 위의 ‘최신 누적 데이터 불러오기’를 눌러 주세요.")
    elif power_admin_df.empty:
        st.warning("Google Sheets에 저장된 전원 정밀점검 자료가 아직 없습니다.")
    else:
        source_df = power_admin_df.copy().fillna("")

        for required_column in POWER_INSPECTION_HEADERS:
            if required_column not in source_df.columns:
                source_df[required_column] = ""

        source_df["_저장일시_dt"] = pd.to_datetime(source_df["저장일시"], errors="coerce")
        source_df = source_df.sort_values("_저장일시_dt", ascending=False, na_position="last")

        st.markdown("#### 🔎 담당 지역 및 조회범위 선택")
        filter_col1, filter_col2, filter_col3 = st.columns(3)

        mother_values = sorted(
            value for value in source_df["모국"].astype(str).str.strip().unique().tolist() if value
        )
        with filter_col1:
            selected_mother = st.selectbox(
                "관리 모국",
                ["전체 모국"] + mother_values,
                key="power_admin_filter_mother",
            )

        if selected_mother == "전체 모국":
            locality_source = source_df
        else:
            locality_source = source_df[source_df["모국"].astype(str).str.strip() == selected_mother]

        locality_values = sorted(
            value for value in locality_source["국소"].astype(str).str.strip().unique().tolist() if value
        )
        with filter_col2:
            selected_local = st.selectbox(
                "관리 국소",
                ["전체 국소"] + locality_values,
                key="power_admin_filter_local",
            )

        period_options = {
            "최근 1개월": 31,
            "최근 6개월": 183,
            "최근 1년": 365,
            "최근 2년": 730,
            "전체 기간": None,
        }
        with filter_col3:
            selected_period = st.selectbox(
                "조회 기간",
                list(period_options.keys()),
                index=2,
                key="power_admin_filter_period",
            )

        search_term = st.text_input(
            "측정자료 검색",
            placeholder="점검자, 운용조, 모국, 국소, 특이사항 등",
            key="power_admin_search_term",
        ).strip()

        filtered_df = source_df.copy()
        if selected_mother != "전체 모국":
            filtered_df = filtered_df[
                filtered_df["모국"].astype(str).str.strip() == selected_mother
            ]
        if selected_local != "전체 국소":
            filtered_df = filtered_df[
                filtered_df["국소"].astype(str).str.strip() == selected_local
            ]

        period_days = period_options[selected_period]
        if period_days is not None:
            cutoff = _korea_now().replace(tzinfo=None) - datetime.timedelta(days=period_days)
            filtered_df = filtered_df[
                filtered_df["_저장일시_dt"].notna()
                & (filtered_df["_저장일시_dt"] >= cutoff)
            ]

        if search_term:
            searchable_columns = [column for column in filtered_df.columns if column != "_저장일시_dt"]
            search_mask = filtered_df[searchable_columns].apply(
                lambda row: row.astype(str).str.contains(search_term, case=False, na=False).any(),
                axis=1,
            )
            filtered_df = filtered_df[search_mask]

        total_records = int(len(filtered_df))
        unique_mothers = int(filtered_df["모국"].astype(str).str.strip().replace("", pd.NA).dropna().nunique())
        unique_locals = int(filtered_df["국소"].astype(str).str.strip().replace("", pd.NA).dropna().nunique())
        unique_inspectors = int(filtered_df["점검자"].astype(str).str.strip().replace("", pd.NA).dropna().nunique())
        latest_saved_at = "-"
        if not filtered_df.empty and filtered_df["_저장일시_dt"].notna().any():
            latest_saved_at = filtered_df["_저장일시_dt"].max().strftime("%Y-%m-%d %H:%M:%S")

        metric1, metric2, metric3, metric4, metric5 = st.columns(5)
        metric1.metric("조회 건수", f"{total_records:,}건")
        metric2.metric("모국", f"{unique_mothers:,}개")
        metric3.metric("국소", f"{unique_locals:,}개")
        metric4.metric("점검자", f"{unique_inspectors:,}명")
        metric5.metric("최근 측정일시", latest_saved_at)

        display_df = filtered_df.drop(columns=["_저장일시_dt"], errors="ignore")

        st.markdown("#### 📋 전원 정밀점검 측정자료")
        st.caption(
            "화면에는 선택한 모국·국소·기간 조건의 자료만 표시됩니다. "
            "CSV와 Excel 다운로드에도 동일한 필터가 적용됩니다."
        )
        st.dataframe(display_df, use_container_width=True, hide_index=True, height=520)

        photo_records_df = filtered_df[
            filtered_df.get("사진파일ID목록", pd.Series(index=filtered_df.index, dtype=str)).astype(str).str.strip() != ""
        ].copy()
        if not photo_records_df.empty:
            with st.expander("📷 정밀점검 현장사진 조회·다운로드", expanded=False):
                photo_record_indices = photo_records_df.index.tolist()

                def _power_admin_photo_label(row_index):
                    row = photo_records_df.loc[row_index]
                    count = len([v for v in str(row.get("사진파일ID목록", "") or "").split("|") if v.strip()])
                    return (
                        f"{row.get('저장일시','')} · {row.get('모국','')} / {row.get('국소','')} · "
                        f"{row.get('점검자','')} · 사진 {count}장"
                    )

                selected_photo_row_index = st.selectbox(
                    "사진을 확인할 정밀점검 기록",
                    options=photo_record_indices,
                    format_func=_power_admin_photo_label,
                    key="power_admin_photo_record",
                )
                selected_photo_record = photo_records_df.loc[selected_photo_row_index]
                _render_power_photo_download(
                    selected_photo_record,
                    key_prefix=f"power_admin_{selected_photo_record.get('점검ID', selected_photo_row_index)}",
                )

        summary_df = pd.DataFrame(columns=["모국", "국소", "점검건수", "최근측정일시", "최근점검자"])
        if not filtered_df.empty:
            summary_source = filtered_df.copy()
            summary_source["점검자"] = summary_source["점검자"].astype(str)
            summary_rows = []
            for (mother_name, local_name), group in summary_source.groupby(["모국", "국소"], dropna=False):
                ordered = group.sort_values("_저장일시_dt", ascending=False, na_position="last")
                latest_row = ordered.iloc[0]
                latest_dt = latest_row.get("_저장일시_dt")
                latest_text = latest_dt.strftime("%Y-%m-%d %H:%M:%S") if pd.notna(latest_dt) else str(latest_row.get("저장일시", ""))
                summary_rows.append({
                    "모국": str(mother_name),
                    "국소": str(local_name),
                    "점검건수": int(len(group)),
                    "최근측정일시": latest_text,
                    "최근점검자": str(latest_row.get("점검자", "")),
                })
            summary_df = pd.DataFrame(summary_rows).sort_values(["모국", "국소"])

        safe_mother = "전체" if selected_mother == "전체 모국" else re.sub(r"[^0-9A-Za-z가-힣_-]", "_", selected_mother)
        safe_local = "전체" if selected_local == "전체 국소" else re.sub(r"[^0-9A-Za-z가-힣_-]", "_", selected_local)
        exported_at = _korea_now().strftime("%Y%m%d_%H%M%S")
        file_base = f"전원정밀점검_{safe_mother}_{safe_local}_{exported_at}"

        download_col1, download_col2 = st.columns(2)
        with download_col1:
            csv_bytes = display_df.to_csv(index=False).encode("utf-8-sig")
            st.download_button(
                "📥 현재 조회자료 CSV 다운로드",
                data=csv_bytes,
                file_name=f"{file_base}.csv",
                mime="text/csv",
                use_container_width=True,
                key="power_admin_download_csv",
                disabled=display_df.empty,
            )

        with download_col2:
            try:
                from io import BytesIO

                excel_output = BytesIO()
                with pd.ExcelWriter(excel_output, engine="openpyxl") as writer:
                    display_df.to_excel(writer, index=False, sheet_name="전원정밀점검_자료")
                    summary_df.to_excel(writer, index=False, sheet_name="국소별_요약")
                st.download_button(
                    "📥 현재 조회자료 Excel 다운로드",
                    data=excel_output.getvalue(),
                    file_name=f"{file_base}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="power_admin_download_excel",
                    disabled=display_df.empty,
                )
            except Exception as excel_error:
                st.warning(f"Excel 파일을 생성하지 못했습니다. CSV 다운로드를 이용해 주세요. ({excel_error})")

        with st.expander("📌 데이터 이용 및 운영 안내", expanded=False):
            st.markdown(
                "- 이 화면은 **최신 누적 데이터 불러오기**를 누른 시점의 Google Sheets 자료를 사용합니다.\n"
                "- 실시간 최신자료가 필요하면 조회 버튼을 다시 눌러 새로 읽어오면 됩니다.\n"
                "- 지역 사용자는 담당 모국과 국소를 선택한 뒤 화면 확인 또는 파일 다운로드를 이용할 수 있습니다.\n"
                "- CSV는 범용 공유용, Excel은 원본자료와 국소별 요약을 함께 제공하는 업무용 형식입니다.\n"
                "- 지역별 열람권한을 기술적으로 제한해야 하는 경우에는 사용자 계정과 담당 모국을 연결하는 별도 권한 설정이 필요합니다."
            )
